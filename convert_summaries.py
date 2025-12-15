#!/usr/bin/env python3
"""
Convert earnings transcript summaries to Word and PDF formats
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess


def create_word_document(input_file: str, output_file: str):
    """
    Convert a text file to a formatted Word document
    """
    # Read the text file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Split content into lines
    lines = content.split('\n')

    # Process each line
    for line in lines:
        line = line.rstrip()

        # Empty line
        if not line:
            doc.add_paragraph()
            continue

        # Detect headers (lines with === or ---)
        if '=' * 20 in line or '-' * 20 in line:
            doc.add_paragraph()
            continue

        # Main title
        if (line.isupper() and len(line) > 10) or \
                line.startswith('EARNINGS TRANSCRIPT') or \
                line.startswith('SUMMARY BY') or \
                line.startswith("CLAUDE'S") or \
                line.startswith("CHATGPT'S"):
            p = doc.add_paragraph(line)
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 128)
            continue

        # Section headers
        if (line and line[0].isdigit() and '. ' in line[:5]) or \
                (line.startswith('Company:') or line.startswith('Industry:') or
                 line.startswith('Sector:') or line.startswith('Analysis Date:')):
            p = doc.add_paragraph(line)
            if line[0].isdigit():
                p.style = 'Heading 2'
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            else:
                run = p.runs[0]
                run.font.bold = True
            continue

        # Bullet points
        if line.strip().startswith(('•', '-', '*', '○')):
            p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
            continue

        # Regular paragraph
        doc.add_paragraph(line)

    # Save the document
    doc.save(output_file)
    print(f"   ✓ Created Word: {os.path.basename(output_file)}")


def convert_to_pdf(word_file: str, pdf_file: str):
    """
    Convert Word document to PDF
    """
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
             os.path.dirname(pdf_file), word_file],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode == 0:
            print(f"   ✓ Created PDF: {os.path.basename(pdf_file)}")
            return True
        else:
            print(f"   ⚠️  LibreOffice not available")
            return False
    except FileNotFoundError:
        print(f"   ⚠️  LibreOffice not installed - PDF conversion skipped")
        print(f"   💡 Open .docx in Word and use 'Save As PDF'")
        return False
    except Exception as e:
        print(f"   ⚠️  PDF conversion failed: {e}")
        return False


def main():
    INPUT_DIR = "./output"
    OUTPUT_DIR = "./output"

    print("\n" + "=" * 80)
    print("📄 EARNINGS SUMMARY CONVERTER - TXT to DOCX & PDF")
    print("=" * 80 + "\n")

    input_path = Path(INPUT_DIR)
    txt_files = list(input_path.glob("*_summary.txt")) + list(input_path.glob("*_comparison.txt"))

    if not txt_files:
        print(f"❌ No summary files found in {INPUT_DIR}")
        sys.exit(1)

    print(f"Found {len(txt_files)} file(s) to convert:\n")

    converted_word = 0
    converted_pdf = 0

    for txt_file in txt_files:
        print(f"📝 Converting: {txt_file.name}")

        base_name = txt_file.stem
        word_file = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
        pdf_file = os.path.join(OUTPUT_DIR, f"{base_name}.pdf")

        try:
            create_word_document(str(txt_file), word_file)
            converted_word += 1

            if convert_to_pdf(word_file, pdf_file):
                converted_pdf += 1

            print()

        except Exception as e:
            print(f"   ❌ Error: {e}\n")

    print("=" * 80)
    print("✅ CONVERSION COMPLETE")
    print("=" * 80)
    print(f"📊 Word documents created: {converted_word}/{len(txt_files)}")
    print(f"📊 PDF documents created: {converted_pdf}/{len(txt_files)}")
    print(f"\n📁 Output location: {os.path.abspath(OUTPUT_DIR)}\n")


if __name__ == "__main__":
    main()