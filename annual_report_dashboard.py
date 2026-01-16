"""
Annual Report Analysis Dashboard

Streamlit app for analyzing company 10-K annual reports with AI.
Deployable on share.streamlit.io
"""

import streamlit as st
import json
import io
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime
from dotenv import load_dotenv

# Document generation libraries
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import the analyzer
from annual_report_analysis import AnnualReportAnalyzer

# Load environment variables
load_dotenv()

# Logo path
LOGO_PATH = "company_logo.png"

# Email configuration
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

# Page configuration
st.set_page_config(
    page_title="Annual Report Analysis",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    .stExpander {
        border: 1px solid #e0e0e0;
        border-radius: 5px;
        margin-bottom: 1rem;
    }
    .section-header {
        background-color: #1f77b4;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .risk-high {
        color: #d62728;
        font-weight: bold;
    }
    .risk-medium {
        color: #ff7f0e;
    }
    .risk-low {
        color: #2ca02c;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 15px;
        border-radius: 8px;
        border-left: 4px solid #1f77b4;
    }
</style>
""", unsafe_allow_html=True)


def init_analyzer():
    """Initialize the analyzer with caching."""
    if 'analyzer' not in st.session_state:
        st.session_state.analyzer = AnnualReportAnalyzer(enable_ai=True)
    return st.session_state.analyzer


def generate_pdf_report(results: dict) -> bytes:
    """Generate a PDF report with company logo."""
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, spaceAfter=20, alignment=TA_CENTER)
    h1_style = ParagraphStyle('Heading1', parent=styles['Heading1'], fontSize=14, spaceAfter=10, spaceBefore=15)
    h2_style = ParagraphStyle('Heading2', parent=styles['Heading2'], fontSize=12, spaceAfter=8, spaceBefore=12)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, spaceAfter=6, leading=14)
    center_style = ParagraphStyle('Center', parent=styles['Normal'], alignment=TA_CENTER)

    elements = []

    # Add logo if exists
    if os.path.exists(LOGO_PATH):
        logo = Image(LOGO_PATH, width=2.5*inch, height=1*inch)
        logo.hAlign = 'CENTER'
        elements.append(logo)
        elements.append(Spacer(1, 20))

    symbol = results.get('symbol', 'N/A')
    years = results.get('summary', {}).get('years_covered', [])

    # Title
    elements.append(Paragraph(f'{symbol} Annual Report Analysis', title_style))
    elements.append(Paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', center_style))
    if years:
        elements.append(Paragraph(f'Reports Analyzed: FY{years[0]} - FY{years[-1]}', center_style))
    elements.append(Spacer(1, 20))

    # Executive Summary
    exec_summary = results.get('executive_summary', '')
    if exec_summary:
        elements.append(Paragraph('EXECUTIVE SUMMARY', h1_style))
        # Clean markdown
        exec_summary = exec_summary.replace('**', '').replace('##', '').replace('#', '').replace('*', '')
        for para in exec_summary.split('\n\n'):
            if para.strip():
                # Escape special XML characters
                clean_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                try:
                    elements.append(Paragraph(clean_para, body_style))
                    elements.append(Spacer(1, 6))
                except:
                    pass

    elements.append(PageBreak())

    # Year-over-Year Comparisons
    yoy_comparisons = results.get('yoy_comparisons', [])
    if yoy_comparisons:
        elements.append(Paragraph('YEAR-OVER-YEAR COMPARISONS', h1_style))
        for comp in yoy_comparisons:
            elements.append(Paragraph(comp.get('years', ''), h2_style))
            analysis = comp.get('analysis', '').replace('**', '').replace('##', '').replace('#', '').replace('*', '')
            for para in analysis.split('\n\n'):
                if para.strip():
                    clean_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')[:2000]
                    try:
                        elements.append(Paragraph(clean_para, body_style))
                        elements.append(Spacer(1, 4))
                    except:
                        pass

    elements.append(PageBreak())

    # Individual Reports
    elements.append(Paragraph('INDIVIDUAL REPORT ANALYSES', h1_style))
    for report in results.get('reports', []):
        fy = report.get('fiscal_year', 'N/A')
        elements.append(Paragraph(f'FY{fy} (Filed: {report.get("filing_date", "N/A")})', h2_style))

        ai = report.get('ai_analysis', {})
        if ai.get('mda_analysis'):
            elements.append(Paragraph('MD&A Analysis:', ParagraphStyle('Bold', parent=body_style, fontName='Helvetica-Bold')))
            mda = ai['mda_analysis'].replace('**', '').replace('##', '').replace('#', '').replace('*', '')
            for para in mda.split('\n\n')[:5]:
                if para.strip():
                    clean_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')[:1500]
                    try:
                        elements.append(Paragraph(clean_para, body_style))
                    except:
                        pass
        elements.append(Spacer(1, 15))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word_report(results: dict) -> bytes:
    """Generate a Word document with company logo."""
    doc = Document()

    symbol = results.get('symbol', 'N/A')
    years = results.get('summary', {}).get('years_covered', [])

    # Add logo if exists
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.5))
        doc.add_paragraph()

    # Title
    title = doc.add_heading(f'{symbol} Annual Report Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    if years:
        meta.add_run(f'\nReports Analyzed: FY{years[0]} - FY{years[-1]}')
    meta.add_run(f'\nAI Analysis: Enabled')

    doc.add_paragraph()

    # Executive Summary
    exec_summary = results.get('executive_summary', '')
    if exec_summary:
        doc.add_heading('Executive Summary', level=1)
        exec_summary = exec_summary.replace('**', '').replace('##', '').replace('# ', '')
        for para in exec_summary.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # Year-over-Year Comparisons
    yoy_comparisons = results.get('yoy_comparisons', [])
    if yoy_comparisons:
        doc.add_heading('Year-over-Year Comparisons', level=1)
        for comp in yoy_comparisons:
            doc.add_heading(comp.get('years', ''), level=2)
            analysis = comp.get('analysis', '').replace('**', '').replace('##', '').replace('# ', '')
            for para in analysis.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # Individual Reports
    doc.add_heading('Individual Report Analyses', level=1)
    for report in results.get('reports', []):
        fy = report.get('fiscal_year', 'N/A')
        doc.add_heading(f'FY{fy} (Filed: {report.get("filing_date", "N/A")})', level=2)

        ai = report.get('ai_analysis', {})

        if ai.get('mda_analysis'):
            doc.add_heading('MD&A Analysis', level=3)
            mda = ai['mda_analysis'].replace('**', '').replace('##', '').replace('# ', '')
            for para in mda.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.space_after = Pt(6)

        if ai.get('business_summary') and 'I notice' not in ai.get('business_summary', ''):
            doc.add_heading('Business Analysis', level=3)
            biz = ai['business_summary'].replace('**', '').replace('##', '').replace('# ', '')
            for para in biz.split('\n\n')[:3]:
                if para.strip():
                    p = doc.add_paragraph(para.strip())

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def send_email_with_reports(symbol: str, pdf_data: bytes, word_data: bytes,
                           recipient_email: str = None) -> bool:
    """
    Send email with PDF and Word report attachments.

    Args:
        symbol: Stock ticker symbol
        pdf_data: PDF file bytes
        word_data: Word file bytes
        recipient_email: Email to send to (defaults to EMAIL_ADDRESS from .env)

    Returns:
        True if email sent successfully, False otherwise
    """
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        raise ValueError("Email credentials not configured in .env file")

    recipient = recipient_email or EMAIL_ADDRESS

    # Create message
    msg = MIMEMultipart()
    msg['From'] = EMAIL_ADDRESS
    msg['To'] = recipient
    msg['Subject'] = f'{symbol} Annual Report Analysis - {datetime.now().strftime("%Y-%m-%d")}'

    # Email body
    body = f"""
Annual Report Analysis for {symbol}

Please find attached the analysis reports:
- {symbol}_Annual_Report_Analysis.pdf
- {symbol}_Annual_Report_Analysis.docx

This analysis includes:
- Executive Summary
- Year-over-Year Comparisons
- Individual Report Analyses (Business, Risks, MD&A)

Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}

---
Automated report from Annual Report Analysis Dashboard
    """
    msg.attach(MIMEText(body, 'plain'))

    # Attach PDF
    pdf_attachment = MIMEBase('application', 'pdf')
    pdf_attachment.set_payload(pdf_data)
    encoders.encode_base64(pdf_attachment)
    pdf_attachment.add_header('Content-Disposition',
                              f'attachment; filename="{symbol}_Annual_Report_Analysis.pdf"')
    msg.attach(pdf_attachment)

    # Attach Word document
    word_attachment = MIMEBase('application',
                               'vnd.openxmlformats-officedocument.wordprocessingml.document')
    word_attachment.set_payload(word_data)
    encoders.encode_base64(word_attachment)
    word_attachment.add_header('Content-Disposition',
                               f'attachment; filename="{symbol}_Annual_Report_Analysis.docx"')
    msg.attach(word_attachment)

    # Send email
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        raise Exception(f"Failed to send email: {str(e)}")


@st.cache_data(ttl=3600, show_spinner=False)
def analyze_company_cached(symbol: str, num_reports: int, run_ai: bool, run_multi_agent: bool = True):
    """Cached analysis to avoid re-running for same parameters."""
    analyzer = AnnualReportAnalyzer(enable_ai=run_ai)
    return analyzer.analyze_company(symbol, num_reports, run_ai_analysis=run_ai, run_multi_agent=run_multi_agent)


def display_executive_summary(summary: str):
    """Display the executive summary in a formatted way."""
    if not summary:
        st.warning("No executive summary available.")
        return

    st.markdown("### Executive Summary")
    st.markdown(summary)


def display_report_analysis(report_data: dict, report_idx: int):
    """Display analysis for a single annual report."""
    fiscal_year = report_data.get('fiscal_year', 'Unknown')
    filing_date = report_data.get('filing_date', 'Unknown')

    with st.expander(f"📄 FY{fiscal_year} Annual Report (Filed: {filing_date})", expanded=(report_idx == 0)):
        if report_data.get('error'):
            st.error(f"Error: {report_data['error']}")
            return

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Sections Extracted", len(report_data.get('sections_extracted', [])))
        with col2:
            st.metric("Content Length", f"{report_data.get('content_length', 0):,} chars")
        with col3:
            st.metric("Filing URL", "📎")
            if report_data.get('filing_url'):
                st.markdown(f"[View SEC Filing]({report_data['filing_url']})")

        # AI Analysis sections
        ai_analysis = report_data.get('ai_analysis', {})

        if ai_analysis.get('business_summary'):
            st.markdown("---")
            st.markdown("#### 🏢 Business Analysis")
            st.markdown(ai_analysis['business_summary'])

        if ai_analysis.get('risk_analysis'):
            st.markdown("---")
            st.markdown("#### ⚠️ Risk Analysis")
            st.markdown(ai_analysis['risk_analysis'])

        if ai_analysis.get('mda_analysis'):
            st.markdown("---")
            st.markdown("#### 📈 MD&A Analysis")
            st.markdown(ai_analysis['mda_analysis'])

        # Key metrics if available
        if report_data.get('key_metrics'):
            st.markdown("---")
            st.markdown("#### 📊 Key Metrics Mentioned")
            for key, value in report_data['key_metrics'].items():
                st.write(f"- **{key.replace('_', ' ').title()}**: {value}")

        # Top risks extracted
        if report_data.get('top_risks'):
            st.markdown("---")
            st.markdown("#### 🎯 Top Risk Factors")
            for i, risk in enumerate(report_data['top_risks'], 1):
                st.write(f"{i}. {risk}")


def display_yoy_comparisons(comparisons: list):
    """Display year-over-year comparisons."""
    if not comparisons:
        return

    st.markdown("### 📊 Year-over-Year Comparisons")

    for comp in comparisons:
        years = comp.get('years', 'Unknown')
        analysis = comp.get('analysis', '')

        with st.expander(f"🔄 {years}", expanded=True):
            st.markdown(analysis)


def display_agent_analysis(agent_analysis: dict):
    """Display multi-agent analysis results."""
    if not agent_analysis:
        return

    st.markdown("### 🤖 Multi-Agent Analysis")
    st.markdown("*10 specialized AI agents analyzed the most recent annual report*")

    # Define agent display order and groupings
    agent_groups = {
        "Business & Strategy": ["business_model_analyst", "strategy_analyst", "competitive_position_analyst"],
        "Financial Analysis": ["financial_health_analyst", "debt_credit_analyst"],
        "Risk & Governance": ["risk_deep_dive_analyst", "red_flag_analyst", "management_governance_analyst"],
        "ESG & Investment": ["esg_analyst", "investment_strategist"]
    }

    # Create tabs for agent groups
    tab_names = list(agent_groups.keys())
    tabs = st.tabs(tab_names)

    for tab, (group_name, agent_ids) in zip(tabs, agent_groups.items()):
        with tab:
            for agent_id in agent_ids:
                if agent_id in agent_analysis:
                    agent_data = agent_analysis[agent_id]
                    emoji = agent_data.get('emoji', '🤖')
                    name = agent_data.get('agent_name', agent_id)
                    analysis = agent_data.get('analysis', '')
                    error = agent_data.get('error', '')

                    with st.expander(f"{emoji} {name}", expanded=False):
                        if error:
                            st.error(f"Error: {error}")
                        elif analysis:
                            st.markdown(analysis)
                        else:
                            st.info("No analysis available from this agent.")

    # Summary metrics for agents
    st.markdown("---")
    st.markdown("##### Agent Summary")

    col1, col2, col3 = st.columns(3)

    # Count successful analyses
    successful = sum(1 for a in agent_analysis.values() if a.get('analysis') and not a.get('error'))
    errors = sum(1 for a in agent_analysis.values() if a.get('error'))
    no_content = sum(1 for a in agent_analysis.values() if 'No relevant content' in a.get('analysis', ''))

    with col1:
        st.metric("Agents Run", len(agent_analysis))
    with col2:
        st.metric("Successful", successful)
    with col3:
        st.metric("Data Gaps", no_content)


def main():
    """Main Streamlit app."""
    st.title("📊 Annual Report Analysis")
    st.markdown("*AI-powered analysis of SEC 10-K filings*")

    # Sidebar
    with st.sidebar:
        st.header("Settings")

        symbol = st.text_input(
            "Stock Ticker",
            value="AAPL",
            max_chars=10,
            help="Enter a stock ticker symbol (e.g., AAPL, MSFT, GOOGL)"
        ).upper().strip()

        num_reports = st.slider(
            "Number of Reports",
            min_value=1,
            max_value=5,
            value=3,
            help="Number of annual reports to analyze"
        )

        run_ai = st.checkbox(
            "Enable AI Analysis",
            value=True,
            help="Use AI to analyze report sections"
        )

        run_multi_agent = st.checkbox(
            "Enable Multi-Agent Analysis",
            value=True,
            help="Run 10 specialized AI agents on the most recent report (Business, Risk, Financial, Strategy, ESG, etc.)"
        )

        st.markdown("---")
        st.markdown("##### Auto-Email")

        auto_email = st.checkbox(
            "Email reports automatically",
            value=False,
            help="Automatically email PDF and Word reports after analysis"
        )

        if auto_email:
            auto_email_address = st.text_input(
                "Email to",
                value=EMAIL_ADDRESS or "",
                key="auto_email_address"
            )
        else:
            auto_email_address = EMAIL_ADDRESS

        analyze_button = st.button(
            "🔍 Analyze Reports",
            type="primary",
            use_container_width=True
        )

        st.markdown("---")
        st.markdown("### About")
        st.markdown("""
        This tool fetches and analyzes annual reports (10-K filings) from SEC EDGAR.

        **Features:**
        - Business overview & changes
        - Risk factor analysis
        - MD&A insights
        - Year-over-year comparisons
        - Executive summary
        - **Multi-Agent Analysis** (10 specialized AI agents)

        **AI Agents:**
        - Business Model Analyst
        - Financial Health Analyst
        - Risk Deep Dive Analyst
        - Competitive Position Analyst
        - Strategy Analyst
        - Management & Governance Analyst
        - Debt & Credit Analyst
        - ESG Analyst
        - Red Flag Analyst
        - Investment Strategist

        **AI Models Used:**
        - Claude (Anthropic)
        - GPT-4 (OpenAI fallback)
        """)

    # Main content
    if analyze_button and symbol:
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()

        status_text.text(f"Fetching annual reports for {symbol}...")
        progress_bar.progress(10)

        try:
            # Run analysis
            time_estimate = "2-4 minutes" if run_multi_agent else "1-2 minutes"
            with st.spinner(f"Analyzing {symbol}... This may take {time_estimate}."):
                results = analyze_company_cached(symbol, num_reports, run_ai, run_multi_agent)

            progress_bar.progress(100)
            status_text.empty()

            # Check for errors
            if results.get('errors'):
                for error in results['errors']:
                    st.error(error)
                return

            # Display results
            st.success(f"✅ Analysis complete for {symbol}")

            # Summary metrics
            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.metric("Symbol", results.get('symbol', 'N/A'))
            with col2:
                st.metric("Reports Analyzed", results.get('summary', {}).get('total_reports', 0))
            with col3:
                years = results.get('summary', {}).get('years_covered', [])
                st.metric("Years Covered", f"FY{years[0] if years else 'N/A'} - FY{years[-1] if years else 'N/A'}")
            with col4:
                ai_enabled = results.get('summary', {}).get('ai_analysis_enabled', False)
                st.metric("AI Analysis", "Enabled" if ai_enabled else "Disabled")
            with col5:
                agents_count = results.get('summary', {}).get('agents_count', 0)
                st.metric("Agents Run", agents_count)

            st.markdown("---")

            # Executive Summary (most important - show first)
            if results.get('executive_summary'):
                display_executive_summary(results['executive_summary'])
                st.markdown("---")

            # Year-over-Year Comparisons
            if results.get('yoy_comparisons'):
                display_yoy_comparisons(results['yoy_comparisons'])
                st.markdown("---")

            # Multi-Agent Analysis
            if results.get('agent_analysis'):
                display_agent_analysis(results['agent_analysis'])
                st.markdown("---")

            # Individual Report Analysis
            st.markdown("### 📁 Individual Report Analysis")
            reports = results.get('reports', [])
            for idx, report_data in enumerate(reports):
                display_report_analysis(report_data, idx)

            # Export options
            st.markdown("---")
            st.markdown("### 💾 Export Reports")

            # Generate documents once
            pdf_data = None
            word_data = None

            try:
                pdf_data = generate_pdf_report(results)
                word_data = generate_word_report(results)

                # Auto-email if enabled
                if auto_email and auto_email_address and pdf_data and word_data:
                    with st.spinner(f"Sending reports to {auto_email_address}..."):
                        try:
                            send_email_with_reports(symbol, pdf_data, word_data, auto_email_address)
                            st.success(f"📧 Reports automatically sent to {auto_email_address}")
                        except Exception as email_error:
                            st.warning(f"Auto-email failed: {str(email_error)}")

            except Exception as e:
                st.error(f"Document generation error: {e}")

            col1, col2, col3, col4 = st.columns(4)

            with col1:
                # Download PDF
                if pdf_data:
                    st.download_button(
                        label="📄 Download PDF",
                        data=pdf_data,
                        file_name=f"{symbol}_Annual_Report_Analysis.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

            with col2:
                # Download Word
                if word_data:
                    st.download_button(
                        label="📝 Download Word",
                        data=word_data,
                        file_name=f"{symbol}_Annual_Report_Analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            with col3:
                # JSON export
                json_str = json.dumps(results, indent=2, default=str)
                st.download_button(
                    label="📊 Download JSON",
                    data=json_str,
                    file_name=f"{symbol}_annual_report_analysis.json",
                    mime="application/json",
                    use_container_width=True
                )

            with col4:
                # Markdown summary
                md_content = f"# {symbol} Annual Report Analysis\n\n"
                md_content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
                if results.get('executive_summary'):
                    md_content += f"## Executive Summary\n\n{results['executive_summary']}\n\n"
                if results.get('yoy_comparisons'):
                    md_content += "## Year-over-Year Comparisons\n\n"
                    for comp in results['yoy_comparisons']:
                        md_content += f"### {comp.get('years', '')}\n\n{comp.get('analysis', '')}\n\n"

                st.download_button(
                    label="📋 Download MD",
                    data=md_content,
                    file_name=f"{symbol}_analysis_summary.md",
                    mime="text/markdown",
                    use_container_width=True
                )

            # Email section
            st.markdown("---")
            st.markdown("### 📧 Email Reports")

            email_col1, email_col2 = st.columns([3, 1])

            with email_col1:
                recipient_email = st.text_input(
                    "Recipient Email",
                    value=EMAIL_ADDRESS or "",
                    placeholder="Enter email address",
                    key="email_recipient"
                )

            with email_col2:
                st.write("")  # Spacer for alignment
                st.write("")
                email_button = st.button(
                    "📧 Send Email",
                    type="primary",
                    use_container_width=True,
                    disabled=(not pdf_data or not word_data)
                )

            if email_button:
                if not recipient_email:
                    st.error("Please enter a recipient email address")
                elif pdf_data and word_data:
                    with st.spinner("Sending email..."):
                        try:
                            send_email_with_reports(symbol, pdf_data, word_data, recipient_email)
                            st.success(f"✅ Reports sent successfully to {recipient_email}")
                        except Exception as e:
                            st.error(f"❌ Failed to send email: {str(e)}")

        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"Error analyzing {symbol}: {str(e)}")
            st.exception(e)

    elif not symbol:
        st.info("👈 Enter a stock ticker and click 'Analyze Reports' to begin.")

    else:
        # Show placeholder content
        st.info("👈 Enter a stock ticker and click 'Analyze Reports' to begin.")

        st.markdown("---")
        st.markdown("### How it works")

        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            **1. Fetch Reports**

            We retrieve the last 3 annual reports (10-K filings) from SEC EDGAR via FMP API.
            """)
        with col2:
            st.markdown("""
            **2. Extract Sections**

            Key sections are extracted:
            - Business Description
            - Risk Factors
            - MD&A
            - Financial Statements
            """)
        with col3:
            st.markdown("""
            **3. AI Analysis**

            Claude/GPT analyzes each section for:
            - Business changes
            - Growing/concerning risks
            - Financial trajectory
            - Management signals
            """)


if __name__ == "__main__":
    main()
