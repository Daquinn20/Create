#!/usr/bin/env python3
"""
FMP Earnings Summarizer - All-in-One with Auto Word/PDF Generation
Analyzes earnings transcripts and automatically creates professional reports
"""

import os
import sys

# Fix Windows console encoding for emoji support
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import argparse
import json
import requests
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from typing import List, Dict, Optional
from datetime import datetime
from pathlib import Path
import anthropic
import openai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import shutil
import subprocess


def load_api_keys_from_config():
    """Load API keys from config.json file"""
    config_file = Path(__file__).parent / 'config.json'

    if config_file.exists():
        print("📁 Loading API keys from config.json...")
        with open(config_file, 'r') as f:
            config = json.load(f)

        if 'FMP_API_KEY' in config:
            os.environ['FMP_API_KEY'] = config['FMP_API_KEY']
        if 'ANTHROPIC_API_KEY' in config:
            os.environ['ANTHROPIC_API_KEY'] = config['ANTHROPIC_API_KEY']
        if 'OPENAI_API_KEY' in config:
            os.environ['OPENAI_API_KEY'] = config['OPENAI_API_KEY']

        print("   ✓ API keys loaded from config.json\n")
        return True
    else:
        print("⚠️  No config.json file found")
        return False


def load_signature_config():
    """Load signature/branding config from signature.json"""
    config_file = Path(__file__).parent / 'signature.json'
    default_config = {
        "company_name": "Targeted Equity Consulting Group",
        "analyst_name": "David Quinn",
        "email": "daquinn@targetedequityconsulting.com",
        "phone": "617-905-7415",
        "logo_file": "company_logo.png"
    }

    if config_file.exists():
        try:
            with open(config_file, 'r') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️  Could not load signature.json: {e}")
    return default_config


def send_email_report(symbol: str, pdf_path: str, company_name: str = None):
    """Send the earnings analysis report via email"""
    config_file = Path(__file__).parent / 'config.json'

    if not config_file.exists():
        print("   ⚠️ Email skipped (no config.json with email settings)")
        return False

    with open(config_file, 'r') as f:
        config = json.load(f)

    email_address = config.get('email_address')
    password = config.get('password')
    recipient = config.get('email_recipient', email_address)
    smtp_server = config.get('smtp_server', 'smtp.gmail.com')
    smtp_port = config.get('smtp_port', 587)

    if not email_address or not password:
        print("   ⚠️ Email skipped (missing email credentials in config.json)")
        return False

    # Load signature for email body
    signature = load_signature_config()
    analyst_name = signature.get('analyst_name', 'Analyst')
    company = signature.get('company_name', '')

    # Create subject
    display_name = company_name if company_name else symbol
    subject = f"{symbol} ({display_name}) Earnings Transcript Analysis"

    # Create message
    msg = MIMEMultipart()
    msg['Subject'] = subject
    msg['From'] = email_address
    msg['To'] = recipient

    # Email body
    body = f"""Hi,

Please find attached the {symbol} Earnings Transcript Analysis report.

This report includes:
- Guidance changes analysis (revenue, margins, debt)
- Management tone analysis (bullish/bearish shifts)
- Positive and negative highlights
- Quarter-over-quarter changes
- Investment implications

Best regards,
{analyst_name}
{company}
"""
    msg.attach(MIMEText(body, 'plain'))

    # Attach PDF
    if os.path.exists(pdf_path):
        with open(pdf_path, 'rb') as f:
            pdf = MIMEApplication(f.read(), _subtype='pdf')
            pdf.add_header('Content-Disposition', 'attachment',
                          filename=f'{symbol}_Earnings_Analysis.pdf')
            msg.attach(pdf)

    # Send
    try:
        print(f"\n📧 Sending report to {recipient}...")
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(email_address, password)
            server.send_message(msg)
        print(f"   ✓ Email sent successfully!")
        return True
    except Exception as e:
        print(f"   ❌ Email failed: {e}")
        return False


class FMPEarningsSummarizer:
    def __init__(self, fmp_api_key: str, anthropic_api_key: str, openai_api_key: str):
        self.fmp_api_key = fmp_api_key
        self.anthropic_client = anthropic.Anthropic(api_key=anthropic_api_key)
        self.openai_client = openai.OpenAI(api_key=openai_api_key)
        self.fmp_v4_base_url = "https://financialmodelingprep.com/api/v4"
        self.fmp_base_url = "https://financialmodelingprep.com/api/v3"
        self.signature = load_signature_config()

    def _fetch_transcripts_for_year(self, symbol: str, year: int) -> List[Dict]:
        """Fetch transcripts for a single year from the FMP batch API"""
        url = f"{self.fmp_v4_base_url}/batch_earning_call_transcript/{symbol.upper()}"
        params = {'year': year, 'apikey': self.fmp_api_key}

        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        if isinstance(data, dict) and 'Error Message' in data:
            return []
        if not isinstance(data, list):
            return []
        return data

    def fetch_recent_transcripts(self, symbol: str, num_transcripts: int = 4) -> List[Dict]:
        """Fetch transcripts using v4 batch API, querying current and prior year"""
        print(f"📊 Fetching earnings transcripts for {symbol}...")

        current_year = datetime.now().year

        try:
            # Fetch current year first
            all_items = self._fetch_transcripts_for_year(symbol, current_year)

            # If we don't have enough, also fetch prior year
            if len(all_items) < num_transcripts:
                prior_items = self._fetch_transcripts_for_year(symbol, current_year - 1)
                all_items.extend(prior_items)

            if not all_items:
                print(f"   ❌ No transcripts available")
                return []

            # Sort by year desc, quarter desc to get most recent first
            all_items.sort(key=lambda x: (x.get('year', 0), x.get('quarter', 0)), reverse=True)

            print(f"   ✓ Found {len(all_items)} transcripts")

            transcripts = []
            for item in all_items[:num_transcripts]:
                quarter = item.get('quarter')
                year = item.get('year')
                content_text = item.get('content', '')
                date = item.get('date', 'Unknown')

                if content_text:
                    word_count = len(content_text.split())
                    transcripts.append({
                        'symbol': symbol,
                        'year': year,
                        'quarter': quarter,
                        'date': date,
                        'content': content_text,
                        'word_count': word_count
                    })
                    print(f"   📄 Q{quarter} {year} ✓ ({word_count:,} words)")

            return transcripts

        except Exception as e:
            print(f"   ❌ Error: {e}")
            return []

    def get_company_profile(self, symbol: str) -> Optional[Dict]:
        url = f"{self.fmp_base_url}/profile/{symbol}"
        params = {'apikey': self.fmp_api_key}

        try:
            response = requests.get(url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            return data[0] if isinstance(data, list) and len(data) > 0 else None
        except Exception as e:
            print(f"   ⚠️ Could not fetch company profile: {e}")
            return None

    def load_user_views(self, symbol: str) -> Optional[str]:
        """Load user's investment views from views/{symbol}.txt or .docx"""
        views_dir = Path(__file__).parent / 'views'

        # Check for .txt file first
        txt_file = views_dir / f'{symbol.upper()}.txt'
        if txt_file.exists():
            print(f"📝 Loading your views from views/{symbol.upper()}.txt...")
            with open(txt_file, 'r', encoding='utf-8') as f:
                views = f.read().strip()
            if views:
                print(f"   ✓ Loaded {len(views.split()):,} words of investment views\n")
                return views

        # Check for .docx file
        docx_file = views_dir / f'{symbol.upper()}.docx'
        if docx_file.exists():
            print(f"📝 Loading your views from views/{symbol.upper()}.docx...")
            try:
                doc = Document(docx_file)
                views = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                if views:
                    print(f"   ✓ Loaded {len(views.split()):,} words of investment views\n")
                    return views
            except Exception as e:
                print(f"   ⚠️ Could not read .docx file: {e}")

        # Also search for any file containing the symbol name
        for file in views_dir.glob('*'):
            if symbol.upper() in file.name.upper() and file.suffix in ['.txt', '.docx'] and not file.name.startswith('~$'):
                print(f"📝 Loading your views from views/{file.name}...")
                try:
                    if file.suffix == '.txt':
                        with open(file, 'r', encoding='utf-8') as f:
                            views = f.read().strip()
                    else:  # .docx
                        doc = Document(file)
                        views = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    if views:
                        print(f"   ✓ Loaded {len(views.split()):,} words of investment views\n")
                        return views
                except Exception as e:
                    print(f"   ⚠️ Could not read file: {e}")

        return None

    def load_prior_analysis(self, symbol: str, output_dir: str) -> Optional[str]:
        """Load the most recent prior analysis from OneDrive reports folder or output directory"""
        symbol_upper = symbol.upper()

        # Primary: OneDrive earnings reports folder (naming: SYMBOL_date.docx)
        onedrive_dir = Path.home() / 'OneDrive' / 'Documents' / 'Targeted Equity Consulting Group' / 'TECG Earnings Report Analysis'
        if onedrive_dir.exists():
            # Find files matching the symbol prefix
            matches = sorted(
                [f for f in onedrive_dir.glob('*') if f.name.upper().startswith(symbol_upper + '_') and f.suffix in ['.docx', '.txt']],
                key=lambda f: f.stat().st_mtime,
                reverse=True
            )
            if matches:
                latest = matches[0]
                print(f"📋 Loading prior analysis from {latest.name}...")
                try:
                    if latest.suffix == '.docx':
                        doc = Document(latest)
                        content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    else:
                        with open(latest, 'r', encoding='utf-8') as f:
                            content = f.read().strip()
                    if content:
                        print(f"   ✓ Loaded {len(content.split()):,} words of prior analysis\n")
                        return content
                except Exception as e:
                    print(f"   ⚠️ Could not read prior analysis: {e}")

        # Fallback: output directory (naming: SYMBOL_claude_summary.txt)
        prior_file = Path(output_dir) / f'{symbol_upper}_claude_summary.txt'
        if prior_file.exists():
            print(f"📋 Loading prior analysis from {prior_file.name}...")
            try:
                with open(prior_file, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                if content:
                    print(f"   ✓ Loaded {len(content.split()):,} words of prior analysis\n")
                    return content
            except Exception as e:
                print(f"   ⚠️ Could not read prior analysis: {e}")

        return None

    def create_summary_prompt(self, symbol: str, transcripts: List[Dict],
                              company_info: Optional[Dict] = None,
                              user_views: Optional[str] = None,
                              prior_analysis: Optional[str] = None) -> str:
        header = f"COMPANY: {symbol}\n"
        if company_info:
            header += f"Name: {company_info.get('companyName', 'N/A')}\n"
            header += f"Industry: {company_info.get('industry', 'N/A')}\n"
            header += f"Sector: {company_info.get('sector', 'N/A')}\n"
        header += "\n"

        combined_text = ""
        for i, transcript in enumerate(transcripts, 1):
            combined_text += f"\n\n{'=' * 80}\n"
            combined_text += f"TRANSCRIPT {i}: Q{transcript['quarter']} {transcript['year']}\n"
            combined_text += f"Date: {transcript['date']}\n"
            combined_text += f"Word Count: {transcript['word_count']:,}\n"
            combined_text += f"{'=' * 80}\n\n"
            combined_text += transcript['content']

        # Build the analysis request - CHANGE DETECTION FRAMEWORK
        analysis_sections = """
== EXECUTIVE SUMMARY (Put this FIRST) ==

Provide a clear verdict at the very top:

VERDICT: [POSITIVE / NEGATIVE / NEUTRAL]
VERDICT REASON: [One sentence explaining the single most important factor]

STRATEGIC SITUATION: [2-3 sentences on what's MOST IMPORTANT happening at this company right now. If something is CHANGING - lead with that. If steady execution - describe what they're compounding on.]

KEY POSITIVES:
• [Most important positive with evidence]
• [Second positive]
• [Third positive]

KEY CONCERNS:
• [Most important concern with evidence]
• [Second concern]
• [Third concern]

BOTTOM LINE: [1-2 sentences - direct, actionable takeaway. Is this interesting, risky, boring, or broken?]

== DETAILED ANALYSIS ==

1. WHAT IS CHANGING? (MOST IMPORTANT - Analyze this first)

   Your PRIMARY job is to DETECT CHANGE. What is DIFFERENT from prior quarters?

   NARRATIVE SHIFTS:
   - Is management using NEW language to describe what the company does?
   - Are they claiming a NEW or EXPANDED addressable market?
   - Are they distancing themselves from their legacy business?
   - What do they talk about MORE now vs. previous quarters?
   - What do they talk about LESS or avoid discussing?
   - Are they emphasizing DIFFERENT metrics than before?

   BUSINESS MODEL SHIFTS:
   - Revenue mix changing between segments/products?
   - Pricing model evolution?
   - Customer mix shifting (size, industry, geography)?
   - Margin profile changing and WHY?

   STRATEGIC REPOSITIONING:
   - New markets, geographies, or customer segments being targeted?
   - M&A activity or divestitures signaling new direction?
   - New partnerships or ecosystem plays?
   - Exiting or de-emphasizing legacy businesses?

   If there is NO significant change, say so clearly - that's useful information too.

2. GUIDANCE CHANGES (Compare across quarters):
   - Revenue guidance: Any raises, cuts, or narrowing of ranges?
   - Margin guidance: Gross margin, operating margin expectations
   - Capital allocation: Changes in leverage targets, buybacks, dividends
   - Segment-specific guidance changes
   - Full-year vs quarterly outlook shifts

3. MANAGEMENT & LEADERSHIP:
   - Any executive changes (CEO, CFO, key departures/hires)? What does it signal?
   - Changes in who presents or answers questions
   - Tone shifts: More BULLISH or BEARISH vs prior quarters?
   - Confidence level - hedging language vs confident language

4. KEY BUSINESS DRIVERS (Use the company's OWN metrics):
   - What metrics does MANAGEMENT keep highlighting? Those are their KPIs
   - Don't apply generic templates - listen to what THEY emphasize
   - Report the actual metrics and values they discuss
   - Note which metrics are growing faster/slower than before

5. Q&A SESSION DEEP DIVE (Critical - unscripted reveals the most):
   - FOCUS ON THE MOST RECENT QUARTER'S Q&A
   - Evasive or deflective answers: Which questions did management dodge?
   - Surprising disclosures: What surfaced ONLY because an analyst asked?
   - Tone shifts: Less confident or defensive on certain topics?
   - Analyst pushback: Where did analysts challenge the narrative?
   - Repeated themes: What topics did multiple analysts probe?
   - Compare to prior Q&As - are analysts asking NEW questions?

6. POSITIVE HIGHLIGHTS:
   - Guidance raises or beats
   - New growth drivers or opportunities
   - Market share gains
   - Margin expansion signals
   - Evidence the strategy is working

7. NEGATIVE HIGHLIGHTS / RED FLAGS:
   - Guidance cuts or misses
   - Margin compression signals
   - Competitive pressures mentioned
   - Evasive answers to analyst questions
   - Metrics getting worse or being de-emphasized

8. INVESTMENT IMPLICATIONS:
   - Bull case: What has to go RIGHT?
   - Bear case: What could go WRONG?
   - Key metrics to monitor that prove/disprove the thesis
   - What would change the outlook?"""

        # Add user views section if provided
        user_views_section = ""
        if user_views:
            user_views_section = f"""

{'=' * 80}
ANALYST'S INVESTMENT VIEWS (Please evaluate against transcript evidence):
{'=' * 80}
{user_views}
{'=' * 80}

9. Views Evaluation - Assess the analyst's views above:
   - Which views are SUPPORTED by the transcript evidence?
   - Which views are CHALLENGED or contradicted?
   - What new information from the transcripts should update these views?
   - Rate overall alignment: Strong Support / Partial Support / Mixed / Contradicted"""

        # Add prior analysis section if available
        prior_analysis_section = ""
        if prior_analysis:
            prior_analysis_section = f"""

{'=' * 80}
PRIOR ANALYSIS (Your previous analysis of this company — compare and update):
{'=' * 80}
{prior_analysis}
{'=' * 80}

IMPORTANT: Compare your new analysis against the prior analysis above.
   - What has CHANGED since the last analysis?
   - Which prior concerns have been RESOLVED or WORSENED?
   - What NEW developments were not in the prior analysis?
   - Has the investment thesis STRENGTHENED or WEAKENED?
   - Call out any narrative shifts or reversals from the prior analysis."""

        prompt = f"""{header}
Please analyze these {len(transcripts)} earnings call transcripts for {symbol} and provide a comprehensive investment-focused summary.

CRITICAL INSTRUCTIONS:
1. START WITH THE EXECUTIVE SUMMARY - verdict, strategic situation, key positives/concerns, bottom line
2. Your PRIMARY job is to DETECT CHANGE - what is DIFFERENT from prior quarters? If management is repositioning the company, pivoting strategy, or changing their narrative - that's the most important thing to capture.
3. The Q&A SESSION DEEP DIVE is MANDATORY - this is where management gives unscripted responses and often reveals more than prepared remarks.
4. Use the company's OWN language and metrics - don't apply generic templates.
5. FOCUS HEAVILY ON THE MOST RECENT QUARTER — it is the freshest and most actionable data.

{analysis_sections}{user_views_section}{prior_analysis_section}

{combined_text}

Provide detailed, objective analysis for investment decision-making. Lead with what's CHANGING or what's MOST IMPORTANT. Be specific with quotes, metrics, and evidence."""

        return prompt

    def summarize_with_claude(self, symbol: str, transcripts: List[Dict],
                              company_info: Optional[Dict] = None,
                              user_views: Optional[str] = None,
                              prior_analysis: Optional[str] = None) -> str:
        prompt = self.create_summary_prompt(symbol, transcripts, company_info, user_views, prior_analysis)
        message = self.anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            system="""You are an expert equity research analyst. Your PRIMARY job is to DETECT CHANGE.

When analyzing earnings transcripts:
1. START with an EXECUTIVE SUMMARY (verdict, strategic situation, key positives/concerns, bottom line)
2. Focus on what is CHANGING - narrative shifts, business model changes, strategic repositioning
3. If management is pivoting or transforming the company, that's the MOST important thing to capture
4. Use the company's OWN metrics and language - don't apply generic templates
5. Include a dedicated Q&A Deep Dive section - unscripted answers reveal the most
6. Be objective - your verdict can be Positive, Negative, or Neutral based on evidence
7. If nothing significant is changing, say so - steady execution is useful information too""",
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text

    def summarize_with_chatgpt(self, symbol: str, transcripts: List[Dict],
                               company_info: Optional[Dict] = None,
                               user_views: Optional[str] = None,
                               prior_analysis: Optional[str] = None) -> str:
        prompt = self.create_summary_prompt(symbol, transcripts, company_info, user_views, prior_analysis)
        response = self.openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """You are an expert equity research analyst. Your PRIMARY job is to DETECT CHANGE.

When analyzing earnings transcripts:
1. START with an EXECUTIVE SUMMARY (verdict, strategic situation, key positives/concerns, bottom line)
2. Focus on what is CHANGING - narrative shifts, business model changes, strategic repositioning
3. If management is pivoting or transforming the company, that's the MOST important thing to capture
4. Use the company's OWN metrics and language - don't apply generic templates
5. Include a dedicated Q&A Deep Dive section - unscripted answers reveal the most
6. Be objective - your verdict can be Positive, Negative, or Neutral based on evidence
7. If nothing significant is changing, say so - steady execution is useful information too"""},
                {"role": "user", "content": prompt}
            ],
            max_tokens=8000
        )
        return response.choices[0].message.content

    def create_word_document(self, text_content: str, output_file: str, title: str):
        """Create a formatted Word document from text content"""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add company logo at top
        logo_file = self.signature.get('logo_file', 'company_logo.png')
        logo_path = Path(__file__).parent / logo_file
        if logo_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(4.0))
            doc.add_paragraph()

        # Split content into lines
        lines = text_content.split('\n')

        # Process each line
        for line in lines:
            line = line.rstrip()

            if not line:
                doc.add_paragraph()
                continue

            # Skip separator lines (=== or --- or ___)
            if '=' * 10 in line or '-' * 10 in line or '_' * 10 in line:
                continue

            # Skip lines that are just dashes or equals
            stripped = line.strip()
            if stripped and all(c in '-=_' for c in stripped):
                continue

            # Markdown ### headers - make bold
            if line.startswith('###'):
                text = line.lstrip('#').strip()
                p = doc.add_paragraph(text)
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(11)
                continue

            # Markdown ## headers - make bold and slightly larger
            if line.startswith('##'):
                text = line.lstrip('#').strip()
                p = doc.add_paragraph(text)
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 51, 102)
                continue

            # Headers
            if (line.isupper() and len(line) > 10) or \
                    line.startswith('EARNINGS TRANSCRIPT') or \
                    line.startswith('SUMMARY BY') or \
                    line.startswith("CLAUDE'S") or \
                    line.startswith("CHATGPT'S"):
                p = doc.add_paragraph(line)
                p.style = 'Heading 1'
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                continue

            # Section headers
            if (line and line[0].isdigit() and '. ' in line[:5]) or \
                    (line.startswith(('Company:', 'Industry:', 'Sector:', 'Analysis Date:'))):
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.font.bold = True
                if line[0].isdigit():
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 102, 204)
                continue

            # Bullet points - handle various formats
            stripped = line.strip()
            if stripped.startswith(('• ', '- ', '* ', '○ ', '– ', '— ')):
                # Remove bullet character and space
                text = stripped[2:].strip()
                if text:
                    p = doc.add_paragraph(text, style='List Bullet')
                continue

            # Handle double dash bullets (-- text)
            if stripped.startswith('--') and len(stripped) > 2 and stripped[2:].strip():
                text = stripped[2:].strip()
                p = doc.add_paragraph(text, style='List Bullet')
                continue

            # Bold text marked with **
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.font.bold = True
                continue

            # Regular paragraph
            doc.add_paragraph(line)

        # Add signature footer
        doc.add_paragraph()
        doc.add_paragraph()

        # Separator line
        p = doc.add_paragraph("_" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Analyst name
        analyst_name = self.signature.get('analyst_name', 'David Quinn')
        p = doc.add_paragraph(analyst_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

        # Company name in footer
        footer_company = self.signature.get('company_name', 'Targeted Equity Consulting Group')
        p = doc.add_paragraph(footer_company)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Contact info
        email = self.signature.get('email', '')
        phone = self.signature.get('phone', '')
        if email or phone:
            contact_line = " | ".join(filter(None, [email, phone]))
            p = doc.add_paragraph(contact_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.save(output_file)
        print(f"      ✓ Word: {os.path.basename(output_file)}")

    def convert_to_pdf(self, word_file: str, pdf_file: str):
        """Convert Word document to PDF using docx2pdf (Windows/Mac) or LibreOffice"""
        # Try docx2pdf first (uses Microsoft Word on Windows)
        try:
            from docx2pdf import convert
            convert(word_file, pdf_file)
            print(f"      ✓ PDF: {os.path.basename(pdf_file)}")
            return True
        except ImportError:
            pass  # docx2pdf not installed
        except Exception as e:
            print(f"      ⚠️ docx2pdf error: {e}")

        # Fallback to LibreOffice (resolve full path first)
        soffice_path = shutil.which('soffice')
        if soffice_path:
            try:
                result = subprocess.run(
                    [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir',
                     os.path.dirname(pdf_file), word_file],
                    capture_output=True,
                    text=True,
                    timeout=30
                )

                if result.returncode == 0:
                    print(f"      ✓ PDF: {os.path.basename(pdf_file)}")
                    return True
            except Exception as e:
                print(f"      ⚠️ LibreOffice conversion failed: {e}")

        print(f"      ⚠️ PDF skipped (install docx2pdf: pip install docx2pdf)")
        return False

    def save_all_formats(self, symbol: str, claude_summary: Optional[str],
                         chatgpt_summary: Optional[str],
                         transcripts: List[Dict], company_info: Optional[Dict],
                         output_dir: str):
        """Save in all formats: TXT, DOCX, and PDF. Only generates reports for models that ran."""
        os.makedirs(output_dir, exist_ok=True)

        # Create header
        header = f"EARNINGS TRANSCRIPT ANALYSIS: {symbol}\n"
        if company_info:
            header += f"Company: {company_info.get('companyName', 'N/A')}\n"
            header += f"Industry: {company_info.get('industry', 'N/A')}\n"
            header += f"Sector: {company_info.get('sector', 'N/A')}\n"
        header += f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        header += "=" * 80 + "\n\n"
        header += "TRANSCRIPTS ANALYZED:\n"
        for i, t in enumerate(transcripts, 1):
            header += f"{i}. Q{t['quarter']} {t['year']} - {t['date']} ({t['word_count']:,} words)\n"
        header += "=" * 80 + "\n\n"

        print("\n💾 Saving reports...")

        has_claude = claude_summary is not None
        has_chatgpt = chatgpt_summary is not None

        # Claude summary (only if generated)
        if has_claude:
            claude_base = os.path.join(output_dir, f'{symbol}_claude_summary')
            claude_txt = claude_base + '.txt'
            claude_docx = claude_base + '.docx'
            claude_pdf = claude_base + '.pdf'

            with open(claude_txt, 'w', encoding='utf-8') as f:
                f.write(header + "SUMMARY BY CLAUDE\n" + "=" * 80 + "\n\n" + claude_summary)
            print(f"   📄 Claude Summary:")
            print(f"      ✓ Text: {os.path.basename(claude_txt)}")

            self.create_word_document(header + claude_summary, claude_docx,
                                      f"{symbol} - Claude Analysis")
            self.convert_to_pdf(claude_docx, claude_pdf)

        # ChatGPT summary (only if generated)
        if has_chatgpt:
            chatgpt_base = os.path.join(output_dir, f'{symbol}_chatgpt_summary')
            chatgpt_txt = chatgpt_base + '.txt'
            chatgpt_docx = chatgpt_base + '.docx'
            chatgpt_pdf = chatgpt_base + '.pdf'

            with open(chatgpt_txt, 'w', encoding='utf-8') as f:
                f.write(header + "SUMMARY BY CHATGPT\n" + "=" * 80 + "\n\n" + chatgpt_summary)
            print(f"   📄 ChatGPT Summary:")
            print(f"      ✓ Text: {os.path.basename(chatgpt_txt)}")

            self.create_word_document(header + chatgpt_summary, chatgpt_docx,
                                      f"{symbol} - ChatGPT Analysis")
            self.convert_to_pdf(chatgpt_docx, chatgpt_pdf)

        # Comparison (only if both models ran)
        if has_claude and has_chatgpt:
            comparison_base = os.path.join(output_dir, f'{symbol}_comparison')
            comparison_txt = comparison_base + '.txt'
            comparison_docx = comparison_base + '.docx'
            comparison_pdf = comparison_base + '.pdf'

            comparison_content = header + "\n" + "=" * 80 + "\nCLAUDE'S ANALYSIS\n" + "=" * 80 + "\n\n" + \
                                 claude_summary + "\n\n" + "=" * 80 + "\nCHATGPT'S ANALYSIS\n" + "=" * 80 + \
                                 "\n\n" + chatgpt_summary

            with open(comparison_txt, 'w', encoding='utf-8') as f:
                f.write(comparison_content)
            print(f"   📄 Comparison Report:")
            print(f"      ✓ Text: {os.path.basename(comparison_txt)}")

            self.create_word_document(comparison_content, comparison_docx,
                                      f"{symbol} - Comparison Report")
            self.convert_to_pdf(comparison_docx, comparison_pdf)

        # JSON
        json_path = os.path.join(output_dir, f'{symbol}_transcripts.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump({
                'symbol': symbol,
                'company_info': company_info,
                'transcripts': transcripts
            }, f, indent=2)

        return (
            claude_txt if has_claude else None,
            chatgpt_txt if has_chatgpt else None,
            comparison_txt if (has_claude and has_chatgpt) else None
        )


def main():
    parser = argparse.ArgumentParser(
        description='Analyze earnings transcripts with auto Word/PDF generation'
    )

    parser.add_argument('symbol', type=str, help='Stock ticker (e.g., AAPL, MSFT)')
    parser.add_argument('-q', '--quarters', type=int, default=4,
                        help='Number of quarters (default: 4)')
    parser.add_argument('-o', '--output', type=str, default='./output',
                        help='Output directory (default: ./output)')
    parser.add_argument('--claude-only', action='store_true',
                        help='Generate only Claude summary')
    parser.add_argument('--chatgpt-only', action='store_true',
                        help='Generate only ChatGPT summary')
    parser.add_argument('--email', action='store_true',
                        help='Email the PDF report after generation')

    args = parser.parse_args()

    # Validate ticker symbol
    if not re.match(r'^[A-Za-z]{1,5}$', args.symbol):
        print(f"❌ Invalid ticker symbol: '{args.symbol}' (expected 1-5 letters, e.g. AAPL)")
        sys.exit(1)

    print("\n" + "=" * 80)
    print(f"📈 EARNINGS ANALYZER: {args.symbol.upper()}")
    print("=" * 80 + "\n")

    if not load_api_keys_from_config():
        sys.exit(1)

    fmp_api_key = os.environ.get('FMP_API_KEY')
    anthropic_api_key = os.environ.get('ANTHROPIC_API_KEY')
    openai_api_key = os.environ.get('OPENAI_API_KEY')

    missing_keys = []
    if not fmp_api_key:
        missing_keys.append('FMP_API_KEY')
    if not args.chatgpt_only and not anthropic_api_key:
        missing_keys.append('ANTHROPIC_API_KEY')
    if not args.claude_only and not openai_api_key:
        missing_keys.append('OPENAI_API_KEY')

    if missing_keys:
        print("❌ Missing API keys in config.json:")
        for key in missing_keys:
            print(f"   - {key}")
        sys.exit(1)

    print("✓ API keys configured\n")

    summarizer = FMPEarningsSummarizer(fmp_api_key, anthropic_api_key or "", openai_api_key or "")

    company_info = summarizer.get_company_profile(args.symbol.upper())
    if company_info:
        print(f"🏢 {company_info.get('companyName', args.symbol)}")
        print(f"   Industry: {company_info.get('industry', 'N/A')}")
        print(f"   Sector: {company_info.get('sector', 'N/A')}\n")
    else:
        print(f"⚠️  Could not fetch company profile for {args.symbol.upper()} — report header will omit company details\n")

    # Load user's investment views if available
    user_views = summarizer.load_user_views(args.symbol.upper())

    # Load prior analysis if available
    prior_analysis = summarizer.load_prior_analysis(args.symbol.upper(), args.output)

    transcripts = summarizer.fetch_recent_transcripts(args.symbol.upper(), num_transcripts=args.quarters)

    if not transcripts:
        print(f"\n❌ Could not fetch transcripts for {args.symbol.upper()}")
        sys.exit(1)

    total_words = sum(t['word_count'] for t in transcripts)
    print(f"\n✓ Loaded {len(transcripts)} transcripts ({total_words:,} total words)\n")

    claude_summary = None
    chatgpt_summary = None

    if not args.chatgpt_only:
        print("🤖 Generating Claude analysis...")
        try:
            claude_summary = summarizer.summarize_with_claude(args.symbol.upper(), transcripts, company_info, user_views, prior_analysis)
            print("   ✓ Complete\n")
        except Exception as e:
            print(f"   ❌ Error: {e}\n")
            claude_summary = f"Error: {e}"

    if not args.claude_only:
        print("🤖 Generating ChatGPT analysis...")
        try:
            chatgpt_summary = summarizer.summarize_with_chatgpt(args.symbol.upper(), transcripts, company_info, user_views, prior_analysis)
            print("   ✓ Complete\n")
        except Exception as e:
            print(f"   ❌ Error: {e}\n")
            chatgpt_summary = f"Error: {e}"

    summarizer.save_all_formats(
        args.symbol.upper(), claude_summary, chatgpt_summary, transcripts, company_info, args.output
    )

    print("\n" + "=" * 80)
    print("✅ ANALYSIS COMPLETE!")
    print("=" * 80)
    print(f"📁 All reports saved to: {os.path.abspath(args.output)}")
    print(f"📊 Formats: TXT, DOCX, PDF\n")

    # Send email if requested
    if args.email:
        # Email whichever report was generated (prefer Claude, fall back to ChatGPT)
        if claude_summary is not None:
            pdf_path = os.path.join(args.output, f'{args.symbol.upper()}_claude_summary.pdf')
        else:
            pdf_path = os.path.join(args.output, f'{args.symbol.upper()}_chatgpt_summary.pdf')
        company_name = company_info.get('companyName') if company_info else None
        send_email_report(args.symbol.upper(), pdf_path, company_name)


if __name__ == "__main__":
    main()