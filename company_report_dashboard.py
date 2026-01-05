"""
Company Report Dashboard - Streamlit Version
Imports all functions from company_report_backend.py to preserve exact functionality
"""
import streamlit as st
import os
import sys
from datetime import datetime
from typing import Dict, Any
import pandas as pd
from concurrent.futures import ThreadPoolExecutor, as_completed
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from io import BytesIO

# Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os as os_module

# Import all functions from the Flask backend
# This preserves exact same metrics and PDF generation
from company_report_backend import (
    # API functions
    fmp_get,
    FMP_API_KEY,
    ANTHROPIC_API_KEY,
    OPENAI_API_KEY,
    # Data fetching functions
    get_business_overview,
    get_revenue_segments,
    get_competitive_advantages,
    get_key_metrics_data,
    get_valuations,
    get_risks,
    get_recent_highlights,
    get_competition,
    get_management,
    get_balance_sheet_metrics,
    get_technical_analysis,
    get_investment_thesis,
    get_competitive_analysis_ai,
    # PDF generation
    generate_pdf_report,
    # Multi-agent system
    SPECIALIZED_AGENTS,
    run_all_agents_parallel,
)

# Page configuration
st.set_page_config(
    page_title="Company Report Generator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS with black border and centered logo
st.markdown("""
<style>
    .main .block-container {
        border: 1.5px solid black;
        padding: 30px;
        border-radius: 10px;
        background-color: white;
        max-width: 1200px;
        margin: auto;
    }
    .stMetric {
        background-color: #f0f2f6;
        padding: 10px;
        border-radius: 5px;
    }
    .section-header {
        background-color: #2c2c2c;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .logo-container {
        text-align: center;
        margin-bottom: 20px;
    }
    .logo-container img {
        max-width: 400px;
        height: auto;
    }
</style>
""", unsafe_allow_html=True)


def get_secret(key):
    """Get secret from environment or Streamlit secrets"""
    value = os.getenv(key)
    if not value:
        try:
            value = st.secrets.get(key)
        except:
            pass
    return value

EMAIL_ADDRESS = get_secret("EMAIL_ADDRESS")
EMAIL_PASSWORD = get_secret("EMAIL_PASSWORD")


def send_report_email(pdf_buffer, word_buffer, symbol, recipient_email):
    """Send PDF and Word report via email."""
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        return False, "Email credentials not configured"

    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = recipient_email
        msg['Subject'] = f"Company Report - {symbol} - {datetime.now().strftime('%B %d, %Y')}"

        body_text = f"""Your Company Report for {symbol} is attached.

Both PDF and Word document versions are included.

Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

Targeted Equity Consulting Group
"""
        msg.attach(MIMEText(body_text, 'plain'))

        date_str = datetime.now().strftime('%Y%m%d')

        # Attach PDF
        pdf_buffer.seek(0)
        pdf_attachment = MIMEBase('application', 'octet-stream')
        pdf_attachment.set_payload(pdf_buffer.read())
        encoders.encode_base64(pdf_attachment)
        pdf_attachment.add_header('Content-Disposition', 'attachment',
                            filename=f"{symbol}_Company_Report_{date_str}.pdf")
        msg.attach(pdf_attachment)

        # Attach Word document
        word_buffer.seek(0)
        word_attachment = MIMEBase('application', 'octet-stream')
        word_attachment.set_payload(word_buffer.read())
        encoders.encode_base64(word_attachment)
        word_attachment.add_header('Content-Disposition', 'attachment',
                            filename=f"{symbol}_Company_Report_{date_str}.docx")
        msg.attach(word_attachment)

        # Send via Gmail SMTP
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.sendmail(EMAIL_ADDRESS, recipient_email, msg.as_string())
        server.quit()

        return True, "Email sent successfully with PDF and Word attachments!"
    except Exception as e:
        return False, f"Email error: {str(e)}"


def set_table_keep_together(table):
    """Set table to keep together on one page (prevent page breaks)."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.get_or_add_pPr()
                keepNext = OxmlElement('w:keepNext')
                keepLines = OxmlElement('w:keepLines')
                pPr.append(keepNext)
                pPr.append(keepLines)


def generate_word_report(report_data: Dict[str, Any]) -> BytesIO:
    """Generate Word document report matching PDF format."""
    doc = Document()

    symbol = report_data.get('symbol', 'N/A')
    overview = report_data.get('business_overview', {})

    # Add company logo if it exists (1.3x size = 3.9 inches)
    logo_path = 'company_logo.png'
    if os_module.path.exists(logo_path):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(3.9))
        except:
            pass  # Skip if logo can't be added

    # Title
    title = doc.add_heading(f"{overview.get('company_name', symbol)} ({symbol})", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Generated date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section 1: Company Details
    doc.add_heading("1. Company Details", level=1)

    def fmt_num(val):
        if val is None or val == 0:
            return "N/A"
        if abs(val) >= 1e9:
            return f"${val/1e9:.2f}B"
        elif abs(val) >= 1e6:
            return f"${val/1e6:.2f}M"
        return f"${val:,.0f}"

    # Format values
    price = f"${overview.get('price', 0):.2f}" if overview.get('price') else "N/A"
    market_cap = fmt_num(overview.get('market_cap', 0))
    high_52 = f"${overview.get('week_52_high'):.2f}" if isinstance(overview.get('week_52_high'), (int, float)) else "N/A"
    low_52 = f"${overview.get('week_52_low'):.2f}" if isinstance(overview.get('week_52_low'), (int, float)) else "N/A"
    beta = overview.get('beta')
    beta_str = f"{beta:.2f}" if isinstance(beta, (int, float)) and beta else "N/A"
    employees = overview.get('employees')
    # Handle int, float, or string
    if isinstance(employees, int):
        employees_str = f"{employees:,}"
    elif isinstance(employees, float):
        employees_str = f"{int(employees):,}"
    elif isinstance(employees, str) and employees.isdigit():
        employees_str = f"{int(employees):,}"
    else:
        employees_str = str(employees) if employees and employees != "N/A" else "N/A"

    # Format enterprise value
    ev = overview.get('enterprise_value', 0)
    if ev and ev > 0:
        if ev >= 1e9:
            ev_str = f"${ev/1e9:.2f}B"
        elif ev >= 1e6:
            ev_str = f"${ev/1e6:.2f}M"
        else:
            ev_str = f"${ev:,.0f}"
    else:
        ev_str = "N/A"

    # Format dividend yield
    div_yield = overview.get('dividend_yield', 0)
    div_yield_str = f"{div_yield:.2f}%" if isinstance(div_yield, (int, float)) and div_yield else "N/A"

    headquarters = overview.get('headquarters', 'N/A')

    # 6-column table layout
    details_table = doc.add_table(rows=3, cols=6)
    details_table.style = 'Table Grid'

    # Row 1: Headers
    headers = ["Ticker", "Current Price", "Market Cap", "Enterprise Value", "52-Week High", "52-Week Low"]
    for i, header in enumerate(headers):
        details_table.rows[0].cells[i].text = header

    # Row 2: Values
    values_row1 = [symbol, price, market_cap, ev_str, high_52, low_52]
    for i, value in enumerate(values_row1):
        details_table.rows[1].cells[i].text = str(value)

    # Row 3: Second set of headers
    headers2 = ["Industry", "Sector", "Headquarters", "Beta", "Employees", "Dividend Yield"]
    for i, header in enumerate(headers2):
        details_table.rows[2].cells[i].text = header

    set_table_keep_together(details_table)

    # Second table for profile values
    profile_table = doc.add_table(rows=1, cols=6)
    profile_table.style = 'Table Grid'
    values_row2 = [
        str(overview.get('industry', 'N/A')),
        str(overview.get('sector', 'N/A')),
        str(headquarters),
        beta_str,
        employees_str,
        div_yield_str
    ]
    for i, value in enumerate(values_row2):
        profile_table.rows[0].cells[i].text = value
    set_table_keep_together(profile_table)

    doc.add_paragraph()

    # Section 2: Business Overview
    doc.add_heading("2. Business Overview", level=1)
    doc.add_paragraph(overview.get('description', 'No description available'))

    # Section 3: Revenue by Segment
    doc.add_heading("3. Revenue by Segment", level=1)
    revenue_data = report_data.get('revenue_data', {})

    # Historical Margins Table (10 years)
    historical_margins = revenue_data.get('historical_margins', [])
    if historical_margins:
        doc.add_paragraph("Margins - 10 Year History", style='Heading 2')

        # Build table with periods as columns
        periods = [m.get('period', 'N/A') for m in historical_margins[:11]]
        num_cols = len(periods) + 1  # +1 for Metric column

        margin_table = doc.add_table(rows=4, cols=num_cols)
        margin_table.style = 'Table Grid'

        # Header row
        margin_table.rows[0].cells[0].text = "Metric"
        for i, period in enumerate(periods):
            margin_table.rows[0].cells[i + 1].text = str(period)

        # Gross Margin row
        margin_table.rows[1].cells[0].text = "Gross Margin"
        for i, m in enumerate(historical_margins[:11]):
            margin_table.rows[1].cells[i + 1].text = f"{m.get('gross_margin', 0):.1f}%"

        # Operating Margin row
        margin_table.rows[2].cells[0].text = "Operating Margin"
        for i, m in enumerate(historical_margins[:11]):
            margin_table.rows[2].cells[i + 1].text = f"{m.get('operating_margin', 0):.1f}%"

        # Net Margin row
        margin_table.rows[3].cells[0].text = "Net Margin"
        for i, m in enumerate(historical_margins[:11]):
            margin_table.rows[3].cells[i + 1].text = f"{m.get('net_margin', 0):.1f}%"

        set_table_keep_together(margin_table)
        doc.add_paragraph()
    else:
        # Fallback to simple margins
        margins = revenue_data.get('margins', {})
        if margins:
            doc.add_paragraph(f"Gross Margin: {margins.get('gross_margin', 0):.2f}%")
            doc.add_paragraph(f"Operating Margin: {margins.get('operating_margin', 0):.2f}%")
            doc.add_paragraph(f"Net Margin: {margins.get('net_margin', 0):.2f}%")

    # Segments table
    segments = revenue_data.get('segments', [])
    if segments:
        doc.add_paragraph("Segments", style='Heading 2')
        total_rev = sum(s.get('revenue', 0) or 0 for s in segments)

        seg_table = doc.add_table(rows=1, cols=3)
        seg_table.style = 'Table Grid'
        seg_table.rows[0].cells[0].text = "Segment"
        seg_table.rows[0].cells[1].text = "Revenue"
        seg_table.rows[0].cells[2].text = "% of Total"

        for seg in segments[:10]:
            name = seg.get('name', 'N/A')
            revenue = seg.get('revenue', 0)
            if revenue and revenue > 0:
                pct = (revenue / total_rev * 100) if total_rev > 0 else 0
                row = seg_table.add_row()
                row.cells[0].text = name
                row.cells[1].text = fmt_num(revenue)
                row.cells[2].text = f"{pct:.1f}%"

        set_table_keep_together(seg_table)

    # Section 4: Highlights
    doc.add_heading("4. Highlights from Recent Quarters", level=1)
    highlights = report_data.get('recent_highlights', [])
    for h in highlights[:4]:
        quarter = h.get('quarter', 'N/A')
        doc.add_heading(quarter, level=2)
        for detail in h.get('details', []):
            doc.add_paragraph(f"• {detail}", style='List Bullet')

    # Section 5: Competitive Advantages
    doc.add_heading("5. Competitive Advantages", level=1)
    advantages = report_data.get('competitive_advantages', [])
    for i, adv in enumerate(advantages[:6], 1):
        doc.add_paragraph(f"{i}. {adv}")

    # Section 6: Key Metrics
    doc.add_heading("6. Key Metrics", level=1)
    metrics = report_data.get('key_metrics', {})

    def fmt_pct(val):
        if val is None or (isinstance(val, (int, float)) and val == 0):
            return 'N/A'
        return f"{val:.1f}%"

    metrics_table = doc.add_table(rows=5, cols=6)
    metrics_table.style = 'Table Grid'

    headers = ["Metric", "5 Year Avg", "3 Yr Avg", "TTM", "Est 1 Yr", "Est 2 Yr"]
    for i, h in enumerate(headers):
        metrics_table.rows[0].cells[i].text = h

    metric_rows = [
        ("Revenue Growth", 'revenue_growth_5yr', 'revenue_growth_3yr', 'revenue_growth_ttm', 'revenue_growth_est_1yr', 'revenue_growth_est_2yr'),
        ("Gross Margin", 'gross_margin_5yr', 'gross_margin_3yr', None, 'gross_margin_est_1yr', 'gross_margin_est_2yr'),
        ("Operating Margin", 'operating_margin_5yr', 'operating_margin_3yr', None, 'operating_margin_est_1yr', 'operating_margin_est_2yr'),
        ("Net Margin", 'net_income_margin_5yr', 'net_income_margin_3yr', None, 'net_income_margin_est_1yr', 'net_income_margin_est_2yr'),
    ]

    for row_idx, (label, *keys) in enumerate(metric_rows, 1):
        metrics_table.rows[row_idx].cells[0].text = label
        for col_idx, key in enumerate(keys, 1):
            if key is None:
                val = metrics.get('gross_margin', 0) * 100 if 'Gross' in label else metrics.get('operating_margin', 0) * 100 if 'Operating' in label else metrics.get('net_margin', 0) * 100
            else:
                val = metrics.get(key)
            metrics_table.rows[row_idx].cells[col_idx].text = fmt_pct(val)
    set_table_keep_together(metrics_table)

    doc.add_paragraph()

    # Section 7: Valuations
    doc.add_heading("7. Valuations", level=1)
    valuations = report_data.get('valuations', {})
    current_val = valuations.get('current', valuations)

    val_table = doc.add_table(rows=3, cols=4)
    val_table.style = 'Table Grid'
    val_data = [
        ("P/E Ratio", f"{current_val.get('pe_ratio', 0):.2f}" if current_val.get('pe_ratio') else 'N/A',
         "Price/Book", f"{current_val.get('price_to_book', 0):.2f}" if current_val.get('price_to_book') else 'N/A'),
        ("Price/Sales", f"{current_val.get('price_to_sales', 0):.2f}" if current_val.get('price_to_sales') else 'N/A',
         "EV/EBITDA", f"{current_val.get('ev_to_ebitda', 0):.2f}" if current_val.get('ev_to_ebitda') else 'N/A'),
        ("PEG Ratio", f"{current_val.get('peg_ratio', 0):.2f}" if current_val.get('peg_ratio') else 'N/A',
         "Price/FCF", f"{current_val.get('price_to_fcf', 0):.2f}" if current_val.get('price_to_fcf') else 'N/A'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(val_data):
        val_table.rows[i].cells[0].text = l1
        val_table.rows[i].cells[1].text = str(v1)
        val_table.rows[i].cells[2].text = l2
        val_table.rows[i].cells[3].text = str(v2)
    set_table_keep_together(val_table)

    doc.add_paragraph()

    # Section 8: Balance Sheet
    doc.add_heading("8. Balance Sheet / Credit Metrics", level=1)
    balance = report_data.get('balance_sheet_metrics', {})
    current_bs = balance.get('current', {})
    if current_bs:
        bs_table = doc.add_table(rows=7, cols=2)
        bs_table.style = 'Table Grid'
        bs_items = [
            ("Total Assets", fmt_num(current_bs.get('total_assets'))),
            ("Total Liabilities", fmt_num(current_bs.get('total_liabilities'))),
            ("Total Equity", fmt_num(current_bs.get('total_equity'))),
            ("Cash & Equivalents", fmt_num(current_bs.get('cash_and_equivalents'))),
            ("Total Debt", fmt_num(current_bs.get('total_debt'))),
            ("Net Debt", fmt_num(current_bs.get('net_debt'))),
            ("Working Capital", fmt_num(current_bs.get('working_capital'))),
        ]
        for i, (label, value) in enumerate(bs_items):
            bs_table.rows[i].cells[0].text = label
            bs_table.rows[i].cells[1].text = str(value)
        set_table_keep_together(bs_table)

    doc.add_paragraph()

    # Section 9: Technical Analysis
    doc.add_heading("9. Technical Analysis", level=1)
    technical = report_data.get('technical_analysis', {})
    price_data = technical.get('price_data', {})
    moving_avgs = technical.get('moving_averages', {})

    if price_data:
        doc.add_paragraph(f"Current Price: ${price_data.get('current_price', 0):.2f}")
        doc.add_paragraph(f"Day Change: {price_data.get('change_percent', 0):+.2f}%")
        doc.add_paragraph(f"52W High: ${price_data.get('year_high', 0):.2f}")
        doc.add_paragraph(f"52W Low: ${price_data.get('year_low', 0):.2f}")

    if moving_avgs:
        doc.add_heading("Moving Averages", level=2)
        ma_table = doc.add_table(rows=5, cols=3)
        ma_table.style = 'Table Grid'
        ma_items = [
            ("SMA 10", f"${moving_avgs.get('sma_10', 0):.2f}", f"{moving_avgs.get('price_vs_sma_10', 0):+.2f}%"),
            ("SMA 20", f"${moving_avgs.get('sma_20', 0):.2f}", f"{moving_avgs.get('price_vs_sma_20', 0):+.2f}%"),
            ("SMA 50", f"${moving_avgs.get('sma_50', 0):.2f}", f"{moving_avgs.get('price_vs_sma_50', 0):+.2f}%"),
            ("SMA 100", f"${moving_avgs.get('sma_100', 0):.2f}", "N/A"),
            ("SMA 200", f"${moving_avgs.get('sma_200', 0):.2f}", f"{moving_avgs.get('price_vs_sma_200', 0):+.2f}%"),
        ]
        for i, (ma, val, vs) in enumerate(ma_items):
            ma_table.rows[i].cells[0].text = ma
            ma_table.rows[i].cells[1].text = val
            ma_table.rows[i].cells[2].text = vs
        set_table_keep_together(ma_table)

    doc.add_paragraph()

    # Section 10: Risks
    doc.add_heading("10. Risks and Red Flags", level=1)
    risks = report_data.get('risks', {})

    doc.add_heading("A) Company Red Flags", level=2)
    for i, risk in enumerate(risks.get('company_specific', [])[:8], 1):
        doc.add_paragraph(f"{i}. {risk}")

    doc.add_heading("B) General Risks", level=2)
    for i, risk in enumerate(risks.get('general', [])[:8], 1):
        doc.add_paragraph(f"{i}. {risk}")

    # Section 11: Management
    doc.add_heading("11. Management", level=1)
    management = report_data.get('management', {})
    key_execs = management.get('key_executives', []) if isinstance(management, dict) else management

    if key_execs:
        exec_table = doc.add_table(rows=min(len(key_execs), 10) + 1, cols=3)
        exec_table.style = 'Table Grid'
        exec_table.rows[0].cells[0].text = "Name"
        exec_table.rows[0].cells[1].text = "Title"
        exec_table.rows[0].cells[2].text = "Pay"

        for i, exec in enumerate(key_execs[:10], 1):
            if isinstance(exec, dict):
                exec_table.rows[i].cells[0].text = exec.get('name', 'N/A')
                exec_table.rows[i].cells[1].text = exec.get('title', 'N/A')
                pay = exec.get('pay')
                exec_table.rows[i].cells[2].text = f"${pay:,.0f}" if pay else 'N/A'
        set_table_keep_together(exec_table)

    doc.add_paragraph()

    # Investment Thesis
    thesis = report_data.get('investment_thesis', {})
    if thesis:
        doc.add_heading("Investment Thesis", level=1)

        if thesis.get('summary'):
            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(thesis['summary'])

        if thesis.get('bull_case'):
            doc.add_heading("Bull Case", level=2)
            for point in thesis['bull_case']:
                doc.add_paragraph(f"• {point}", style='List Bullet')

        if thesis.get('bear_case'):
            doc.add_heading("Bear Case", level=2)
            for point in thesis['bear_case']:
                doc.add_paragraph(f"• {point}", style='List Bullet')

    # Competitive Analysis
    comp_analysis = report_data.get('competitive_analysis', {})
    if comp_analysis:
        doc.add_heading("Competitive Analysis", level=1)

        if comp_analysis.get('moat_analysis'):
            doc.add_heading("Moat Analysis", level=2)
            doc.add_paragraph(comp_analysis['moat_analysis'])

        if comp_analysis.get('competitive_position'):
            doc.add_heading("Competitive Position", level=2)
            doc.add_paragraph(comp_analysis['competitive_position'])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Targeted Equity Consulting Group")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def format_large_number(value):
    """Format large numbers with B/M suffix"""
    if value is None or value == 0:
        return "N/A"
    if abs(value) >= 1e9:
        return f"${value/1e9:.2f}B"
    elif abs(value) >= 1e6:
        return f"${value/1e6:.2f}M"
    else:
        return f"${value:,.0f}"


def format_percent(value, multiplier=1):
    """Format percentage values"""
    if value is None or value == 0:
        return "N/A"
    return f"{value * multiplier:.1f}%"


def display_company_details(overview: Dict[str, Any]):
    """Display Section 1: Company Details"""
    st.markdown("### 1. Company Details")

    # Format values
    price = f"${overview.get('price', 0):.2f}" if overview.get('price') else "N/A"
    market_cap = format_large_number(overview.get('market_cap', 0))
    high_52 = f"${overview.get('week_52_high'):.2f}" if isinstance(overview.get('week_52_high'), (int, float)) else "N/A"
    low_52 = f"${overview.get('week_52_low'):.2f}" if isinstance(overview.get('week_52_low'), (int, float)) else "N/A"
    beta = overview.get('beta')
    beta_str = f"{beta:.2f}" if isinstance(beta, (int, float)) and beta else "N/A"
    employees = overview.get('employees')
    # Handle int, float, or string
    if isinstance(employees, int):
        employees_str = f"{employees:,}"
    elif isinstance(employees, float):
        employees_str = f"{int(employees):,}"
    elif isinstance(employees, str) and employees.isdigit():
        employees_str = f"{int(employees):,}"
    else:
        employees_str = str(employees) if employees and employees != "N/A" else "N/A"

    # Format enterprise value
    ev = overview.get('enterprise_value', 0)
    if ev and ev > 0:
        if ev >= 1e9:
            ev_str = f"${ev/1e9:.2f}B"
        elif ev >= 1e6:
            ev_str = f"${ev/1e6:.2f}M"
        else:
            ev_str = f"${ev:,.0f}"
    else:
        ev_str = "N/A"

    # Format dividend yield
    div_yield = overview.get('dividend_yield', 0)
    div_yield_str = f"{div_yield:.2f}%" if isinstance(div_yield, (int, float)) and div_yield else "N/A"

    headquarters = overview.get('headquarters', 'N/A')

    # 6-column layout
    col1, col2, col3, col4, col5, col6 = st.columns(6)

    with col1:
        st.metric("Ticker", overview.get('ticker', 'N/A'))
        st.metric("Industry", overview.get('industry', 'N/A'))

    with col2:
        st.metric("Current Price", price)
        st.metric("Sector", overview.get('sector', 'N/A'))

    with col3:
        st.metric("Market Cap", market_cap)
        st.metric("Headquarters", headquarters)

    with col4:
        st.metric("Enterprise Value", ev_str)
        st.metric("Beta", beta_str)

    with col5:
        st.metric("52-Week High", high_52)
        st.metric("Employees", employees_str)

    with col6:
        st.metric("52-Week Low", low_52)
        st.metric("Dividend Yield", div_yield_str)


def display_business_overview(overview: Dict[str, Any]):
    """Display Section 2: Business Overview"""
    st.markdown("### 2. Business Overview")

    description = overview.get('description', 'No description available')
    st.markdown(description)


def display_revenue_segments(revenue_data: Dict[str, Any]):
    """Display Section 3: Revenue by Segment"""
    st.markdown("### 3. Revenue by Segment")

    # Historical Margins Table (10 years)
    historical_margins = revenue_data.get('historical_margins', [])
    if historical_margins:
        st.markdown("**Margins - 10 Year History**")

        # Build DataFrame for display
        periods = [m.get('period', 'N/A') for m in historical_margins[:11]]
        gross_margins = [f"{m.get('gross_margin', 0):.1f}%" for m in historical_margins[:11]]
        operating_margins = [f"{m.get('operating_margin', 0):.1f}%" for m in historical_margins[:11]]
        net_margins = [f"{m.get('net_margin', 0):.1f}%" for m in historical_margins[:11]]

        margin_df = pd.DataFrame({
            'Metric': ['Gross Margin', 'Operating Margin', 'Net Margin'],
            **{period: [gross_margins[i], operating_margins[i], net_margins[i]] for i, period in enumerate(periods)}
        })
        st.dataframe(margin_df, use_container_width=True, hide_index=True)
    else:
        # Fallback to simple margins
        margins = revenue_data.get('margins', {})
        if margins:
            st.markdown("**Margins**")
            margin_df = pd.DataFrame([
                {"Metric": "Gross Margin", "Value": f"{margins.get('gross_margin', 0):.2f}%"},
                {"Metric": "Operating Margin", "Value": f"{margins.get('operating_margin', 0):.2f}%"},
                {"Metric": "Net Margin", "Value": f"{margins.get('net_margin', 0):.2f}%"},
            ])
            st.dataframe(margin_df, use_container_width=True, hide_index=True)

    # Segments
    st.markdown("**Segments**")
    segments = revenue_data.get('segments', [])
    if segments:
        # Build segment table
        segment_data = []
        total_rev = sum(s.get('revenue', 0) or 0 for s in segments)
        for segment in segments[:10]:
            name = segment.get('name', 'N/A')
            revenue = segment.get('revenue', 0)
            if revenue and revenue > 0:
                pct = (revenue / total_rev * 100) if total_rev > 0 else 0
                segment_data.append({
                    "Segment": name,
                    "Revenue": format_large_number(revenue),
                    "% of Total": f"{pct:.1f}%"
                })
        if segment_data:
            st.dataframe(pd.DataFrame(segment_data), use_container_width=True, hide_index=True)

    # Show AI analysis if available
    if segments and segments[0].get('ai_analysis'):
        with st.expander("AI Segment Analysis"):
            st.markdown(segments[0]['ai_analysis'])


def display_recent_highlights(highlights: list):
    """Display Section 4: Highlights from Recent Quarters"""
    st.markdown("### 4. Highlights from Recent Quarters")

    # Show AI summary if available
    if highlights and highlights[0].get('ai_summary'):
        with st.expander("AI Quarterly Trends Analysis", expanded=True):
            st.markdown(highlights[0]['ai_summary'])

    for highlight in highlights[:4]:
        quarter = highlight.get('quarter', 'N/A')
        details = highlight.get('details', [])

        with st.expander(f"**{quarter}**"):
            for detail in details:
                st.write(f"- {detail}")


def display_competitive_advantages(advantages: list):
    """Display Section 5: Competitive Advantages"""
    st.markdown("### 5. Competitive Advantages")

    for i, advantage in enumerate(advantages[:6], 1):
        st.write(f"{i}. {advantage}")


def display_key_metrics(metrics: Dict[str, Any]):
    """Display Section 6: Key Metrics"""
    st.markdown("### 6. Key Metrics")

    # Revenue Growth & Margins Table
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Growth & Margins**")

        def fmt_pct(val):
            if val is None or (isinstance(val, (int, float)) and val == 0):
                return 'N/A'
            return f"{val:.1f}%"

        growth_data = pd.DataFrame([
            {
                "Metric": "Revenue Growth",
                "5 Year Avg": fmt_pct(metrics.get('revenue_growth_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('revenue_growth_3yr')),
                "TTM": fmt_pct(metrics.get('revenue_growth_ttm')),
                "Est 1 Yr": fmt_pct(metrics.get('revenue_growth_est_1yr')),
                "Est 2 Yr": fmt_pct(metrics.get('revenue_growth_est_2yr')),
            },
            {
                "Metric": "Gross Margin",
                "5 Year Avg": fmt_pct(metrics.get('gross_margin_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('gross_margin_3yr')),
                "TTM": fmt_pct(metrics.get('gross_margin', 0) * 100),
                "Est 1 Yr": fmt_pct(metrics.get('gross_margin_est_1yr')),
                "Est 2 Yr": fmt_pct(metrics.get('gross_margin_est_2yr')),
            },
            {
                "Metric": "Operating Margin",
                "5 Year Avg": fmt_pct(metrics.get('operating_margin_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('operating_margin_3yr')),
                "TTM": fmt_pct(metrics.get('operating_margin', 0) * 100),
                "Est 1 Yr": fmt_pct(metrics.get('operating_margin_est_1yr')),
                "Est 2 Yr": fmt_pct(metrics.get('operating_margin_est_2yr')),
            },
            {
                "Metric": "Net Income Margin",
                "5 Year Avg": fmt_pct(metrics.get('net_income_margin_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('net_income_margin_3yr')),
                "TTM": fmt_pct(metrics.get('net_margin', 0) * 100),
                "Est 1 Yr": fmt_pct(metrics.get('net_income_margin_est_1yr')),
                "Est 2 Yr": fmt_pct(metrics.get('net_income_margin_est_2yr')),
            },
        ])
        st.dataframe(growth_data, use_container_width=True, hide_index=True)

    with col2:
        st.markdown("**Return Metrics**")

        returns_data = pd.DataFrame([
            {
                "Metric": "ROIC",
                "5 Year Avg": fmt_pct(metrics.get('roic_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('roic_3yr')),
                "TTM": fmt_pct(metrics.get('roic')),
            },
            {
                "Metric": "ROE",
                "5 Year Avg": fmt_pct(metrics.get('roe_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('roe_3yr')),
                "TTM": fmt_pct(metrics.get('roe')),
            },
            {
                "Metric": "ROA",
                "5 Year Avg": fmt_pct(metrics.get('roa_5yr')),
                "3 Yr Avg": fmt_pct(metrics.get('roa_3yr')),
                "TTM": fmt_pct(metrics.get('roa')),
            },
            {
                "Metric": "WACC",
                "5 Year Avg": "-",
                "3 Yr Avg": "-",
                "TTM": fmt_pct(metrics.get('wacc')),
            },
        ])
        st.dataframe(returns_data, use_container_width=True, hide_index=True)


def display_valuations(valuations: Dict[str, Any]):
    """Display Section 7: Valuations"""
    st.markdown("### 7. Valuations")

    current_val = valuations.get('current', valuations)

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("P/E Ratio", f"{current_val.get('pe_ratio', 0):.2f}" if current_val.get('pe_ratio') else 'N/A')
        st.metric("Price/Sales", f"{current_val.get('price_to_sales', 0):.2f}" if current_val.get('price_to_sales') else 'N/A')

    with col2:
        st.metric("Price/Book", f"{current_val.get('price_to_book', 0):.2f}" if current_val.get('price_to_book') else 'N/A')
        st.metric("EV/EBITDA", f"{current_val.get('ev_to_ebitda', 0):.2f}" if current_val.get('ev_to_ebitda') else 'N/A')

    with col3:
        st.metric("PEG Ratio", f"{current_val.get('peg_ratio', 0):.2f}" if current_val.get('peg_ratio') else 'N/A')
        st.metric("Price/FCF", f"{current_val.get('price_to_fcf', 0):.2f}" if current_val.get('price_to_fcf') else 'N/A')

    # Historical valuations
    historical = valuations.get('historical', [])
    if historical:
        with st.expander("Historical Valuations (10 Years)"):
            hist_df = pd.DataFrame(historical)
            st.dataframe(hist_df, use_container_width=True, hide_index=True)


def display_balance_sheet(balance_sheet: Dict[str, Any]):
    """Display Section 8: Balance Sheet / Credit Metrics"""
    st.markdown("### 8. Balance Sheet / Credit Metrics")

    current = balance_sheet.get('current', {})

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Balance Sheet Summary**")
        if current:
            bs_data = pd.DataFrame([
                {"Item": "Total Assets", "Value": format_large_number(current.get('total_assets'))},
                {"Item": "Total Liabilities", "Value": format_large_number(current.get('total_liabilities'))},
                {"Item": "Total Equity", "Value": format_large_number(current.get('total_equity'))},
                {"Item": "Cash & Equivalents", "Value": format_large_number(current.get('cash_and_equivalents'))},
                {"Item": "Total Debt", "Value": format_large_number(current.get('total_debt'))},
                {"Item": "Net Debt", "Value": format_large_number(current.get('net_debt'))},
                {"Item": "Working Capital", "Value": format_large_number(current.get('working_capital'))},
            ])
            st.dataframe(bs_data, use_container_width=True, hide_index=True)

    with col2:
        st.markdown("**Liquidity Ratios (TTM)**")
        liquidity = balance_sheet.get('liquidity_ratios', {})
        if liquidity:
            liq_data = pd.DataFrame([
                {"Ratio": "Current Ratio", "Value": f"{liquidity.get('current_ratio', 0):.2f}"},
                {"Ratio": "Quick Ratio", "Value": f"{liquidity.get('quick_ratio', 0):.2f}"},
                {"Ratio": "Cash Ratio", "Value": f"{liquidity.get('cash_ratio', 0):.2f}"},
                {"Ratio": "DSO", "Value": f"{liquidity.get('days_sales_outstanding', 0):.0f} days"},
                {"Ratio": "DIO", "Value": f"{liquidity.get('days_inventory_outstanding', 0):.0f} days"},
                {"Ratio": "Cash Conversion Cycle", "Value": f"{liquidity.get('cash_conversion_cycle', 0):.0f} days"},
            ])
            st.dataframe(liq_data, use_container_width=True, hide_index=True)

    # Historical ratios
    liq_hist = balance_sheet.get('liquidity_ratios_historical', [])
    credit_hist = balance_sheet.get('credit_ratios_historical', [])

    if liq_hist:
        with st.expander("Liquidity Ratios (10-Year History)"):
            st.dataframe(pd.DataFrame(liq_hist), use_container_width=True, hide_index=True)

    if credit_hist:
        with st.expander("Credit Ratios (10-Year History)"):
            st.dataframe(pd.DataFrame(credit_hist), use_container_width=True, hide_index=True)


def display_technical_analysis(technical: Dict[str, Any]):
    """Display Section 9: Technical Analysis"""
    st.markdown("### 9. Technical Analysis")

    col1, col2, col3 = st.columns(3)

    # Price Data
    price_data = technical.get('price_data', {})
    with col1:
        st.markdown("**Price Summary**")
        if price_data:
            st.metric("Current Price", f"${price_data.get('current_price', 0):.2f}")
            st.metric("Day Change", f"{price_data.get('change_percent', 0):+.2f}%")
            st.metric("52W High", f"${price_data.get('year_high', 0):.2f}")
            st.metric("52W Low", f"${price_data.get('year_low', 0):.2f}")
            st.metric("% from 52W High", f"{price_data.get('pct_from_52w_high', 0):+.1f}%")

    # Moving Averages
    moving_avgs = technical.get('moving_averages', {})
    with col2:
        st.markdown("**Moving Averages**")
        if moving_avgs:
            ma_data = pd.DataFrame([
                {"MA": "SMA 10", "Value": f"${moving_avgs.get('sma_10', 0):.2f}", "vs Price": f"{moving_avgs.get('price_vs_sma_10', 0):+.2f}%"},
                {"MA": "SMA 20", "Value": f"${moving_avgs.get('sma_20', 0):.2f}", "vs Price": f"{moving_avgs.get('price_vs_sma_20', 0):+.2f}%"},
                {"MA": "SMA 50", "Value": f"${moving_avgs.get('sma_50', 0):.2f}", "vs Price": f"{moving_avgs.get('price_vs_sma_50', 0):+.2f}%"},
                {"MA": "SMA 100", "Value": f"${moving_avgs.get('sma_100', 0):.2f}", "vs Price": "N/A"},
                {"MA": "SMA 200", "Value": f"${moving_avgs.get('sma_200', 0):.2f}", "vs Price": f"{moving_avgs.get('price_vs_sma_200', 0):+.2f}%"},
            ])
            st.dataframe(ma_data, use_container_width=True, hide_index=True)

    # Momentum & Trend
    momentum = technical.get('momentum_indicators', {})
    trend = technical.get('trend_analysis', {})

    with col3:
        st.markdown("**Momentum & Trend**")

        rsi = momentum.get('rsi', {})
        if rsi:
            st.metric("RSI (14)", f"{rsi.get('value', 50):.1f} - {rsi.get('signal', 'Neutral')}")

        macd = momentum.get('macd', {})
        if macd:
            st.metric("MACD", f"{macd.get('macd_line', 0):.4f} ({macd.get('signal', 'N/A')})")

        stoch = momentum.get('stochastic', {})
        if stoch:
            st.metric("Stochastic %K", f"{stoch.get('k', 50):.1f} ({stoch.get('signal', 'Neutral')})")

        if trend:
            st.metric("Overall Trend", trend.get('overall_trend', 'N/A'))
            st.metric("Golden Cross", "Yes" if trend.get('golden_cross') else "No")

    # Volatility and Support/Resistance
    volatility = technical.get('volatility_indicators', {})
    support_resistance = technical.get('support_resistance', {})

    with st.expander("Volatility & Support/Resistance"):
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("**Volatility Indicators**")
            atr = volatility.get('atr', {})
            bb = volatility.get('bollinger_bands', {})
            if atr:
                st.write(f"ATR (14): {atr.get('value', 0):.2f} ({atr.get('atr_percent', 0):.2f}%)")
            if bb:
                st.write(f"Bollinger Upper: ${bb.get('upper', 0):.2f}")
                st.write(f"Bollinger Middle: ${bb.get('middle', 0):.2f}")
                st.write(f"Bollinger Lower: ${bb.get('lower', 0):.2f}")
                st.write(f"BB Width: {bb.get('width', 0):.2f}%")
                st.write(f"BB %B: {bb.get('percent_b', 0):.2f}% ({bb.get('signal', 'N/A')})")

        with col2:
            st.markdown("**Support/Resistance**")
            if support_resistance:
                st.write(f"Pivot Point: ${support_resistance.get('pivot', 0):.2f}")
                st.write(f"Resistance 1: ${support_resistance.get('resistance_1', 0):.2f}")
                st.write(f"Resistance 2: ${support_resistance.get('resistance_2', 0):.2f}")
                st.write(f"Support 1: ${support_resistance.get('support_1', 0):.2f}")
                st.write(f"Support 2: ${support_resistance.get('support_2', 0):.2f}")


def display_risks(risks: Dict[str, Any]):
    """Display Section 10: Risks and Red Flags"""
    st.markdown("### 10. Risks and Red Flags")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**A) Company Red Flags**")
        company_specific = risks.get('company_specific', [])
        for i, risk in enumerate(company_specific[:8], 1):
            st.write(f"{i}. {risk}")

    with col2:
        st.markdown("**B) General Risks**")
        general = risks.get('general', [])
        for i, risk in enumerate(general[:8], 1):
            st.write(f"{i}. {risk}")


def display_management(management: Dict[str, Any]):
    """Display Section 11: Management"""
    st.markdown("### 11. Management")

    # If management is a list (old format), convert
    if isinstance(management, list):
        key_executives = management
        insider_trading = {}
    else:
        key_executives = management.get('key_executives', management) if isinstance(management, dict) else management
        insider_trading = management.get('insider_trading', {}) if isinstance(management, dict) else {}

    # Key Executives
    if key_executives:
        st.markdown("**Key Executives**")
        exec_data = []
        for exec in key_executives[:10]:
            if isinstance(exec, dict):
                pay = exec.get('pay')
                pay_str = f"${pay:,.0f}" if pay else 'N/A'
                exec_data.append({
                    "Name": exec.get('name', 'N/A'),
                    "Title": exec.get('title', 'N/A'),
                    "Pay": pay_str,
                    "Stock Held": exec.get('stock_held', 'N/A'),
                    "Prior Employers": exec.get('prior_employers', 'N/A')
                })

        if exec_data:
            st.dataframe(pd.DataFrame(exec_data), use_container_width=True, hide_index=True)

    # Insider Trading
    if insider_trading:
        st.markdown("**Insider Trading (Last 3 Months)**")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Buys", f"{insider_trading.get('buys_count', 0)} (${insider_trading.get('buys_value', 0):,.0f})")
        with col2:
            st.metric("Sells", f"{insider_trading.get('sells_count', 0)} (${insider_trading.get('sells_value', 0):,.0f})")


def display_competitive_analysis(comp_analysis: Dict[str, Any]):
    """Display Competitive Analysis (AI-powered)"""
    st.markdown("### Competitive Analysis")

    if comp_analysis.get('moat_analysis'):
        with st.expander("Moat Analysis", expanded=True):
            st.markdown(comp_analysis['moat_analysis'])

    if comp_analysis.get('competitive_position'):
        with st.expander("Competitive Position"):
            st.markdown(comp_analysis['competitive_position'])

    if comp_analysis.get('market_dynamics'):
        with st.expander("Market Dynamics"):
            st.markdown(comp_analysis['market_dynamics'])

    if comp_analysis.get('competitive_advantages'):
        with st.expander("Key Competitive Advantages"):
            for adv in comp_analysis['competitive_advantages']:
                st.write(f"- {adv}")

    if comp_analysis.get('industry_analysis'):
        with st.expander("Industry-Specific Analysis"):
            st.markdown(comp_analysis['industry_analysis'])


def display_investment_thesis(thesis: Dict[str, Any]):
    """Display Investment Thesis"""
    st.markdown("### Investment Thesis")

    if thesis.get('summary'):
        st.markdown("**Executive Summary**")
        st.markdown(thesis['summary'])

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Bull Case**")
        for point in thesis.get('bull_case', []):
            st.write(f"- {point}")

    with col2:
        st.markdown("**Bear Case**")
        for point in thesis.get('bear_case', []):
            st.write(f"- {point}")

    if thesis.get('key_metrics_to_watch'):
        with st.expander("Key Metrics to Watch"):
            for metric in thesis['key_metrics_to_watch']:
                st.write(f"- {metric}")

    if thesis.get('catalysts'):
        with st.expander("Upcoming Catalysts"):
            for catalyst in thesis['catalysts']:
                st.write(f"- {catalyst}")


def display_multi_agent_analysis(agent_results: Dict[str, Any]):
    """Display analysis from all 10 specialized agents."""
    st.markdown("### Multi-Agent Analysis")
    st.markdown("*10 specialized AI agents analyzed this company in parallel*")

    # Create 2 columns for agents
    col1, col2 = st.columns(2)

    agent_list = list(agent_results.items())

    for i, (agent_id, result) in enumerate(agent_list):
        col = col1 if i % 2 == 0 else col2

        with col:
            status_icon = "✅" if result.get("status") == "success" else "❌"
            with st.expander(f"{result.get('emoji', '🤖')} {result.get('agent_name', agent_id)} {status_icon}", expanded=False):
                if result.get("status") == "success":
                    st.markdown(result.get("analysis", "No analysis available"))
                else:
                    st.error(result.get("analysis", "Analysis failed"))


def main():
    st.sidebar.title("Company Report Generator")

    # Check API keys
    if not FMP_API_KEY:
        st.sidebar.error("FMP API key not configured")
        st.error("Please configure your FMP_API_KEY in .env file or Streamlit secrets")
        return

    st.sidebar.success("FMP API: Configured")
    st.sidebar.info(f"Anthropic API: {'Configured' if ANTHROPIC_API_KEY else 'Not configured'}")
    st.sidebar.info(f"OpenAI API: {'Configured' if OPENAI_API_KEY else 'Not configured'}")

    # Ticker input
    symbol = st.sidebar.text_input("Enter Ticker Symbol", value="AAPL", max_chars=10).upper().strip()

    generate_button = st.sidebar.button("Generate Report", type="primary", use_container_width=True)

    if generate_button and symbol:
        with st.spinner(f"Generating comprehensive report for {symbol}..."):
            progress_bar = st.progress(0)
            status_text = st.empty()

            try:
                # Fetch all data in PARALLEL for speed
                status_text.text("Fetching all data in parallel...")
                progress_bar.progress(10)

                # Define all data fetch tasks
                fetch_tasks = {
                    "business_overview": lambda: get_business_overview(symbol),
                    "revenue_data": lambda: get_revenue_segments(symbol),
                    "competitive_advantages": lambda: get_competitive_advantages(symbol),
                    "recent_highlights": lambda: get_recent_highlights(symbol),
                    "key_metrics": lambda: get_key_metrics_data(symbol),
                    "valuations": lambda: get_valuations(symbol),
                    "balance_sheet": lambda: get_balance_sheet_metrics(symbol),
                    "technical": lambda: get_technical_analysis(symbol),
                    "risks": lambda: get_risks(symbol),
                    "management_list": lambda: get_management(symbol),
                    "competitive_analysis": lambda: get_competitive_analysis_ai(symbol),
                }

                results = {}
                completed_count = 0
                total_tasks = len(fetch_tasks)

                # Execute all fetches in parallel
                with ThreadPoolExecutor(max_workers=10) as executor:
                    future_to_key = {executor.submit(func): key for key, func in fetch_tasks.items()}

                    for future in as_completed(future_to_key):
                        key = future_to_key[future]
                        try:
                            results[key] = future.result()
                        except Exception as e:
                            print(f"Error fetching {key}: {e}")
                            results[key] = {} if key not in ["competitive_advantages", "recent_highlights", "risks"] else []

                        completed_count += 1
                        progress = int(10 + (completed_count / total_tasks) * 60)  # 10-70%
                        progress_bar.progress(progress)
                        status_text.text(f"Fetched {completed_count}/{total_tasks}: {key}...")

                # Unpack results
                business_overview = results.get("business_overview", {})
                revenue_data = results.get("revenue_data", {})
                competitive_advantages = results.get("competitive_advantages", [])
                recent_highlights = results.get("recent_highlights", [])
                key_metrics = results.get("key_metrics", {})
                valuations = results.get("valuations", {})
                balance_sheet = results.get("balance_sheet", {})
                technical = results.get("technical", {})
                risks = results.get("risks", [])
                management_list = results.get("management_list", [])
                competitive_analysis = results.get("competitive_analysis", {})

                # Wrap management in dict format expected by PDF generator
                management = {"key_executives": management_list if isinstance(management_list, list) else []}

                progress_bar.progress(75)

                # Build report data
                report_data = {
                    "symbol": symbol,
                    "generated_at": datetime.now().isoformat(),
                    "business_overview": business_overview,
                    "revenue_data": revenue_data,
                    "competitive_advantages": competitive_advantages,
                    "recent_highlights": recent_highlights,
                    "key_metrics": key_metrics,
                    "valuations": valuations,
                    "balance_sheet_metrics": balance_sheet,
                    "technical_analysis": technical,
                    "risks": risks,
                    "management": management,
                    "competitive_analysis": competitive_analysis,
                }

                status_text.text("Generating investment thesis...")
                progress_bar.progress(85)
                report_data["investment_thesis"] = get_investment_thesis(symbol, report_data)

                # Run 10 specialized agents in parallel
                status_text.text("Running 10 AI agents in parallel...")
                progress_bar.progress(90)

                # Build company data for agents
                company_data_for_agents = {
                    "company_name": business_overview.get("company_name", symbol),
                    "industry": business_overview.get("industry", "N/A"),
                    "sector": business_overview.get("sector", "N/A"),
                    "market_cap": business_overview.get("market_cap", 0),
                    "price": business_overview.get("price", 0),
                    "revenue_growth_ttm": key_metrics.get("revenue_growth_ttm", "N/A"),
                    "gross_margin": key_metrics.get("gross_margin", "N/A"),
                    "operating_margin": key_metrics.get("operating_margin", "N/A"),
                    "net_margin": key_metrics.get("net_margin", "N/A"),
                    "roe": key_metrics.get("roe", "N/A"),
                    "roic": key_metrics.get("roic", "N/A"),
                    "pe_ratio": valuations.get("current", valuations).get("pe_ratio", "N/A"),
                    "ev_to_ebitda": valuations.get("current", valuations).get("ev_to_ebitda", "N/A"),
                    "debt_to_equity": balance_sheet.get("current", {}).get("debt_to_equity", "N/A"),
                    "beta": business_overview.get("beta", "N/A"),
                    "week_52_high": business_overview.get("week_52_high", "N/A"),
                    "week_52_low": business_overview.get("week_52_low", "N/A"),
                    "ytd_return": technical.get("price_data", {}).get("ytd_return", "N/A"),
                    "description": business_overview.get("description", "N/A"),
                }

                agent_results = run_all_agents_parallel(symbol, company_data_for_agents)
                report_data["agent_analysis"] = agent_results

                progress_bar.progress(100)
                status_text.empty()
                progress_bar.empty()

                # Store in session state
                st.session_state.report_data = report_data

            except Exception as e:
                st.error(f"Error generating report: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
                return

    # Display report if available
    if "report_data" in st.session_state:
        report_data = st.session_state.report_data
        overview = report_data.get("business_overview", {})

        if "error" in overview:
            st.error(overview["error"])
            return

        # Logo at top center
        col_logo1, col_logo2, col_logo3 = st.columns([1, 2, 1])
        with col_logo2:
            try:
                st.image("company_logo.png", use_container_width=True)
            except:
                pass  # Skip if logo not found

        st.markdown("---")

        # Header
        col1, col2 = st.columns([3, 1])
        with col1:
            st.title(f"{overview.get('company_name', 'N/A')} ({report_data['symbol']})")
            st.caption(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

        with col2:
            # PDF and Word Download buttons
            try:
                date_str = datetime.now().strftime('%Y%m%d')

                # Generate PDF
                pdf_buffer = generate_pdf_report(report_data)
                st.download_button(
                    label="Download PDF",
                    data=pdf_buffer,
                    file_name=f"{report_data['symbol']}_Company_Report_{date_str}.pdf",
                    mime="application/pdf",
                    type="primary"
                )

                # Generate Word document
                word_buffer = generate_word_report(report_data)
                st.download_button(
                    label="Download Word",
                    data=word_buffer,
                    file_name=f"{report_data['symbol']}_Company_Report_{date_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Email button (sends both PDF and Word)
                if st.button("📧 Email Reports"):
                    with st.spinner("Sending email with PDF and Word..."):
                        pdf_buffer.seek(0)
                        word_buffer.seek(0)
                        success, message = send_report_email(
                            pdf_buffer,
                            word_buffer,
                            report_data['symbol'],
                            "daquinn@targetedequityconsulting.com"
                        )
                        if success:
                            st.success(message)
                        else:
                            st.error(message)
            except Exception as e:
                st.error(f"Error generating reports: {e}")

        st.divider()

        # Display all sections
        display_company_details(overview)
        st.divider()

        display_business_overview(overview)
        st.divider()

        display_revenue_segments(report_data.get("revenue_data", {}))
        st.divider()

        display_recent_highlights(report_data.get("recent_highlights", []))
        st.divider()

        display_competitive_advantages(report_data.get("competitive_advantages", []))
        st.divider()

        display_key_metrics(report_data.get("key_metrics", {}))
        st.divider()

        display_valuations(report_data.get("valuations", {}))
        st.divider()

        display_balance_sheet(report_data.get("balance_sheet_metrics", {}))
        st.divider()

        display_technical_analysis(report_data.get("technical_analysis", {}))
        st.divider()

        display_risks(report_data.get("risks", {}))
        st.divider()

        display_management(report_data.get("management", {}))
        st.divider()

        display_competitive_analysis(report_data.get("competitive_analysis", {}))
        st.divider()

        display_investment_thesis(report_data.get("investment_thesis", {}))

        # Display Multi-Agent Analysis
        if report_data.get("agent_analysis"):
            st.divider()
            display_multi_agent_analysis(report_data.get("agent_analysis", {}))

    else:
        st.info("Enter a ticker symbol and click 'Generate Report' to get started.")

        st.markdown("""
        ### Features (Same as Original)
        - **Section 1**: Company Details (price, market cap, 52W range)
        - **Section 2**: Business Overview (AI-enhanced from annual reports)
        - **Section 3**: Revenue by Segment with margins
        - **Section 4**: Highlights from Recent Quarters (AI analysis)
        - **Section 5**: Competitive Advantages
        - **Section 6**: Key Metrics (5yr, 3yr, TTM, estimates)
        - **Section 7**: Valuations with 10-year history
        - **Section 8**: Balance Sheet / Credit Metrics (10-year history)
        - **Section 9**: Technical Analysis (SMAs, RSI, MACD, Bollinger, etc.)
        - **Section 10**: Risks and Red Flags (AI-powered)
        - **Section 11**: Management with insider trading
        - **Competitive Analysis**: AI-powered moat and market analysis
        - **Investment Thesis**: AI-generated bull/bear cases
        - **PDF Export**: Same format as original
        """)


if __name__ == "__main__":
    main()
