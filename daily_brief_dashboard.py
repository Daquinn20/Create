"""
Daily Brief Dashboard
Interactive Streamlit dashboard for market overview and daily briefing
"""
import streamlit as st
import yfinance as yf
import pandas as pd
import requests
from datetime import datetime, timedelta
import os
from dotenv import load_dotenv
import anthropic
from openai import OpenAI
from docx import Document
import io

load_dotenv()

st.set_page_config(
    page_title="Daily Brief",
    page_icon="📋",
    layout="wide"
)

# API Keys
def get_api_key(key_name):
    key = os.getenv(key_name)
    if key:
        return key
    try:
        return st.secrets.get(key_name)
    except:
        return None

FMP_API_KEY = get_api_key("FMP_API_KEY")
ANTHROPIC_API_KEY = get_api_key("ANTHROPIC_API_KEY")
OPENAI_API_KEY = get_api_key("OPENAI_API_KEY")

# Market indices
INDICES = {
    "S&P 500": "^GSPC",
    "NASDAQ": "^IXIC",
    "Dow Jones": "^DJI",
    "Russell 2000": "^RUT",
    "VIX": "^VIX"
}

GLOBAL_INDICES = {
    "FTSE 100": "^FTSE",
    "DAX": "^GDAXI",
    "Nikkei 225": "^N225",
    "Hang Seng": "^HSI"
}


@st.cache_data(ttl=300)
def fetch_market_data():
    """Fetch current market data."""
    data = {}
    for name, symbol in {**INDICES, **GLOBAL_INDICES}.items():
        try:
            ticker = yf.Ticker(symbol)
            hist = ticker.history(period="2d")
            if len(hist) >= 1:
                current = hist['Close'].iloc[-1]
                prev = hist['Close'].iloc[-2] if len(hist) >= 2 else current
                change = ((current - prev) / prev) * 100
                data[name] = {"price": current, "change": change}
        except:
            pass
    return data


@st.cache_data(ttl=300)
def fetch_sector_performance():
    """Fetch sector performance from FMP."""
    if not FMP_API_KEY:
        return None

    try:
        url = f"https://financialmodelingprep.com/api/v3/sectors-performance?apikey={FMP_API_KEY}"
        response = requests.get(url, timeout=10)
        data = response.json()
        return data if isinstance(data, list) else None
    except:
        return None


@st.cache_data(ttl=300)
def fetch_market_news():
    """Fetch market news from FMP."""
    if not FMP_API_KEY:
        return []

    try:
        url = f"https://financialmodelingprep.com/api/v3/stock_news?limit=10&apikey={FMP_API_KEY}"
        response = requests.get(url, timeout=10)
        return response.json() if response.status_code == 200 else []
    except:
        return []


@st.cache_data(ttl=300)
def fetch_economic_calendar():
    """Fetch economic calendar from FMP."""
    if not FMP_API_KEY:
        return []

    try:
        today = datetime.now().strftime("%Y-%m-%d")
        week_later = (datetime.now() + timedelta(days=7)).strftime("%Y-%m-%d")
        url = f"https://financialmodelingprep.com/api/v3/economic_calendar?from={today}&to={week_later}&apikey={FMP_API_KEY}"
        response = requests.get(url, timeout=10)
        return response.json()[:10] if response.status_code == 200 else []
    except:
        return []


def generate_ai_summary(market_data, sector_data, news):
    """Generate AI summary of market conditions."""
    prompt = f"""Provide a brief (3-4 sentences) market overview for today based on this data:

Market Indices:
{market_data}

Sector Performance:
{sector_data}

Recent Headlines:
{news}

Focus on: key market moves, notable trends, and what investors should watch."""

    if ANTHROPIC_API_KEY:
        try:
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            response = client.messages.create(
                model="claude-3-5-haiku-latest",
                max_tokens=300,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text.strip()
        except:
            pass

    if OPENAI_API_KEY:
        try:
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300
            )
            return response.choices[0].message.content.strip()
        except:
            pass

    return "AI summary unavailable."


def create_word_doc(market_data, sector_data, news, ai_summary):
    """Create Word document with daily brief."""
    doc = Document()
    doc.add_heading("Daily Market Brief", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

    doc.add_heading("AI Market Summary", level=1)
    doc.add_paragraph(ai_summary)

    doc.add_heading("Market Indices", level=1)
    for name, data in market_data.items():
        change_str = f"+{data['change']:.2f}%" if data['change'] >= 0 else f"{data['change']:.2f}%"
        doc.add_paragraph(f"{name}: {data['price']:,.2f} ({change_str})")

    if sector_data:
        doc.add_heading("Sector Performance", level=1)
        for sector in sector_data[:5]:
            doc.add_paragraph(f"{sector['sector']}: {sector['changesPercentage']}")

    doc.add_heading("News Headlines", level=1)
    for item in news[:5]:
        doc.add_paragraph(f"• {item.get('title', 'N/A')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Main App
st.title("📋 Daily Brief Dashboard")
st.markdown("Market overview and daily briefing")

# Check API keys
if not FMP_API_KEY:
    st.warning("FMP_API_KEY not found. Some features will be limited.")

if st.sidebar.button("🔄 Refresh Data", type="primary"):
    st.cache_data.clear()

# Fetch all data
with st.spinner("Fetching market data..."):
    market_data = fetch_market_data()
    sector_data = fetch_sector_performance()
    news = fetch_market_news()
    economic_calendar = fetch_economic_calendar()

# US Markets
st.header("🇺🇸 US Markets")
cols = st.columns(5)
for i, (name, symbol) in enumerate(INDICES.items()):
    if name in market_data:
        data = market_data[name]
        with cols[i]:
            delta_color = "normal" if name == "VIX" else "normal"
            st.metric(
                name,
                f"{data['price']:,.2f}",
                f"{data['change']:+.2f}%",
                delta_color="inverse" if name == "VIX" else "normal"
            )

# Global Markets
st.header("🌍 Global Markets")
cols = st.columns(4)
for i, name in enumerate(GLOBAL_INDICES.keys()):
    if name in market_data:
        data = market_data[name]
        with cols[i]:
            st.metric(name, f"{data['price']:,.2f}", f"{data['change']:+.2f}%")

# Sector Performance
if sector_data:
    st.header("📊 Sector Performance")
    cols = st.columns(3)
    for i, sector in enumerate(sector_data[:9]):
        with cols[i % 3]:
            change = sector.get('changesPercentage', '0%')
            st.metric(sector['sector'], change)

# AI Summary
st.header("🤖 AI Market Summary")
if st.button("Generate AI Summary"):
    with st.spinner("Generating AI summary..."):
        market_str = "\n".join([f"{k}: {v['change']:+.2f}%" for k, v in market_data.items()])
        sector_str = "\n".join([f"{s['sector']}: {s['changesPercentage']}" for s in (sector_data or [])[:5]])
        news_str = "\n".join([n.get('title', '')[:100] for n in news[:5]])
        summary = generate_ai_summary(market_str, sector_str, news_str)
        st.info(summary)

# News
st.header("📰 Market News")
for item in news[:5]:
    with st.expander(item.get('title', 'N/A')[:100]):
        st.write(item.get('text', 'No content')[:500])
        st.caption(f"Source: {item.get('site', 'Unknown')} | {item.get('publishedDate', '')[:10]}")

# Economic Calendar
if economic_calendar:
    st.header("📅 Economic Calendar (Next 7 Days)")
    for event in economic_calendar[:5]:
        st.write(f"**{event.get('event', 'N/A')}** - {event.get('date', '')[:10]}")

# Download
st.divider()
if market_data:
    ai_summary = "Click 'Generate AI Summary' to include AI analysis."
    doc = create_word_doc(market_data, sector_data or [], news, ai_summary)
    st.download_button(
        "📥 Download Daily Brief (Word)",
        doc,
        file_name=f"Daily_Brief_{datetime.now().strftime('%Y-%m-%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
