"""
Earnings Transcript Analysis Dashboard
Interactive Streamlit dashboard for analyzing earnings call transcripts
"""
import streamlit as st
import os
import json
import re
import requests
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path
import anthropic
import openai
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate
from reportlab.platypus.frames import Frame
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Page config
st.set_page_config(
    page_title="Earnings Transcript Analyzer",
    page_icon="📈",
    layout="wide"
)

# Custom CSS
st.markdown("""
    <style>
    .big-font {
        font-size:30px !important;
        font-weight: bold;
    }
    .stTextArea textarea {
        font-family: monospace;
    }
    </style>
    """, unsafe_allow_html=True)

# API Keys from environment or Streamlit secrets
def get_api_key(key_name):
    """Get API key from environment or Streamlit secrets"""
    # Try environment first
    key = os.getenv(key_name)
    if key:
        return key
    # Try Streamlit secrets
    try:
        return st.secrets.get(key_name)
    except Exception:
        return None

FMP_API_KEY = get_api_key("FMP_API_KEY")
ANTHROPIC_API_KEY = get_api_key("ANTHROPIC_API_KEY")
OPENAI_API_KEY = get_api_key("OPENAI_API_KEY")
EMAIL_ADDRESS = get_api_key("EMAIL_ADDRESS")
EMAIL_PASSWORD = get_api_key("EMAIL_PASSWORD")


def fetch_transcripts(symbol: str, num_quarters: int = 4) -> List[Dict]:
    """Fetch earnings transcripts from FMP API"""
    url = f"https://financialmodelingprep.com/api/v4/batch_earning_call_transcript/{symbol.upper()}"
    transcripts = []
    current_year = datetime.now().year

    try:
        # Fetch from current year and previous years until we have enough transcripts
        for year in range(current_year, current_year - 3, -1):
            if len(transcripts) >= num_quarters:
                break

            params = {'year': year, 'apikey': FMP_API_KEY}
            response = requests.get(url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()

            if isinstance(data, dict) and 'Error Message' in data:
                continue

            if not isinstance(data, list) or len(data) == 0:
                continue

            for item in data:
                if len(transcripts) >= num_quarters:
                    break
                content_text = item.get('content', '')
                if content_text:
                    transcripts.append({
                        'symbol': symbol,
                        'year': item.get('year'),
                        'quarter': item.get('quarter'),
                        'date': item.get('date', 'Unknown'),
                        'content': content_text,
                        'word_count': len(content_text.split())
                    })

        return transcripts
    except Exception as e:
        st.error(f"Error fetching transcripts: {e}")
        return []


def get_company_profile(symbol: str) -> Optional[Dict]:
    """Get company profile from FMP"""
    url = f"https://financialmodelingprep.com/api/v3/profile/{symbol}"
    params = {'apikey': FMP_API_KEY}

    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()
        return data[0] if isinstance(data, list) and len(data) > 0 else None
    except Exception:
        return None


def fetch_quarterly_financials(symbol: str, num_quarters: int = 8) -> Optional[pd.DataFrame]:
    """Fetch quarterly income statement data for charts"""
    url = f"https://financialmodelingprep.com/api/v3/income-statement/{symbol}"
    params = {'period': 'quarter', 'limit': num_quarters, 'apikey': FMP_API_KEY}

    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        if not data or isinstance(data, dict):  # API error returns dict
            return None

        # Build DataFrame
        records = []
        for item in reversed(data):  # Reverse to get oldest first
            date = item.get('date', '')
            fiscal_year = item.get('calendarYear', '')
            period = item.get('period', '')
            quarter_label = f"{period} {fiscal_year}"

            records.append({
                'Quarter': quarter_label,
                'Date': date,
                'Revenue': item.get('revenue', 0) / 1_000_000,  # Convert to millions
                'Gross Profit': item.get('grossProfit', 0) / 1_000_000,
                'Operating Income': item.get('operatingIncome', 0) / 1_000_000,
                'Net Income': item.get('netIncome', 0) / 1_000_000,
                'EPS': item.get('eps', 0)
            })

        return pd.DataFrame(records)
    except Exception as e:
        st.warning(f"Could not fetch financial data: {e}")
        return None


def fetch_stock_price_history(symbol: str, years: int = 2) -> Optional[pd.DataFrame]:
    """Fetch historical daily stock prices"""
    url = f"https://financialmodelingprep.com/api/v3/historical-price-full/{symbol}"
    params = {'apikey': FMP_API_KEY}

    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        if not data or 'historical' not in data:
            return None

        # Filter to last N years
        from datetime import timedelta
        cutoff_date = datetime.now() - timedelta(days=years * 365)

        records = []
        for item in data['historical']:
            date = datetime.strptime(item['date'], '%Y-%m-%d')
            if date >= cutoff_date:
                records.append({
                    'Date': date,
                    'Close': item.get('close', 0),
                    'Volume': item.get('volume', 0)
                })

        df = pd.DataFrame(records)
        return df.sort_values('Date') if not df.empty else None
    except Exception:
        return None


def fetch_earnings_surprises(symbol: str, num_quarters: int = 8) -> Optional[pd.DataFrame]:
    """Fetch earnings surprises (beats/misses) for footnotes"""
    url = f"https://financialmodelingprep.com/api/v3/earnings-surprises/{symbol}"
    params = {'apikey': FMP_API_KEY}

    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        if not data:
            return None

        records = []
        for item in data[:num_quarters]:
            date = item.get('date', '')
            actual = item.get('actualEarningResult', 0)
            estimated = item.get('estimatedEarning', 0)

            if estimated and estimated != 0:
                surprise_pct = ((actual - estimated) / abs(estimated)) * 100
            else:
                surprise_pct = 0

            records.append({
                'Date': date,
                'Actual EPS': actual,
                'Estimated EPS': estimated,
                'Surprise %': surprise_pct,
                'Beat/Miss': 'Beat' if surprise_pct > 0 else ('Miss' if surprise_pct < 0 else 'Met')
            })

        return pd.DataFrame(records)
    except Exception:
        return None


def create_financial_charts(symbol: str):
    """Create and display financial charts for the symbol"""
    st.markdown("---")
    st.subheader("📊 Financial Performance (Last 8 Quarters)")

    try:
        financials = fetch_quarterly_financials(symbol)
        if financials is None or financials.empty:
            st.warning(f"No quarterly financial data available for {symbol}")
            return
    except Exception as e:
        st.error(f"Error fetching financial data: {e}")
        return

    # Calculate margins upfront
    financials['Gross Margin %'] = (financials['Gross Profit'] / financials['Revenue'] * 100).round(1)
    financials['Operating Margin %'] = (financials['Operating Income'] / financials['Revenue'] * 100).round(1)

    # Calculate sequential growth rates
    growth_data = []
    for i in range(1, len(financials)):
        prev = financials.iloc[i-1]
        curr = financials.iloc[i]

        rev_growth = ((curr['Revenue'] - prev['Revenue']) / prev['Revenue'] * 100) if prev['Revenue'] != 0 else 0
        gp_growth = ((curr['Gross Profit'] - prev['Gross Profit']) / prev['Gross Profit'] * 100) if prev['Gross Profit'] != 0 else 0

        if prev['Operating Income'] != 0:
            op_growth = ((curr['Operating Income'] - prev['Operating Income']) / abs(prev['Operating Income']) * 100)
        else:
            op_growth = 0

        growth_data.append({
            'Quarter': curr['Quarter'],
            'Revenue': f"{rev_growth:+.1f}%",
            'Gross Profit': f"{gp_growth:+.1f}%",
            'Op. Income': f"{op_growth:+.1f}%"
        })

    # Layout: Charts stacked on left, Growth table on right
    col_charts, col_table = st.columns([2, 1])

    with col_charts:
        # Revenue Chart
        fig_rev = go.Figure()
        fig_rev.add_trace(go.Bar(
            x=financials['Quarter'],
            y=financials['Revenue'],
            marker_color='#4472C4',
            name='Revenue',
            text=[f'${v:,.0f}' for v in financials['Revenue']],
            textposition='outside'
        ))
        fig_rev.update_layout(
            title='Revenue ($M)',
            xaxis_tickangle=-45,
            height=280,
            margin=dict(l=40, r=40, t=40, b=60)
        )
        st.plotly_chart(fig_rev, use_container_width=True)

        # Gross Profit & Margin Chart
        fig_gp = make_subplots(specs=[[{"secondary_y": True}]])
        fig_gp.add_trace(
            go.Bar(
                x=financials['Quarter'],
                y=financials['Gross Profit'],
                marker_color='#70AD47',
                name='Gross Profit ($M)',
                text=[f'${v:,.0f}' for v in financials['Gross Profit']],
                textposition='outside'
            ),
            secondary_y=False
        )
        fig_gp.add_trace(
            go.Scatter(
                x=financials['Quarter'],
                y=financials['Gross Margin %'],
                mode='lines+markers',
                name='Gross Margin %',
                line=dict(color='#000000', width=2),
                marker=dict(size=6)
            ),
            secondary_y=True
        )
        fig_gp.update_layout(
            title='Gross Profit & Margin',
            xaxis_tickangle=-45,
            height=280,
            margin=dict(l=40, r=40, t=40, b=60),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        fig_gp.update_yaxes(title_text="$M", secondary_y=False)
        fig_gp.update_yaxes(title_text="%", secondary_y=True)
        st.plotly_chart(fig_gp, use_container_width=True)

        # Operating Income & Margin Chart
        fig_op = make_subplots(specs=[[{"secondary_y": True}]])
        colors_op = ['#ED7D31' if val >= 0 else '#C00000' for val in financials['Operating Income']]
        fig_op.add_trace(
            go.Bar(
                x=financials['Quarter'],
                y=financials['Operating Income'],
                marker_color=colors_op,
                name='Operating Income ($M)',
                text=[f'${v:,.0f}' for v in financials['Operating Income']],
                textposition='outside'
            ),
            secondary_y=False
        )
        fig_op.add_trace(
            go.Scatter(
                x=financials['Quarter'],
                y=financials['Operating Margin %'],
                mode='lines+markers',
                name='Operating Margin %',
                line=dict(color='#000000', width=2),
                marker=dict(size=6)
            ),
            secondary_y=True
        )
        fig_op.update_layout(
            title='Operating Income & Margin',
            xaxis_tickangle=-45,
            height=280,
            margin=dict(l=40, r=40, t=40, b=60),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        fig_op.update_yaxes(title_text="$M", secondary_y=False)
        fig_op.update_yaxes(title_text="%", secondary_y=True)
        st.plotly_chart(fig_op, use_container_width=True)

    with col_table:
        st.markdown("### Sequential Growth (QoQ)")
        if growth_data:
            growth_df = pd.DataFrame(growth_data)
            st.dataframe(growth_df, use_container_width=True, hide_index=True, height=750)

    # Stock Price Chart (2 years) - Full width below
    st.subheader("📉 Stock Price (2 Years)")
    price_data = fetch_stock_price_history(symbol, years=2)
    if price_data is not None and not price_data.empty:
        fig_price = go.Figure()
        fig_price.add_trace(go.Scatter(
            x=price_data['Date'],
            y=price_data['Close'],
            mode='lines',
            name='Close Price',
            line=dict(color='#4472C4', width=2),
            fill='tozeroy',
            fillcolor='rgba(68, 114, 196, 0.1)'
        ))
        fig_price.update_layout(
            title=f'{symbol} Daily Close Price',
            xaxis_title='Date',
            yaxis_title='Price ($)',
            height=350,
            hovermode='x unified',
            margin=dict(l=40, r=40, t=40, b=40)
        )
        st.plotly_chart(fig_price, use_container_width=True)
    else:
        st.info("Stock price data not available")

    # Earnings Surprises Footnotes
    surprises = fetch_earnings_surprises(symbol)
    if surprises is not None and not surprises.empty:
        st.subheader("📋 Earnings Performance Notes")

        beats = surprises[surprises['Surprise %'] > 0]
        misses = surprises[surprises['Surprise %'] < 0]

        col_beat, col_miss = st.columns(2)

        with col_beat:
            if not beats.empty:
                st.markdown("**✅ Beats:**")
                for _, row in beats.iterrows():
                    st.markdown(f"- {row['Date']}: Beat by {row['Surprise %']:.1f}% (${row['Actual EPS']:.2f} vs ${row['Estimated EPS']:.2f} est)")

        with col_miss:
            if not misses.empty:
                st.markdown("**❌ Misses:**")
                for _, row in misses.iterrows():
                    st.markdown(f"- {row['Date']}: Missed by {abs(row['Surprise %']):.1f}% (${row['Actual EPS']:.2f} vs ${row['Estimated EPS']:.2f} est)")


def load_prior_analysis(symbol: str) -> Optional[str]:
    """Load the most recent prior analysis from OneDrive reports folder or output directory"""
    symbol_upper = symbol.upper()

    # Primary: OneDrive earnings reports folder
    onedrive_dir = Path.home() / 'OneDrive' / 'Documents' / 'Targeted Equity Consulting Group' / 'TECG Earnings Report Analysis'
    if onedrive_dir.exists():
        matches = sorted(
            [f for f in onedrive_dir.glob('*') if f.name.upper().startswith(symbol_upper + '_') and f.suffix in ['.docx', '.txt']],
            key=lambda f: f.stat().st_mtime,
            reverse=True
        )
        if matches:
            latest = matches[0]
            try:
                if latest.suffix == '.docx':
                    doc = Document(latest)
                    content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                else:
                    with open(latest, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                if content:
                    return content, latest.name
            except Exception:
                pass

    # Fallback: local output directory
    output_dir = Path(__file__).parent / 'output'
    prior_file = output_dir / f'{symbol_upper}_claude_summary.txt'
    if prior_file.exists():
        try:
            with open(prior_file, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            if content:
                return content, prior_file.name
        except Exception:
            pass

    return None, None


def create_analysis_prompt(symbol: str, transcripts: List[Dict], company_info: Optional[Dict] = None,
                           prior_analysis: Optional[str] = None) -> str:
    """Create the analysis prompt"""
    header = f"COMPANY: {symbol}\n"
    if company_info:
        header += f"Name: {company_info.get('companyName', 'N/A')}\n"
        header += f"Industry: {company_info.get('industry', 'N/A')}\n"
        header += f"Sector: {company_info.get('sector', 'N/A')}\n"
    header += "\n"

    combined_text = ""
    for i, transcript in enumerate(transcripts, 1):
        combined_text += f"\n\n{'=' * 80}\n"
        combined_text += f"TRANSCRIPT {i}: Q{transcript['quarter']} {transcript['year']}\n"
        combined_text += f"Date: {transcript['date']}\n"
        combined_text += f"{'=' * 80}\n\n"
        combined_text += transcript['content']

    prompt = f"""{header}
Please analyze these {len(transcripts)} earnings call transcripts for {symbol} and provide a comprehensive investment-focused summary.

CRITICAL INSTRUCTIONS:
1. START WITH THE EXECUTIVE SUMMARY - verdict, strategic situation, key positives/concerns, bottom line
2. Your PRIMARY job is to DETECT CHANGE - what is DIFFERENT from prior quarters? If management is repositioning the company, pivoting strategy, or changing their narrative - that's the most important thing to capture.
3. The Q&A SESSION DEEP DIVE is MANDATORY - this is where management gives unscripted responses and often reveals more than prepared remarks.
4. Use the company's OWN language and metrics - don't apply generic templates.
5. FOCUS HEAVILY ON THE MOST RECENT QUARTER — it is the freshest and most actionable data.

== EXECUTIVE SUMMARY (Put this FIRST) ==

Provide a clear verdict at the very top:

VERDICT: [POSITIVE / NEGATIVE / NEUTRAL]
VERDICT REASON: [One sentence explaining the single most important factor]

STRATEGIC SITUATION: [2-3 sentences on what's MOST IMPORTANT happening at this company right now. If something is CHANGING - lead with that. If steady execution - describe what they're compounding on.]

KEY POSITIVES:
• [Most important positive with evidence]
• [Second positive]
• [Third positive]

KEY CONCERNS:
• [Most important concern with evidence]
• [Second concern]
• [Third concern]

BOTTOM LINE: [1-2 sentences - direct, actionable takeaway. Is this interesting, risky, boring, or broken?]

== DETAILED ANALYSIS ==

1. WHAT IS CHANGING? (MOST IMPORTANT - Analyze this FIRST and in DETAIL)

   Your PRIMARY job is to DETECT CHANGE. What is DIFFERENT from prior quarters? This section should be COMPREHENSIVE.

   A) NARRATIVE SHIFTS - How management describes the company:
   - Is management using NEW language to describe what the company does or its identity?
   - Are they claiming a NEW or EXPANDED total addressable market (TAM)?
   - Are they distancing themselves from their legacy business or old identity?
   - What do they talk about MORE now vs. previous quarters? (Count the emphasis)
   - What do they talk about LESS or avoid discussing entirely?
   - Are they emphasizing DIFFERENT metrics than before? Which ones?
   - Key QUOTES that signal a change in direction - cite them directly
   - Example: "We're no longer just a [X] company - we're now a [Y] platform"

   B) BUSINESS MODEL SHIFTS - How they make money:
   - Revenue mix changing between segments/products? Show the numbers
   - Pricing model evolution (one-time → recurring, usage-based, subscription)?
   - Customer mix shifting (enterprise vs SMB, new verticals, new geographies)?
   - Channel mix changing (direct vs partners, digital vs physical)?
   - Margin profile changing - WHY? Is it strategic or forced?
   - Unit economics improving or deteriorating?

   C) STRATEGIC REPOSITIONING - Where they're heading:
   - New markets, geographies, or customer segments being targeted?
   - M&A activity or divestitures signaling new direction?
   - New partnerships, alliances, or ecosystem plays?
   - Exiting, de-emphasizing, or sunsetting legacy businesses?
   - Capital allocation shifts (more R&D? more buybacks? different capex priorities?)
   - New products, platforms, or capabilities that represent a strategic bet?

   D) LEADERSHIP & CULTURE SHIFTS:
   - New CEO, CFO, or key executives? What does their background signal?
   - Changes in who presents or leads the earnings call?
   - Tone of the company changing (startup energy vs mature operator)?
   - Different analysts being addressed or different investor base being courted?

   IMPORTANT: If there is NO significant change happening, say so clearly and explain what steady-state execution looks like for this company. Not every company is pivoting - some are compounding on a working strategy.

2. GUIDANCE CHANGES (Compare across quarters):
   - Revenue guidance: Any raises, cuts, or narrowing of ranges?
   - Margin guidance: Gross margin, operating margin expectations
   - Capital allocation: Changes in leverage targets, buybacks, dividends
   - Segment-specific guidance changes
   - Full-year vs quarterly outlook shifts

3. MANAGEMENT & LEADERSHIP:
   - Any executive changes (CEO, CFO, key departures/hires)? What does it signal?
   - Changes in who presents or answers questions
   - Tone shifts: More BULLISH or BEARISH vs prior quarters?
   - Confidence level - hedging language vs confident language

4. KEY BUSINESS DRIVERS (Use the company's OWN metrics):
   - What metrics does MANAGEMENT keep highlighting? Those are their KPIs
   - Don't apply generic templates - listen to what THEY emphasize
   - Report the actual metrics and values they discuss
   - Note which metrics are growing faster/slower than before

5. Q&A SESSION DEEP DIVE (Critical - This is unscripted and reveals the most):
   - FOCUS HEAVILY ON THE MOST RECENT QUARTER'S Q&A — this is the freshest and most actionable data
   - Pay VERY close attention to the analyst Q&A section — management responses here are unscripted and often reveal more than prepared remarks
   - Evasive or deflective answers: Which questions did management dodge, redirect, or give vague answers to? These are red flags
   - Surprising disclosures: What new information surfaced ONLY because an analyst asked about it?
   - Tone shifts: Did management sound less confident or more defensive on certain topics vs their prepared remarks?
   - Analyst pushback: Where did analysts challenge management's narrative? What were they skeptical about?
   - Repeated themes: What topics did multiple analysts probe? Consensus concerns signal key investor debates
   - Off-script admissions: Any comments that contradicted or softened the prepared remarks
   - Follow-up intensity: Topics where analysts asked follow-ups suggest areas of high investor concern
   - Compare the most recent Q&A to prior quarters — are analysts asking NEW questions? Are old concerns resolved or growing?

6. TONE ANALYSIS:
   - Overall management tone: More BULLISH or BEARISH vs prior quarters?
   - Confidence level in delivery and Q&A responses
   - Use of hedging language ("uncertain", "challenging", "cautious") vs confident language ("strong", "accelerating", "exceeding")

7. POSITIVE HIGHLIGHTS:
   - Guidance raises or beats
   - New growth drivers or opportunities mentioned
   - Market share gains
   - Margin expansion signals
   - Strong forward indicators
   - Evidence the strategy is working

8. NEGATIVE HIGHLIGHTS / RED FLAGS:
   - Guidance cuts or misses
   - Margin compression signals
   - Competitive pressures mentioned
   - Macro headwinds cited
   - Unusual executive departures
   - Evasive answers to analyst questions
   - Metrics getting worse or being de-emphasized

9. QUARTER-OVER-QUARTER CHANGES:
   - What's NEW this quarter that wasn't discussed before?
   - What topics are management AVOIDING that they discussed before?
   - Shifting narrative or strategic pivots

10. INVESTMENT IMPLICATIONS:
   - Bull case: What has to go RIGHT?
   - Bear case: What could go WRONG?
   - Key metrics to monitor that prove/disprove the thesis
   - What would change the outlook?
"""

    # Add prior analysis section if available
    prior_section = ""
    if prior_analysis:
        prior_section = f"""
{'=' * 80}
PRIOR ANALYSIS (Your previous analysis of this company — compare and update):
{'=' * 80}
{prior_analysis}
{'=' * 80}

IMPORTANT: Compare your new analysis against the prior analysis above.
   - What has CHANGED since the last analysis?
   - Which prior concerns have been RESOLVED or WORSENED?
   - What NEW developments were not in the prior analysis?
   - Has the investment thesis STRENGTHENED or WEAKENED?
   - Call out any narrative shifts or reversals from the prior analysis.
"""

    prompt += f"""{prior_section}
{combined_text}

Provide detailed, objective analysis for investment decision-making. Remember: the Q&A Deep Dive section is MANDATORY and must contain specific examples from the analyst Q&A, with emphasis on the most recent quarter."""

    return prompt


def analyze_with_claude(prompt: str) -> str:
    """Analyze using Claude"""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8000,
        system="You are an expert equity research analyst. When analyzing earnings transcripts, you MUST always include a dedicated Q&A Session Deep Dive section. The Q&A is where management gives unscripted answers — it reveals more than prepared remarks. Never omit this section.",
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def analyze_with_chatgpt(symbol: str, transcripts: List[Dict], company_info: Optional[Dict] = None,
                         prior_analysis: Optional[str] = None) -> tuple:
    """Analyze using ChatGPT with automatic fallback for rate limits"""
    client = openai.OpenAI(api_key=OPENAI_API_KEY)

    # Try with all transcripts first
    quarters_to_try = len(transcripts)

    while quarters_to_try >= 1:
        try:
            prompt = create_analysis_prompt(symbol, transcripts[:quarters_to_try], company_info, prior_analysis)
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert equity research analyst. When analyzing earnings transcripts, you MUST always include a dedicated Q&A Session Deep Dive section. The Q&A is where management gives unscripted answers — it reveals more than prepared remarks. Never omit this section."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=8000
            )
            result = response.choices[0].message.content
            if quarters_to_try < len(transcripts):
                return result, f"(Analyzed {quarters_to_try} quarters due to rate limits)"
            return result, None
        except openai.RateLimitError as e:
            if quarters_to_try > 1:
                quarters_to_try -= 1
                st.warning(f"Rate limit hit. Retrying with {quarters_to_try} quarters...")
            else:
                raise e

    raise Exception("Could not complete analysis due to rate limits")


def add_page_border(doc):
    """Add a black rectangular border around the page"""
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # 1.5pt = 12 eighths of a point
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')  # Black
            pgBorders.append(border)

        sectPr.append(pgBorders)


def create_word_document(content: str, symbol: str, ai_model: str) -> io.BytesIO:
    """Create Word document and return as bytes"""
    doc = Document()

    # Add page border
    add_page_border(doc)

    # Add company logo at top center
    logo_path = Path(__file__).parent / "company_logo.png"
    if logo_path.exists():
        try:
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_paragraph.add_run()
            run.add_picture(str(logo_path), width=Inches(3))
            doc.add_paragraph()  # Spacer after logo
        except Exception as e:
            pass  # Skip logo if there's an error

    # Title
    title = doc.add_paragraph(f"{symbol} Earnings Transcript Analysis")
    title.style = 'Heading 1'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_line = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Content
    for line in content.split('\n'):
        if line.strip():
            # Check if line starts with markdown headers (#, ##, ###, etc.)
            if line.strip().startswith('#'):
                # Remove # symbols and make bold
                clean_text = line.strip().lstrip('#').strip()
                p = doc.add_paragraph()
                run = p.add_run(clean_text)
                run.bold = True
                run.font.size = Pt(12)
            else:
                # Handle **bold** markdown syntax
                p = doc.add_paragraph()
                # Split by **text** pattern
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text - remove ** and make bold
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        # Regular text
                        if part:
                            p.add_run(part)

    # Add signature at the end
    doc.add_paragraph()
    doc.add_paragraph("David A Quinn")
    doc.add_paragraph("Targeted Equity Consulting")
    doc.add_paragraph("daquinn@targetedequityconsulting.com")
    doc.add_paragraph("617-905-7415")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def draw_page_border(canvas, doc):
    """Draw a black border around the page"""
    canvas.saveState()
    canvas.setStrokeColor(colors.black)
    canvas.setLineWidth(1.5)
    # Draw rectangle with margin from edge
    margin = 0.4 * inch
    canvas.rect(margin, margin,
                letter[0] - 2*margin, letter[1] - 2*margin)
    canvas.restoreState()


def create_pdf_charts(symbol: str) -> list:
    """Create chart images for PDF export using matplotlib. Returns list of (title, image_bytes) tuples."""
    charts = []

    try:
        financials = fetch_quarterly_financials(symbol)
        if financials is None or financials.empty:
            return charts

        # Calculate margins
        financials['Gross Margin %'] = (financials['Gross Profit'] / financials['Revenue'] * 100).round(1)
        financials['Operating Margin %'] = (financials['Operating Income'] / financials['Revenue'] * 100).round(1)

        quarters = financials['Quarter'].tolist()
        x_pos = range(len(quarters))

        # Chart 1: Revenue
        fig, ax = plt.subplots(figsize=(7, 3.5))
        bars = ax.bar(x_pos, financials['Revenue'], color='#4472C4')
        ax.set_xticks(x_pos)
        ax.set_xticklabels(quarters, rotation=45, ha='right', fontsize=8)
        ax.set_title('Revenue ($M)', fontsize=12, fontweight='bold')
        ax.set_ylabel('$M')
        # Add value labels
        for bar, val in zip(bars, financials['Revenue']):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                   f'${val:,.0f}', ha='center', va='bottom', fontsize=7)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        charts.append(("Revenue", buf))
        plt.close(fig)

        # Chart 2: Gross Profit & Margin (dual axis)
        fig, ax1 = plt.subplots(figsize=(7, 3.5))
        bars = ax1.bar(x_pos, financials['Gross Profit'], color='#70AD47', label='Gross Profit ($M)')
        ax1.set_xticks(x_pos)
        ax1.set_xticklabels(quarters, rotation=45, ha='right', fontsize=8)
        ax1.set_ylabel('$M', color='#70AD47')
        ax1.tick_params(axis='y', labelcolor='#70AD47')
        # Add value labels
        for bar, val in zip(bars, financials['Gross Profit']):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                    f'${val:,.0f}', ha='center', va='bottom', fontsize=7)

        ax2 = ax1.twinx()
        ax2.plot(x_pos, financials['Gross Margin %'], color='#000000', marker='o',
                linewidth=2, markersize=5, label='Gross Margin %')
        ax2.set_ylabel('%', color='#000000')
        ax2.tick_params(axis='y', labelcolor='#000000')

        ax1.set_title('Gross Profit & Margin', fontsize=12, fontweight='bold')
        lines1, labels1 = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=8)
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        charts.append(("Gross Profit & Margin", buf))
        plt.close(fig)

        # Chart 3: Operating Income & Margin (dual axis)
        fig, ax1 = plt.subplots(figsize=(7, 3.5))
        bar_colors = ['#ED7D31' if val >= 0 else '#C00000' for val in financials['Operating Income']]
        bars = ax1.bar(x_pos, financials['Operating Income'], color=bar_colors, label='Operating Income ($M)')
        ax1.set_xticks(x_pos)
        ax1.set_xticklabels(quarters, rotation=45, ha='right', fontsize=8)
        ax1.set_ylabel('$M', color='#ED7D31')
        ax1.tick_params(axis='y', labelcolor='#ED7D31')
        ax1.axhline(y=0, color='gray', linestyle='-', linewidth=0.5)
        # Add value labels
        for bar, val in zip(bars, financials['Operating Income']):
            y_pos = bar.get_height() + 0.5 if val >= 0 else bar.get_height() - 1
            va = 'bottom' if val >= 0 else 'top'
            ax1.text(bar.get_x() + bar.get_width()/2, y_pos,
                    f'${val:,.0f}', ha='center', va=va, fontsize=7)

        ax2 = ax1.twinx()
        ax2.plot(x_pos, financials['Operating Margin %'], color='#000000', marker='o',
                linewidth=2, markersize=5, label='Operating Margin %')
        ax2.set_ylabel('%', color='#000000')
        ax2.tick_params(axis='y', labelcolor='#000000')

        ax1.set_title('Operating Income & Margin', fontsize=12, fontweight='bold')
        lines1, labels1 = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=8)
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        charts.append(("Operating Income & Margin", buf))
        plt.close(fig)

        # Chart 4: Stock Price (2 years)
        price_data = fetch_stock_price_history(symbol, years=2)
        if price_data is not None and not price_data.empty:
            fig, ax = plt.subplots(figsize=(8, 3.5))
            ax.fill_between(price_data['Date'], price_data['Close'],
                           alpha=0.3, color='#4472C4')
            ax.plot(price_data['Date'], price_data['Close'],
                   color='#4472C4', linewidth=1.5)
            ax.set_title(f'{symbol} Stock Price (2 Years)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Date')
            ax.set_ylabel('Price ($)')
            ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, p: f'${x:,.0f}'))
            plt.xticks(rotation=45, ha='right', fontsize=8)
            plt.tight_layout()
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts.append(("Stock Price", buf))
            plt.close(fig)

    except Exception as e:
        # If chart generation fails, return empty list
        pass

    return charts


def create_pdf_document(content: str, symbol: str, ai_model: str) -> io.BytesIO:
    """Create PDF document and return as bytes"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=0.75*inch, leftMargin=0.75*inch,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)

    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                 fontSize=18, textColor=colors.HexColor('#2c3e50'),
                                 spaceAfter=6, alignment=TA_CENTER)

    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'],
                                    fontSize=11, textColor=colors.HexColor('#555555'),
                                    alignment=TA_CENTER, spaceAfter=20)

    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                   fontSize=12, textColor=colors.HexColor('#34495e'),
                                   spaceAfter=8, spaceBefore=12, fontName='Helvetica-Bold')

    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                fontSize=10, spaceAfter=6, leading=14)

    # Add logo if exists (1.2x larger)
    logo_path = Path(__file__).parent / "company_logo.png"
    if logo_path.exists():
        try:
            logo = Image(str(logo_path), width=3.6*inch, height=1.2*inch, kind='proportional')
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 0.2*inch))
        except Exception:
            pass

    # Title
    story.append(Paragraph(f"{symbol} Earnings Transcript Analysis", title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", subtitle_style))
    story.append(Spacer(1, 0.2*inch))

    # Content
    for line in content.split('\n'):
        if line.strip():
            # Check if line starts with markdown headers
            if line.strip().startswith('#'):
                clean_text = line.strip().lstrip('#').strip()
                story.append(Paragraph(clean_text, heading_style))
            else:
                # Handle **bold** markdown - convert to <b> tags for reportlab
                formatted_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                try:
                    story.append(Paragraph(formatted_line, body_style))
                except Exception:
                    # If parsing fails, add as plain text
                    story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), body_style))

    # Add charts section
    charts = create_pdf_charts(symbol)
    if charts:
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Financial Performance Charts", heading_style))
        story.append(Spacer(1, 0.2*inch))

        for chart_title, chart_buffer in charts:
            try:
                chart_buffer.seek(0)
                chart_img = Image(chart_buffer, width=5.5*inch, height=2.5*inch, kind='proportional')
                chart_img.hAlign = 'CENTER'
                story.append(chart_img)
                story.append(Spacer(1, 0.2*inch))
            except Exception:
                pass

    # Add Sequential Growth Rates Table
    try:
        financials = fetch_quarterly_financials(symbol)
        if financials is not None and not financials.empty and len(financials) > 1:
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph("Sequential Growth Rates (QoQ)", heading_style))
            story.append(Spacer(1, 0.1*inch))

            # Build table data
            table_data = [['Quarter', 'Revenue Growth', 'Gross Profit Growth', 'Operating Income Growth']]

            for i in range(1, len(financials)):
                prev = financials.iloc[i-1]
                curr = financials.iloc[i]

                rev_growth = ((curr['Revenue'] - prev['Revenue']) / prev['Revenue'] * 100) if prev['Revenue'] != 0 else 0
                gp_growth = ((curr['Gross Profit'] - prev['Gross Profit']) / prev['Gross Profit'] * 100) if prev['Gross Profit'] != 0 else 0

                if prev['Operating Income'] != 0:
                    op_growth = ((curr['Operating Income'] - prev['Operating Income']) / abs(prev['Operating Income']) * 100)
                else:
                    op_growth = 0

                table_data.append([
                    curr['Quarter'],
                    f"{rev_growth:+.1f}%",
                    f"{gp_growth:+.1f}%",
                    f"{op_growth:+.1f}%"
                ])

            # Create table with styling
            growth_table = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.8*inch])
            growth_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
                ('TOPPADDING', (0, 1), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')])
            ]))
            story.append(growth_table)
    except Exception:
        pass

    # Add signature at the end
    story.append(Spacer(1, 0.4*inch))
    signature_style = ParagraphStyle('Signature', parent=styles['Normal'],
                                     fontSize=10, spaceAfter=2, leading=14)
    story.append(Paragraph("David A Quinn", signature_style))
    story.append(Paragraph("Targeted Equity Consulting", signature_style))
    story.append(Paragraph("daquinn@targetedequityconsulting.com", signature_style))
    story.append(Paragraph("617-905-7415", signature_style))

    doc.build(story, onFirstPage=draw_page_border, onLaterPages=draw_page_border)
    buffer.seek(0)
    return buffer


def send_email_with_attachments(recipient_email: str, symbol: str, ai_model: str,
                                 pdf_buffer: io.BytesIO = None, word_buffer: io.BytesIO = None) -> tuple:
    """Send email with PDF and/or Word attachments."""
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        return False, "Email credentials not configured"

    try:
        # Clean credentials
        clean_email = EMAIL_ADDRESS.strip().replace('\xa0', '').encode('ascii', 'ignore').decode('ascii')
        clean_password = EMAIL_PASSWORD.strip().replace('\xa0', '').encode('ascii', 'ignore').decode('ascii')

        # Create message
        msg = MIMEMultipart()
        msg['From'] = clean_email
        msg['To'] = recipient_email
        msg['Subject'] = f"{symbol} Earnings Transcript Analysis - {datetime.now().strftime('%B %d, %Y')}"

        # Email body
        body_text = f"""Your {symbol} Earnings Transcript Analysis ({ai_model}) is attached.

Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}

Targeted Equity Consulting Group
Precision Analysis for Informed Investment Decisions
"""
        msg.attach(MIMEText(body_text, 'plain'))

        # Attach PDF if provided
        if pdf_buffer:
            pdf_buffer.seek(0)
            pdf_attachment = MIMEBase('application', 'octet-stream')
            pdf_attachment.set_payload(pdf_buffer.read())
            encoders.encode_base64(pdf_attachment)
            pdf_attachment.add_header('Content-Disposition', 'attachment',
                                     filename=f"{symbol}_transcript_analysis.pdf")
            msg.attach(pdf_attachment)

        # Attach Word if provided
        if word_buffer:
            word_buffer.seek(0)
            word_attachment = MIMEBase('application', 'octet-stream')
            word_attachment.set_payload(word_buffer.read())
            encoders.encode_base64(word_attachment)
            word_attachment.add_header('Content-Disposition', 'attachment',
                                      filename=f"{symbol}_transcript_analysis.docx")
            msg.attach(word_attachment)

        # Send via Gmail SMTP
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(clean_email, clean_password)
        server.sendmail(clean_email, recipient_email, msg.as_string())
        server.quit()

        return True, "Email sent successfully!"
    except Exception as e:
        return False, f"Email error: {str(e)}"


# Main App
st.title("📈 Earnings Transcript Analyzer")
st.markdown("Analyze earnings call transcripts with AI (Claude & ChatGPT)")

# Check API keys
missing_keys = []
if not FMP_API_KEY:
    missing_keys.append("FMP_API_KEY")
if not ANTHROPIC_API_KEY:
    missing_keys.append("ANTHROPIC_API_KEY")
if not OPENAI_API_KEY:
    missing_keys.append("OPENAI_API_KEY")

if missing_keys:
    st.error(f"Missing API keys: {', '.join(missing_keys)}")
    st.info("Add them to your .env file or Streamlit secrets.")
    st.stop()

# Sidebar
st.sidebar.header("Settings")
symbol = st.sidebar.text_input("Stock Ticker", value="AAPL", max_chars=10).upper()
num_quarters = st.sidebar.slider("Number of Quarters", 1, 8, 4)
ai_choice = st.sidebar.selectbox("AI Model", ["Both", "Claude Only", "ChatGPT Only"])

# Main content
if st.sidebar.button("🔍 Analyze Earnings", type="primary"):

    # Validate ticker symbol
    if not re.match(r'^[A-Z]{1,5}$', symbol):
        st.error(f"Invalid ticker symbol: '{symbol}' (expected 1-5 letters, e.g. AAPL)")
        st.stop()

    # Fetch company info
    with st.spinner(f"Fetching {symbol} company info..."):
        company_info = get_company_profile(symbol)

    if company_info:
        col1, col2, col3 = st.columns(3)
        col1.metric("Company", company_info.get('companyName', symbol))
        col2.metric("Sector", company_info.get('sector', 'N/A'))
        col3.metric("Industry", company_info.get('industry', 'N/A'))

    # Fetch transcripts
    with st.spinner(f"Fetching {num_quarters} quarters of transcripts..."):
        transcripts = fetch_transcripts(symbol, num_quarters)

    if not transcripts:
        st.error(f"No transcripts found for {symbol}")
        st.stop()

    # Show transcript info
    st.success(f"Found {len(transcripts)} transcripts")
    total_words = sum(t['word_count'] for t in transcripts)

    with st.expander("📄 Transcript Details"):
        for t in transcripts:
            st.write(f"**Q{t['quarter']} {t['year']}** - {t['date']} ({t['word_count']:,} words)")

    st.info(f"Total words to analyze: {total_words:,}")

    # Load prior analysis if available
    prior_analysis, prior_filename = load_prior_analysis(symbol)
    if prior_analysis:
        st.info(f"📋 Loaded prior analysis from: {prior_filename}")

    # Create prompt
    prompt = create_analysis_prompt(symbol, transcripts, company_info, prior_analysis)

    # Analyze
    claude_result = None
    chatgpt_result = None

    if ai_choice in ["Both", "Claude Only"]:
        with st.spinner("🤖 Claude is analyzing..."):
            try:
                claude_result = analyze_with_claude(prompt)
            except Exception as e:
                st.error(f"Claude error: {e}")

    chatgpt_note = None
    if ai_choice in ["Both", "ChatGPT Only"]:
        with st.spinner("🤖 ChatGPT is analyzing..."):
            try:
                chatgpt_result, chatgpt_note = analyze_with_chatgpt(symbol, transcripts, company_info, prior_analysis)
                if chatgpt_note:
                    st.info(f"ChatGPT {chatgpt_note}")
            except Exception as e:
                if "rate" in str(e).lower():
                    st.warning(f"ChatGPT rate limit exceeded. Using Claude only.")
                    if not claude_result and ai_choice == "ChatGPT Only":
                        st.info("Falling back to Claude...")
                        try:
                            claude_result = analyze_with_claude(prompt)
                        except Exception as ce:
                            st.error(f"Claude fallback error: {ce}")
                else:
                    st.error(f"ChatGPT error: {e}")

    # Store results in session state for persistence across button clicks
    st.session_state['claude_result'] = claude_result
    st.session_state['chatgpt_result'] = chatgpt_result
    st.session_state['analysis_symbol'] = symbol

    st.success("Analysis complete! See results below.")

    # Automatically send email with results
    if claude_result:
        with st.spinner("Sending email..."):
            word_doc = create_word_document(claude_result, symbol, "Claude")
            pdf_doc = create_pdf_document(claude_result, symbol, "Claude")
            success, message = send_email_with_attachments(
                "daquinn@targetedequityconsulting.com",
                symbol, "Claude", pdf_doc, word_doc
            )
            if success:
                st.success(f"📧 {message}")
            else:
                st.error(f"Email failed: {message}")

# Display results from session state (persists across button clicks)
claude_result = st.session_state.get('claude_result')
chatgpt_result = st.session_state.get('chatgpt_result')
analysis_symbol = st.session_state.get('analysis_symbol', symbol)

if claude_result or chatgpt_result:
    st.header("📊 Analysis Results")

    if claude_result and chatgpt_result:
        tab1, tab2 = st.tabs(["Claude Analysis", "ChatGPT Analysis"])

        with tab1:
            st.markdown(claude_result)
            word_doc = create_word_document(claude_result, analysis_symbol, "Claude")
            pdf_doc = create_pdf_document(claude_result, analysis_symbol, "Claude")

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "📥 Download Word",
                    word_doc,
                    file_name=f"{analysis_symbol}_claude_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_word_claude"
                )
            with col2:
                st.download_button(
                    "📥 Download PDF",
                    pdf_doc,
                    file_name=f"{analysis_symbol}_claude_analysis.pdf",
                    mime="application/pdf",
                    key="dl_pdf_claude"
                )

        with tab2:
            st.markdown(chatgpt_result)
            word_doc = create_word_document(chatgpt_result, analysis_symbol, "ChatGPT")
            pdf_doc = create_pdf_document(chatgpt_result, analysis_symbol, "ChatGPT")

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "📥 Download Word",
                    word_doc,
                    file_name=f"{analysis_symbol}_chatgpt_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_word_chatgpt"
                )
            with col2:
                st.download_button(
                    "📥 Download PDF",
                    pdf_doc,
                    file_name=f"{analysis_symbol}_chatgpt_analysis.pdf",
                    mime="application/pdf",
                    key="dl_pdf_chatgpt"
                )

    elif claude_result:
        st.markdown(claude_result)
        word_doc = create_word_document(claude_result, analysis_symbol, "Claude")
        pdf_doc = create_pdf_document(claude_result, analysis_symbol, "Claude")

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📥 Download Word",
                word_doc,
                file_name=f"{analysis_symbol}_claude_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_word_single"
            )
        with col2:
            st.download_button(
                "📥 Download PDF",
                pdf_doc,
                file_name=f"{analysis_symbol}_claude_analysis.pdf",
                mime="application/pdf",
                key="dl_pdf_single"
            )

    elif chatgpt_result:
        st.markdown(chatgpt_result)
        word_doc = create_word_document(chatgpt_result, analysis_symbol, "ChatGPT")
        pdf_doc = create_pdf_document(chatgpt_result, analysis_symbol, "ChatGPT")

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📥 Download Word",
                word_doc,
                file_name=f"{analysis_symbol}_chatgpt_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_word_gpt"
            )
        with col2:
            st.download_button(
                "📥 Download PDF",
                pdf_doc,
                file_name=f"{analysis_symbol}_chatgpt_analysis.pdf",
                mime="application/pdf",
                key="dl_pdf_gpt"
            )

    # Display financial charts at the bottom
    create_financial_charts(analysis_symbol)

else:
    st.info("👈 Enter a stock ticker and click 'Analyze Earnings' to start")

    # Show example
    st.markdown("""
    ### How it works:
    1. Enter any stock ticker (e.g., AAPL, NVDA, MSFT)
    2. Select number of quarters to analyze
    3. Choose AI model (Claude, ChatGPT, or both)
    4. Click Analyze and get detailed investment insights
    5. Download PDF or Word report, or email directly
    """)
