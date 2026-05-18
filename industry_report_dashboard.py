"""
Industry Report Dashboard - Streamlit Version
Generates comprehensive industry/sector analysis reports with AI-powered insights
"""
import streamlit as st
import os
import logging
from pathlib import Path

# Get the directory where this script is located
SCRIPT_DIR = Path(__file__).parent.resolve()

def get_logo_path():
    """Find the company logo path - works on local and Streamlit Cloud."""
    possible_paths = [
        SCRIPT_DIR / 'company_logo.png',
        Path('company_logo.png'),
        Path('/mount/src/create/company_logo.png'),
        Path('/mount/src/Create/company_logo.png'),
        Path('/app/company_logo.png'),
        Path(os.getcwd()) / 'company_logo.png',
    ]
    for p in possible_paths:
        if p.exists():
            return str(p)
    return None

LOGO_PATH = get_logo_path()

# Debug: List all files in script directory
def list_files_debug():
    files = []
    try:
        files = list(SCRIPT_DIR.glob('*.png'))
    except:
        pass
    return files

DEBUG_FILES = list_files_debug()
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
import pandas as pd
import json
from io import BytesIO
import anthropic
from openai import OpenAI
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, Image, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.colors import HexColor
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from pypdf import PdfReader, PdfWriter

load_dotenv()

# Load config for email - check Streamlit secrets, env vars, then config.json
def load_config():
    config = {}

    # Try config.json first
    try:
        with open("config.json", 'r') as f:
            config = json.load(f)
    except:
        pass

    # Override with Streamlit secrets if available
    try:
        if hasattr(st, 'secrets'):
            if 'EMAIL_ADDRESS' in st.secrets:
                config['email_address'] = st.secrets['EMAIL_ADDRESS']
            if 'EMAIL_PASSWORD' in st.secrets:
                config['password'] = st.secrets['EMAIL_PASSWORD']
            if 'email_address' in st.secrets:
                config['email_address'] = st.secrets['email_address']
            if 'email_password' in st.secrets:
                config['password'] = st.secrets['email_password']
    except:
        pass

    # Override with environment variables
    if os.getenv('EMAIL_ADDRESS'):
        config['email_address'] = os.getenv('EMAIL_ADDRESS')
    if os.getenv('EMAIL_PASSWORD'):
        config['password'] = os.getenv('EMAIL_PASSWORD')

    return config

CONFIG = load_config()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Initialize AI clients
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None
openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# Default stock universe path
DEFAULT_UNIVERSE_PATH = r"C:\Users\daqui\OneDrive\Documents\Targeted Equity Consulting Group\AI dashboard Data\master_universe.csv"

# Import from industry_report_generator
from industry_report_generator import (
    # Data fetching
    get_sector_performance,
    get_sector_pe_ratios,
    get_stocks_by_sector,
    get_stocks_by_industry,
    get_company_profile,
    # AI Analysis
    generate_industry_analysis,
    generate_market_view_analysis,
    identify_winners_losers,
    # Data classes
    ResearchNotes,
    ResearchFile,
    Article,
    WinnersLosersAnalysis,
    CompanyTrendPosition,
    create_research_notes,
    # PDF generation
    generate_industry_pdf,
    # Utilities
    format_currency,
)
from track_record import Pick, conviction_from_label, picks_from_winners_losers, record_picks
from evaluate_picks import evaluate_pending, summarize_track_record
from cost_tracker import CostTracker, BudgetExceeded
from industry_report_generator import deep_research_ticker, web_research_industry, TickerDeepDive
from dataclasses import dataclass, field

try:
    from estimates_tracker import EstimatesTracker
    _ESTIMATES_TRACKER_AVAILABLE = True
except Exception as _e:
    _ESTIMATES_TRACKER_AVAILABLE = False
    EstimatesTracker = None


PICKS_MARKER = "===PICKS_JSON==="
END_PICKS_MARKER = "===END_PICKS_JSON==="


def _parse_agent_output(agent_key: str, text: str) -> Tuple[str, List[Pick]]:
    """Split an agent's response into (prose, picks).

    Expected format (JSON-first, prose-after):
        ===PICKS_JSON===
        {"picks": [...]}
        ===END_PICKS_JSON===
        [prose analysis]

    If max_tokens truncates the response, prose gets cut (not the picks).
    Code fences are tolerated. Missing markers fall back to whole-text-as-prose.
    """
    if not text:
        return "", []

    start_idx = text.find(PICKS_MARKER)
    if start_idx < 0:
        return text.strip(), []

    after_start = start_idx + len(PICKS_MARKER)
    end_idx = text.find(END_PICKS_MARKER, after_start)
    if end_idx >= 0:
        json_part = text[after_start:end_idx]
        prose = text[end_idx + len(END_PICKS_MARKER):]
    else:
        # No end marker — common when max_tokens truncates mid-JSON.
        # Slice from first { to last } in the tail so we recover what we can.
        tail = text[after_start:]
        first = tail.find("{")
        last = tail.rfind("}")
        json_part = tail[first:last + 1] if first >= 0 and last > first else tail
        prose = ""

    json_part = json_part.strip()
    if json_part.startswith("```json"):
        json_part = json_part[7:]
    elif json_part.startswith("```"):
        json_part = json_part[3:]
    if json_part.endswith("```"):
        json_part = json_part[:-3]
    json_part = json_part.strip()

    picks: List[Pick] = []
    if json_part:
        try:
            data = json.loads(json_part)
        except json.JSONDecodeError as e:
            logger.warning(f"Agent {agent_key} picks JSON parse failed: {e}")
            data = {}
        for p in data.get("picks", []) or []:
            sym = (p.get("symbol") or "").strip().upper()
            direction = (p.get("direction") or "").strip().lower()
            if not sym or direction not in {"long", "short", "neutral", "avoid"}:
                continue
            picks.append(Pick(
                ticker=sym,
                agent=agent_key,
                direction=direction,
                thesis=(p.get("rationale") or "").strip(),
                trend=(p.get("trend") or "").strip(),
                conviction=conviction_from_label(p.get("confidence", "Medium")),
            ))

    return prose.strip(), picks


def _format_track_record_block(horizon_days: int = 90, since_days: int = 365) -> str:
    """Render per-agent hit-rate stats as a prompt-injectable block.

    Returns empty string when the DB is missing, empty, or unreadable — so
    a fresh install reads identically to today's prompt.
    """
    try:
        rows = summarize_track_record(horizon_days=horizon_days, since_days=since_days)
    except Exception as e:
        logger.warning(f"Track-record summary unavailable: {e}")
        return ""
    if not rows:
        return ""

    lines = [
        f"HISTORICAL AGENT TRACK RECORD ({horizon_days}-day horizon, alpha vs sector ETF, last {since_days}d):",
    ]
    for r in rows:
        lines.append(
            f"- {r['agent']}: n={r['n']}, hit_rate={r['hit_rate']*100:.0f}%, "
            f"avg_alpha={r['avg_alpha']*100:+.1f}%"
        )
    lines.append(
        "Use this track record when weighing the agent views below: give more weight to "
        "agents with higher hit rates and positive alpha. Don't ignore agents with poor "
        "records — they may still surface valid risks — but lean on proven ones when the views conflict."
    )
    return "\n".join(lines) + "\n\n"


def _format_revision_signals_block(tickers: List[str], top_n: int = 15) -> str:
    """Render 30-day EPS/revenue revision signals from the snapshot DB.

    Pulls per-ticker revisions via EstimatesTracker (local DB, no API
    calls), then surfaces the top_n upward and top_n downward movers by
    30d EPS revision %. Returns empty string when the tracker or data
    is unavailable, so report runs degrade gracefully.
    """
    if not _ESTIMATES_TRACKER_AVAILABLE or EstimatesTracker is None:
        return ""
    if not tickers:
        return ""

    try:
        tracker = EstimatesTracker()
    except Exception as e:
        logger.warning(f"EstimatesTracker init failed: {e}")
        return ""

    rows = []
    for t in tickers:
        try:
            s = tracker.get_revisions_summary(t, [30])
        except Exception:
            continue
        eps = s.get('eps_rev_30d')
        rev = s.get('rev_rev_30d')
        if eps is None and rev is None:
            continue
        rows.append({'ticker': t, 'eps_30d': eps, 'rev_30d': rev})

    if not rows:
        return ""

    scored = [r for r in rows if r['eps_30d'] is not None]
    if not scored:
        return ""

    scored.sort(key=lambda r: r['eps_30d'], reverse=True)
    ups = [r for r in scored if r['eps_30d'] > 0][:top_n]
    downs = [r for r in scored if r['eps_30d'] < 0]
    downs.sort(key=lambda r: r['eps_30d'])
    downs = downs[:top_n]

    if not ups and not downs:
        return ""

    def _fmt_row(r):
        eps_s = f"{r['eps_30d']:+.1f}%" if r['eps_30d'] is not None else "n/a"
        rev_s = f"{r['rev_30d']:+.1f}%" if r['rev_30d'] is not None else "n/a"
        return f"- {r['ticker']}: eps {eps_s}, rev {rev_s}"

    lines = [
        f"EARNINGS REVISION SIGNALS (30-day window, real revisions from snapshot DB, "
        f"comparing same fiscal period across time; covers {len(scored)} of "
        f"{len(tickers)} universe tickers):",
    ]
    if ups:
        lines.append("Top upward EPS revisions (estimates rising — bullish):")
        lines.extend(_fmt_row(r) for r in ups)
    if downs:
        lines.append("Top downward EPS revisions (estimates falling — bearish):")
        lines.extend(_fmt_row(r) for r in downs)
    lines.append(
        "Rising revisions often precede outperformance and falling revisions often "
        "precede underperformance — weight these signals when picking winners and losers, "
        "especially when multiple agents already lean the same direction."
    )
    return "\n".join(lines) + "\n\n"


# ============================================
# DOCUMENT READING FUNCTIONS
# ============================================

def read_word_document(file_buffer) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
        doc = Document(file_buffer)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading Word document: {e}")
        return ""


def read_pdf_document(file_buffer) -> str:
    """Extract text from a PDF document."""
    try:
        import pdfplumber
        full_text = []
        with pdfplumber.open(file_buffer) as pdf:
            for page in pdf.pages[:50]:  # Limit to 50 pages
                text = page.extract_text()
                if text:
                    full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading PDF document: {e}")
        return ""


def read_text_document(file_buffer) -> str:
    """Read a plain text file."""
    try:
        return file_buffer.read().decode('utf-8')
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return ""


def read_excel_document(file_buffer) -> str:
    """Extract text from an Excel spreadsheet."""
    try:
        # Read all sheets from the Excel file
        excel_data = pd.read_excel(file_buffer, sheet_name=None, header=None)
        full_text = []

        for sheet_name, df in excel_data.items():
            if df.empty:
                continue

            full_text.append(f"=== Sheet: {sheet_name} ===")

            # Check if first row looks like headers
            first_row = df.iloc[0].astype(str) if len(df) > 0 else []
            has_header = any(
                col.lower() in ['symbol', 'ticker', 'name', 'company', 'date', 'value', 'price', 'revenue', 'sector', 'industry']
                for col in first_row.values if pd.notna(col)
            )

            if has_header:
                # Re-read with headers
                file_buffer.seek(0)
                sheet_df = pd.read_excel(file_buffer, sheet_name=sheet_name)
                # Format as table
                full_text.append(sheet_df.to_string(index=False, max_rows=100))
            else:
                # Format without headers
                full_text.append(df.to_string(index=False, header=False, max_rows=100))

            full_text.append("")  # Empty line between sheets

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading Excel document: {e}")
        return ""


def read_uploaded_document(uploaded_file) -> str:
    """Read an uploaded document (Word, PDF, Excel, or TXT)."""
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()
    if filename.endswith('.docx'):
        return read_word_document(uploaded_file)
    elif filename.endswith('.pdf'):
        return read_pdf_document(uploaded_file)
    elif filename.endswith('.txt'):
        return read_text_document(uploaded_file)
    elif filename.endswith(('.xlsx', '.xls')):
        return read_excel_document(uploaded_file)
    else:
        return ""


def get_file_type(filename: str) -> str:
    """Determine file type from filename."""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return "pdf"
    elif filename_lower.endswith('.docx'):
        return "word"
    elif filename_lower.endswith(('.xlsx', '.xls')):
        return "excel"
    elif filename_lower.endswith('.txt'):
        return "text"
    return "unknown"


def summarize_research_file(
    filename: str,
    content: str,
    industry_context: str = "",
    ai_provider: str = "anthropic"
) -> str:
    """Generate AI summary of a research document."""
    if not content or len(content.strip()) < 100:
        return "Document too short to summarize."

    # Truncate content if too long
    max_chars = 12000
    truncated_content = content[:max_chars] if len(content) > max_chars else content

    context_note = f"\nIndustry Context: {industry_context}" if industry_context else ""

    prompt = f"""Analyze this research document and provide a concise summary.{context_note}

DOCUMENT: {filename}

CONTENT:
{truncated_content}

Please provide:
1. **Key Findings** (3-5 bullet points of the most important insights)
2. **Relevance** (1-2 sentences on how this relates to investment analysis)
3. **Data Points** (any key statistics, metrics, or figures mentioned)

Keep the summary focused and actionable for investment research purposes.
Format in markdown."""

    try:
        if ai_provider == "anthropic" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        elif ai_provider == "openai" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.choices[0].message.content
        else:
            return "AI summarization not available - no API key configured."
    except Exception as e:
        logger.error(f"Error summarizing research file {filename}: {e}")
        return f"Summary unavailable: {str(e)}"


def process_research_files(
    uploaded_files: list,
    industry_context: str = "",
    ai_provider: str = "anthropic",
    progress_callback=None
) -> List[ResearchFile]:
    """Process multiple uploaded research files and generate summaries."""
    research_files = []

    if not uploaded_files:
        return research_files

    total_files = len(uploaded_files)

    for i, uploaded_file in enumerate(uploaded_files):
        filename = uploaded_file.name
        file_type = get_file_type(filename)

        if progress_callback:
            progress_callback(f"Processing file {i+1}/{total_files}: {filename}", (i + 0.3) / total_files)

        # Read content
        uploaded_file.seek(0)
        content = read_uploaded_document(uploaded_file)

        if not content:
            logger.warning(f"Could not extract content from {filename}")
            research_files.append(ResearchFile(
                filename=filename,
                file_type=file_type,
                content="",
                summary="Could not extract content from this file."
            ))
            continue

        if progress_callback:
            progress_callback(f"Summarizing: {filename}", (i + 0.7) / total_files)

        # Generate summary
        summary = summarize_research_file(filename, content, industry_context, ai_provider)

        research_files.append(ResearchFile(
            filename=filename,
            file_type=file_type,
            content=content,
            summary=summary
        ))

    return research_files


# ============================================
# PDF GENERATION FOR WINNERS/LOSERS
# ============================================

def generate_winners_losers_word(
    industry_name: str,
    winners_losers: WinnersLosersAnalysis,
    original_file_buffer: BytesIO = None,
    note_content: str = None,
    research_files: List[ResearchFile] = None,
    user_directions: str = None
) -> BytesIO:
    """Generate Word document by copying original and appending winners/losers."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # If we have original file, copy it; otherwise create new
    if original_file_buffer:
        original_file_buffer.seek(0)
        doc = Document(original_file_buffer)
    else:
        doc = Document()
        # Add logo at very top for new documents
        if LOGO_PATH:
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(LOGO_PATH, width=Inches(4.5))
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")

        # Add title
        title = doc.add_heading(f'{industry_name} Analysis', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add user directions if provided
        if user_directions and user_directions.strip():
            doc.add_heading('Analysis Directions', level=1)
            directions_para = doc.add_paragraph()
            directions_para.add_run(user_directions).italic = True
            doc.add_paragraph()

        # Add the research note content if provided
        if note_content:
            doc.add_heading('Research Note', level=1)
            # Split content into paragraphs and add them
            for para_text in note_content.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

    # Add page break before Winners/Losers section
    doc.add_page_break()

    # Add company logo at top of Winners/Losers section
    if LOGO_PATH:
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(LOGO_PATH, width=Inches(4.5))
            doc.add_paragraph()  # Spacer
        except Exception as e:
            logger.warning(f"Could not add logo to Word doc: {e}")

    # Add Winners/Losers header
    header = doc.add_heading('Winners & Losers Analysis', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add summary if exists
    if winners_losers.summary:
        summary_para = doc.add_paragraph()
        summary_para.add_run(winners_losers.summary).italic = True
        doc.add_paragraph()

    # WINNERS section
    if winners_losers.winners:
        winners_heading = doc.add_heading('WINNERS', level=2)
        for run in winners_heading.runs:
            run.font.color.rgb = RGBColor(21, 87, 36)  # Dark green

        # Create winners table
        table = doc.add_table(rows=1, cols=3)
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass  # Style not available, use default

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Symbol'
        header_cells[1].text = 'Company'
        header_cells[2].text = 'Rationale'

        # Style header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'd4edda')
            cell._tc.get_or_add_tcPr().append(shading)

        # Add winner rows
        for w in winners_losers.winners:
            row = table.add_row().cells
            row[0].text = w.symbol
            row[1].text = w.company_name
            row[2].text = w.rationale

        doc.add_paragraph()

    # LOSERS section
    if winners_losers.losers:
        losers_heading = doc.add_heading('LOSERS', level=2)
        for run in losers_heading.runs:
            run.font.color.rgb = RGBColor(114, 28, 36)  # Dark red

        # Create losers table
        table = doc.add_table(rows=1, cols=3)
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass  # Style not available, use default

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Symbol'
        header_cells[1].text = 'Company'
        header_cells[2].text = 'Rationale'

        # Style header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'f8d7da')
            cell._tc.get_or_add_tcPr().append(shading)

        # Add loser rows
        for l in winners_losers.losers:
            row = table.add_row().cells
            row[0].text = l.symbol
            row[1].text = l.company_name
            row[2].text = l.rationale

    # Add Research Documents section if provided
    if research_files:
        doc.add_page_break()

        # Add logo at top of Research Documents section
        if LOGO_PATH:
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(LOGO_PATH, width=Inches(4.5))
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Could not add logo to research section: {e}")

        research_heading = doc.add_heading('Research Documents', level=1)
        research_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"{len(research_files)} document(s) analyzed")
        doc.add_paragraph()

        for rf in research_files:
            # File type icon text
            file_icon = {"pdf": "[PDF]", "word": "[WORD]", "excel": "[EXCEL]", "text": "[TXT]"}.get(rf.file_type, "[FILE]")

            # File heading
            file_heading = doc.add_heading(f"{file_icon} {rf.filename}", level=2)
            for run in file_heading.runs:
                run.font.size = Pt(12)

            # Summary
            if rf.summary:
                # Split summary into paragraphs
                for para_text in rf.summary.split('\n'):
                    if para_text.strip():
                        # Handle basic markdown
                        clean_text = para_text.strip()
                        p = doc.add_paragraph()
                        if clean_text.startswith('**') and '**' in clean_text[2:]:
                            # Bold text
                            clean_text = clean_text.replace('**', '')
                            run = p.add_run(clean_text)
                            run.bold = True
                        elif clean_text.startswith('- ') or clean_text.startswith('* '):
                            # Bullet point
                            p.add_run(clean_text)
                            p.style = 'List Bullet'
                        else:
                            p.add_run(clean_text)

            doc.add_paragraph()  # Spacer between files

    # Add signature at end
    doc.add_paragraph()
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = sig_para.add_run("David A Quinn")
    sig_run.bold = True

    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para.add_run("Targeted Equity Consulting Group")

    phone_para = doc.add_paragraph()
    phone_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    phone_para.add_run("617-905-7415")

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_winners_losers_pdf(
    industry_name: str,
    trends_data: Dict[str, Any],
    winners_losers: WinnersLosersAnalysis,
    original_note_content: str = None,
    user_directions: str = None
) -> BytesIO:
    """Generate simple PDF with just winners and losers tables (for appending to original)."""

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=5,
        textColor=HexColor('#1a1a2e')
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=8,
        textColor=HexColor('#2E86AB')
    )
    rationale_style = ParagraphStyle(
        'Rationale',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#333333'),
        leftIndent=10,
        spaceAfter=6,
        leading=11
    )

    elements = []

    # Add company logo if exists
    if LOGO_PATH:
        try:
            img = Image(LOGO_PATH, width=4.68*inch, height=1.56*inch)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(Spacer(1, 0.1*inch))
        except Exception as e:
            logger.warning(f"Could not load logo: {e}")

    # Add title
    elements.append(Paragraph(f"{industry_name} Analysis", title_style))
    elements.append(Spacer(1, 0.2*inch))

    # Add user directions if provided
    if user_directions and user_directions.strip():
        directions_heading_style = ParagraphStyle(
            'DirectionsHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=5,
            textColor=HexColor('#2E86AB')
        )
        directions_body_style = ParagraphStyle(
            'DirectionsBody',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#555555'),
            spaceAfter=15,
            leading=14,
            fontName='Helvetica-Oblique'
        )
        elements.append(Paragraph("Analysis Directions", directions_heading_style))
        # Escape special characters for ReportLab
        safe_directions = user_directions.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
        elements.append(Paragraph(safe_directions, directions_body_style))
        elements.append(Spacer(1, 0.1*inch))

    # Add research note content if provided
    if original_note_content:
        note_style = ParagraphStyle(
            'NoteStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#333333'),
            spaceAfter=8,
            leading=14
        )
        elements.append(Paragraph("Research Note", heading_style))
        elements.append(Spacer(1, 0.1*inch))
        # Split content into paragraphs
        for para_text in original_note_content.split('\n'):
            if para_text.strip():
                # Escape special characters for ReportLab
                safe_text = para_text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(safe_text, note_style))
        elements.append(Spacer(1, 0.3*inch))

    # Style for wrapped text in table cells
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=11
    )

    # Winners Table
    if winners_losers.winners:
        elements.append(Paragraph("WINNERS", heading_style))

        winner_data = [["Symbol", "Company", "Rationale"]]
        for w in winners_losers.winners:
            # Use Paragraph for rationale to enable text wrapping
            rationale_para = Paragraph(w.rationale, cell_style)
            winner_data.append([
                w.symbol,
                (w.company_name[:30] + "...") if len(w.company_name) > 30 else w.company_name,
                rationale_para
            ])

        winner_table = Table(winner_data, colWidths=[55, 120, 285])
        winner_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#d4edda')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#155724')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (1, -1), 'LEFT'),
            ('ALIGN', (2, 0), (2, -1), 'LEFT'),
            ('BOX', (0, 0), (-1, -1), 1.5, HexColor('#155724')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor('#c3e6cb')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor('#f8fff8')]),
        ]))
        elements.append(winner_table)
        elements.append(Spacer(1, 20))

    # Losers Table
    if winners_losers.losers:
        elements.append(Paragraph("LOSERS", heading_style))

        loser_data = [["Symbol", "Company", "Rationale"]]
        for l in winners_losers.losers:
            # Use Paragraph for rationale to enable text wrapping
            rationale_para = Paragraph(l.rationale, cell_style)
            loser_data.append([
                l.symbol,
                (l.company_name[:30] + "...") if len(l.company_name) > 30 else l.company_name,
                rationale_para
            ])

        loser_table = Table(loser_data, colWidths=[55, 120, 285])
        loser_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f8d7da')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#721c24')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (1, -1), 'LEFT'),
            ('ALIGN', (2, 0), (2, -1), 'LEFT'),
            ('BOX', (0, 0), (-1, -1), 1.5, HexColor('#721c24')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor('#f5c6cb')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor('#fff8f8')]),
        ]))
        elements.append(loser_table)

    # Add signature at end
    signature_style = ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=HexColor('#333333')
    )
    elements.append(Spacer(1, 0.5*inch))
    elements.append(Paragraph("<b>David A Quinn</b>", signature_style))
    elements.append(Paragraph("Targeted Equity Consulting Group", signature_style))
    elements.append(Paragraph("617-905-7415", signature_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_winners_losers_appendix_pdf(
    industry_name: str,
    winners_losers: WinnersLosersAnalysis
) -> BytesIO:
    """Generate a Winners/Losers PDF appendix (no research note - for merging with original)."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=8,
        textColor=HexColor('#2E86AB')
    )
    signature_style = ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=HexColor('#333333')
    )

    elements = []

    # Add logo
    if LOGO_PATH:
        try:
            logger.info(f"Adding logo from: {LOGO_PATH}")
            img = Image(LOGO_PATH, width=4.68*inch, height=1.56*inch)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(Spacer(1, 0.2*inch))
            logger.info("Logo added successfully to appendix PDF")
        except Exception as e:
            logger.error(f"Could not load logo: {e}")
            import traceback
            logger.error(traceback.format_exc())
    else:
        logger.warning("LOGO_PATH is None - no logo will be added")

    # Style for wrapped text in table cells
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=11
    )

    # Winners Table
    if winners_losers.winners:
        elements.append(Paragraph("WINNERS", heading_style))
        winner_data = [["Symbol", "Company", "Rationale"]]
        for w in winners_losers.winners:
            # Use Paragraph for rationale to enable text wrapping
            rationale_para = Paragraph(w.rationale, cell_style)
            winner_data.append([
                w.symbol,
                (w.company_name[:30] + "...") if len(w.company_name) > 30 else w.company_name,
                rationale_para
            ])
        winner_table = Table(winner_data, colWidths=[55, 120, 285])
        winner_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#d4edda')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#155724')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, -1), 9),
            ('BOX', (0, 0), (-1, -1), 1.5, HexColor('#155724')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor('#c3e6cb')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(winner_table)
        elements.append(Spacer(1, 20))

    # Losers Table
    if winners_losers.losers:
        elements.append(Paragraph("LOSERS", heading_style))
        loser_data = [["Symbol", "Company", "Rationale"]]
        for l in winners_losers.losers:
            # Use Paragraph for rationale to enable text wrapping
            rationale_para = Paragraph(l.rationale, cell_style)
            loser_data.append([
                l.symbol,
                (l.company_name[:30] + "...") if len(l.company_name) > 30 else l.company_name,
                rationale_para
            ])
        loser_table = Table(loser_data, colWidths=[55, 120, 285])
        loser_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f8d7da')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#721c24')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, -1), 9),
            ('BOX', (0, 0), (-1, -1), 1.5, HexColor('#721c24')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor('#f5c6cb')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(loser_table)

    # Signature
    elements.append(Spacer(1, 0.5*inch))
    elements.append(Paragraph("<b>David A Quinn</b>", signature_style))
    elements.append(Paragraph("Targeted Equity Consulting Group", signature_style))
    elements.append(Paragraph("617-905-7415", signature_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_cover_page_pdf(industry_name: str) -> BytesIO:
    """Generate a cover page with logo and title."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=1.5*inch,
        bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=20,
        textColor=HexColor('#1a1a2e')
    )
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        alignment=TA_CENTER,
        textColor=HexColor('#666666')
    )

    elements = []

    # Add logo at top
    if LOGO_PATH:
        try:
            img = Image(LOGO_PATH, width=5*inch, height=1.67*inch)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(Spacer(1, 0.5*inch))
        except Exception as e:
            logger.warning(f"Could not add logo to cover: {e}")

    # Add title
    elements.append(Spacer(1, 1*inch))
    elements.append(Paragraph(f"{industry_name}", title_style))
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph("Industry Analysis Report", subtitle_style))
    elements.append(Spacer(1, 0.5*inch))

    from datetime import datetime
    elements.append(Paragraph(f"{datetime.now().strftime('%B %d, %Y')}", subtitle_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def merge_pdf_with_appendix(
    original_pdf_bytes: bytes,
    appendix_buffer: BytesIO,
    industry_name: str = "Industry"
) -> BytesIO:
    """Merge cover page + original PDF + Winners/Losers appendix."""
    output = BytesIO()
    writer = PdfWriter()

    # Add cover page with logo first
    cover_buffer = generate_cover_page_pdf(industry_name)
    cover_reader = PdfReader(cover_buffer)
    for page in cover_reader.pages:
        writer.add_page(page)

    # Add original PDF pages
    original_reader = PdfReader(BytesIO(original_pdf_bytes))
    for page in original_reader.pages:
        writer.add_page(page)

    # Add appendix pages (Winners/Losers)
    appendix_buffer.seek(0)
    appendix_reader = PdfReader(appendix_buffer)
    for page in appendix_reader.pages:
        writer.add_page(page)

    writer.write(output)
    output.seek(0)
    return output


def send_email_with_attachment(
    file_buffer: BytesIO,
    industry_name: str,
    file_type: str = "pdf",  # "pdf" or "docx"
    recipient_email: str = None
) -> Tuple[bool, str]:
    """Send email with PDF or Word attachment."""
    try:
        email_address = CONFIG.get('email_address')
        password = CONFIG.get('password')
        smtp_server = CONFIG.get('smtp_server', 'smtp.gmail.com')
        smtp_port = CONFIG.get('smtp_port', 587)
        recipient = recipient_email or CONFIG.get('email_recipient', email_address)

        if not email_address or not password:
            return False, "Email credentials not configured in config.json"

        # Create message
        msg = MIMEMultipart()
        msg['From'] = email_address
        msg['To'] = recipient
        msg['Subject'] = f"{industry_name} - Winners & Losers Analysis"

        # Email body
        body = f"""
{industry_name} Analysis Report

Please find attached the Winners & Losers analysis report.

Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

---
Targeted Equity Consulting Group
Precision Analysis for Informed Investment Decisions
        """
        msg.attach(MIMEText(body, 'plain'))

        # Attach file
        file_buffer.seek(0)
        attachment = MIMEBase('application', 'octet-stream')
        attachment.set_payload(file_buffer.read())
        encoders.encode_base64(attachment)

        safe_name = industry_name.replace(' ', '_').replace('/', '_')
        date_str = datetime.now().strftime('%Y%m%d')
        extension = "docx" if file_type == "docx" else "pdf"
        attachment.add_header(
            'Content-Disposition',
            'attachment',
            filename=f"{safe_name}_Winners_Losers_{date_str}.{extension}"
        )
        msg.attach(attachment)

        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_address, password)
        server.sendmail(email_address, recipient, msg.as_string())
        server.quit()

        return True, f"Email sent successfully to {recipient}!"

    except Exception as e:
        logger.error(f"Email error: {e}")
        return False, f"Email error: {str(e)}"


def send_email_with_pdf(
    pdf_buffer: BytesIO,
    industry_name: str,
    recipient_email: str = None
) -> Tuple[bool, str]:
    """Send email with PDF attachment (backwards compatible wrapper)."""
    return send_email_with_attachment(pdf_buffer, industry_name, "pdf", recipient_email)


# ============================================
# UNIVERSE SCANNING FUNCTIONS
# ============================================

def load_stock_universe(file_path: str = None, uploaded_file=None) -> pd.DataFrame:
    """Load stock universe from CSV file."""
    try:
        if uploaded_file is not None:
            df = pd.read_csv(uploaded_file, header=None)
        elif file_path and os.path.exists(file_path):
            df = pd.read_csv(file_path, header=None)
        else:
            return pd.DataFrame()

        # Check if first row looks like headers
        first_row = df.iloc[0].astype(str)
        has_header = any(col.lower() in ['symbol', 'ticker', 'name', 'company', 'exchange']
                        for col in first_row.values)

        if has_header:
            # Re-read with headers
            if uploaded_file is not None:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file)
            elif file_path:
                df = pd.read_csv(file_path)
        else:
            # Assign default column names: Symbol, Company, Exchange
            if len(df.columns) >= 3:
                df.columns = ['Symbol', 'Company', 'Exchange'] + [f'Col{i}' for i in range(3, len(df.columns))]
            elif len(df.columns) == 2:
                df.columns = ['Symbol', 'Company']
            elif len(df.columns) == 1:
                df.columns = ['Symbol']

        return df
    except Exception as e:
        logger.error(f"Error loading stock universe: {e}")
        return pd.DataFrame()


def extract_trends_from_note(note_content: str, ai_provider: str = "anthropic") -> Dict[str, Any]:
    """Extract key trends and themes from an industry research note using AI."""

    prompt = f"""Analyze this industry research note and extract the following in JSON format:

RESEARCH NOTE:
{note_content[:15000]}  # Limit content length

Please extract and return as JSON:
{{
    "industry": "The main industry or sector discussed",
    "key_trends": ["List of 3-5 key industry trends identified"],
    "bullish_factors": ["Factors that would benefit companies"],
    "bearish_factors": ["Factors that would hurt companies"],
    "key_themes": ["Main investment themes"],
    "summary": "2-3 sentence summary of the note"
}}

Return ONLY valid JSON, no other text."""

    try:
        if ai_provider == "anthropic" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.content[0].text
        elif ai_provider == "openai" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.choices[0].message.content
        else:
            return {"error": "No AI provider available"}

        # Clean up response
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        return json.loads(response_text.strip())
    except Exception as e:
        logger.error(f"Error extracting trends: {e}")
        return {"error": str(e)}


# ============================================
# MULTI-AGENT ANALYSIS SYSTEM
# ============================================

ANALYSIS_AGENTS = {
    "industry_analyst": {
        "name": "Industry Analyst",
        "prompt": """You are a Senior Industry Analyst. Read this research report and identify:
1. Key industry trends and shifts
2. Which business models benefit from these trends
3. Which business models are threatened
4. Market dynamics and competitive changes

Focus on structural changes that create lasting winners and losers."""
    },
    "competitive_intel": {
        "name": "Competitive Intelligence",
        "prompt": """You are a Competitive Intelligence Analyst. Based on this report, identify:
1. Companies gaining competitive advantage
2. Companies losing market position
3. New entrants or disruptors benefiting
4. Incumbents being disrupted

Focus on competitive positioning changes."""
    },
    "financial_analyst": {
        "name": "Financial Analyst",
        "prompt": """You are a Financial Analyst. Based on this report, identify:
1. Companies with financial strength to capitalize on trends
2. Companies with balance sheet risk from these changes
3. Revenue/margin implications for different players
4. Capital requirements and who can fund growth

Focus on financial capacity to win or lose."""
    },
    "valuation_analyst": {
        "name": "Valuation Analyst",
        "prompt": """You are a Valuation Analyst. Based on this report, identify:
1. Companies trading at attractive valuations relative to their positioning
2. Companies that appear overvalued given the risks they face
3. Valuation disconnects - where the market is mispricing the trends
4. Multiple expansion/compression candidates

Focus on price vs. fundamental value given the trends."""
    },
    "technical_momentum": {
        "name": "Technical & Momentum Analyst",
        "prompt": """You are a Technical and Momentum Analyst. Based on this report, identify:
1. Companies likely to see positive earnings revisions and momentum
2. Companies facing negative estimate revisions
3. Sentiment shifts that could drive price action
4. Stocks where fundamentals and technicals could align

Focus on near-term catalysts and momentum factors."""
    },
    "supply_chain_analyst": {
        "name": "Supply Chain Analyst",
        "prompt": """You are a Supply Chain Analyst. Based on this report, identify:
1. Companies with supply chain advantages (vertical integration, key suppliers)
2. Companies vulnerable to supply chain disruption
3. Winners/losers from reshoring, nearshoring, or sourcing shifts
4. Input cost winners and losers

Focus on supply chain positioning and vulnerabilities."""
    },
    "esg_analyst": {
        "name": "ESG & Sustainability Analyst",
        "prompt": """You are an ESG and Sustainability Analyst. Based on this report, identify:
1. Companies positioned to benefit from ESG/sustainability trends
2. Companies facing ESG-related headwinds or regulatory risk
3. Transition winners in clean energy, emissions reduction
4. Companies with governance or social risk exposure

Focus on ESG factors that create investment alpha."""
    },
    "macro_policy_analyst": {
        "name": "Macro & Policy Analyst",
        "prompt": """You are a Macro and Policy Analyst. Based on this report, identify:
1. Companies benefiting from fiscal policy, subsidies, or government spending
2. Companies at risk from regulatory changes or policy shifts
3. Interest rate sensitivity - winners and losers
4. Geopolitical exposure and trade policy impacts

Focus on macro and policy factors affecting specific companies."""
    },
    "ma_strategist": {
        "name": "M&A Strategist",
        "prompt": """You are an M&A and Corporate Strategy Analyst. Based on this report, identify:
1. Likely acquirers who could consolidate the industry
2. Attractive acquisition targets given the trends
3. Companies that need to do deals to remain competitive
4. Spinoff or divestiture candidates

Focus on corporate actions that could unlock or destroy value."""
    },
    "risk_analyst": {
        "name": "Risk Analyst",
        "prompt": """You are a Risk Analyst. Based on this report, identify:
1. Companies most exposed to negative trends
2. Regulatory or legal risks for specific players
3. Technology disruption risks
4. Business model obsolescence risks

Focus on downside risks and vulnerabilities."""
    },
    "investment_strategist": {
        "name": "Investment Strategist",
        "prompt": """You are a Chief Investment Strategist. Synthesize the analysis and provide:
1. Clear WINNERS list with conviction levels
2. Clear LOSERS list with conviction levels
3. Key catalysts to watch
4. Timeframe for thesis to play out

Be decisive - pick clear winners and losers."""
    }
}


def run_agent_analysis(
    agent_key: str,
    report_content: str,
    universe_stocks: str,
    ai_provider: str = "anthropic",
    user_directions: str = ""
) -> Tuple[str, List[Pick]]:
    """Run a single agent and return (prose, structured_picks)."""
    agent = ANALYSIS_AGENTS[agent_key]

    # Include user directions if provided
    directions_section = ""
    if user_directions and user_directions.strip():
        directions_section = f"""

USER ANALYSIS DIRECTIONS (PAY SPECIAL ATTENTION TO THESE):
{user_directions}

"""

    prompt = f"""{agent['prompt']}{directions_section}

RESEARCH REPORT:
{report_content[:12000]}

STOCK UNIVERSE TO CONSIDER:
{universe_stocks}

RESPONSE FORMAT — you MUST follow this exactly:

1. Start your response with the JSON picks block between these markers:

{PICKS_MARKER}
{{
  "picks": [
    {{"symbol": "TICKER", "direction": "long",    "rationale": "1-2 sentence reason", "confidence": "High",   "trend": "trend driving this call"}},
    {{"symbol": "TICKER", "direction": "short",   "rationale": "1-2 sentence reason", "confidence": "Medium", "trend": "trend driving this call"}},
    {{"symbol": "TICKER", "direction": "neutral", "rationale": "1-2 sentence reason", "confidence": "Low",    "trend": "trend driving this call"}}
  ]
}}
{END_PICKS_MARKER}

2. Then write your full prose analysis identifying specific winners and losers, supporting evidence, and structural drivers.

Rules: only use tickers from the STOCK UNIVERSE; direction is "long" for winners, "short" for losers, "neutral" for mixed; confidence is High/Medium/Low. Picks JSON must be valid JSON with no code fences. Do not include any text before the {PICKS_MARKER} marker."""

    stop_reason = None
    in_tok = out_tok = None
    try:
        if ai_provider == "anthropic" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}]
            )
            text = response.content[0].text
            stop_reason = getattr(response, "stop_reason", None)
            usage = getattr(response, "usage", None)
            if usage is not None:
                in_tok = getattr(usage, "input_tokens", None)
                out_tok = getattr(usage, "output_tokens", None)
        elif ai_provider == "openai" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}]
            )
            text = response.choices[0].message.content
            stop_reason = getattr(response.choices[0], "finish_reason", None)
            usage = getattr(response, "usage", None)
            if usage is not None:
                in_tok = getattr(usage, "prompt_tokens", None)
                out_tok = getattr(usage, "completion_tokens", None)
        else:
            return "AI provider not available", []
    except Exception as e:
        logger.error(f"Agent {agent_key} failed: {e}")
        return f"Error: {str(e)}", []

    prose, picks = _parse_agent_output(agent_key, text)
    logger.info(
        f"Agent {agent_key}: stop={stop_reason} in={in_tok} out={out_tok} "
        f"picks={len(picks)} marker={'YES' if PICKS_MARKER in text else 'NO'}"
    )
    if stop_reason == "max_tokens":
        logger.warning(f"Agent {agent_key} hit max_tokens — prose was truncated.")
    return prose, picks


def run_all_agents_and_synthesize(
    report_content: str,
    universe_df: pd.DataFrame = None,
    ai_provider: str = "anthropic",
    progress_callback=None,
    user_directions: str = "",
    theme: str = "",
    sector: Optional[str] = None,
    record_to_track_record: bool = True,
) -> WinnersLosersAnalysis:
    """Run all agents and synthesize into final winners/losers.

    When `record_to_track_record` is True and `theme` is provided, the
    synthesis output is also persisted via record_picks() so future runs
    can compute hit rates.
    """

    # Score any matured picks so the track-record block injected into this
    # run's synthesis prompt is as fresh as possible. Wrapped in try/except
    # because a missing FMP key or transient network issue must not block
    # report generation.
    try:
        evaluate_pending()
    except Exception as eval_err:
        logger.warning(f"Pre-report evaluator pass failed (continuing anyway): {eval_err}")

    # Build stock list if universe provided, otherwise analyze report directly
    universe_tickers: List[str] = []
    if universe_df is not None and not universe_df.empty:
        symbol_col = universe_df.columns[0]
        name_col = universe_df.columns[1] if len(universe_df.columns) > 1 else symbol_col
        stocks_list = []
        for _, row in universe_df.head(150).iterrows():
            sym = str(row[symbol_col]).strip()
            stocks_list.append(f"- {sym} ({row[name_col]})")
            if sym:
                universe_tickers.append(sym)
        universe_stocks = "\n".join(stocks_list)
    else:
        universe_stocks = "(Identify winners/losers directly from the companies mentioned in the report)"

    # Pull 30-day EPS/revenue revision signals from the local snapshot DB
    # so every agent sees the same earnings-revision priors.
    revision_signals_section = _format_revision_signals_block(universe_tickers, top_n=15)
    if revision_signals_section:
        universe_stocks = f"{universe_stocks}\n\n{revision_signals_section}"

    # Run each agent
    agent_results: Dict[str, str] = {}
    agent_picks: Dict[str, List[Pick]] = {}
    agents = list(ANALYSIS_AGENTS.keys())

    for i, agent_key in enumerate(agents):
        if progress_callback:
            progress_callback(f"Running {ANALYSIS_AGENTS[agent_key]['name']}...", (i + 1) / (len(agents) + 1))
        prose, picks = run_agent_analysis(agent_key, report_content, universe_stocks, ai_provider, user_directions)
        agent_results[agent_key] = prose
        agent_picks[agent_key] = picks

    # Final synthesis
    if progress_callback:
        progress_callback("Synthesizing final winners/losers...", 0.95)

    # Include user directions in synthesis if provided
    user_directions_section = ""
    if user_directions and user_directions.strip():
        user_directions_section = f"""
USER ANALYSIS DIRECTIONS (ENSURE THESE ARE ADDRESSED):
{user_directions}

"""

    track_record_section = _format_track_record_block(horizon_days=90, since_days=365)

    synthesis_prompt = f"""Based on the following multi-agent analysis of the research report, create the final WINNERS and LOSERS list.
{user_directions_section}{track_record_section}{revision_signals_section}INDUSTRY ANALYST VIEW:
{agent_results.get('industry_analyst', 'N/A')}

COMPETITIVE INTELLIGENCE VIEW:
{agent_results.get('competitive_intel', 'N/A')}

FINANCIAL ANALYST VIEW:
{agent_results.get('financial_analyst', 'N/A')}

VALUATION ANALYST VIEW:
{agent_results.get('valuation_analyst', 'N/A')}

TECHNICAL & MOMENTUM VIEW:
{agent_results.get('technical_momentum', 'N/A')}

SUPPLY CHAIN ANALYST VIEW:
{agent_results.get('supply_chain_analyst', 'N/A')}

ESG ANALYST VIEW:
{agent_results.get('esg_analyst', 'N/A')}

MACRO & POLICY VIEW:
{agent_results.get('macro_policy_analyst', 'N/A')}

M&A STRATEGIST VIEW:
{agent_results.get('ma_strategist', 'N/A')}

RISK ANALYST VIEW:
{agent_results.get('risk_analyst', 'N/A')}

INVESTMENT STRATEGIST VIEW:
{agent_results.get('investment_strategist', 'N/A')}

Now synthesize into a final JSON output:
{{
    "summary": "2-3 sentence synthesis of the consensus view",
    "winners": [
        {{"symbol": "TICKER", "company_name": "Name", "trend": "Key trend", "rationale": "Why they win", "confidence": "High/Medium/Low"}}
    ],
    "losers": [
        {{"symbol": "TICKER", "company_name": "Name", "trend": "Key trend", "rationale": "Why they lose", "confidence": "High/Medium/Low"}}
    ]
}}

Include stocks where multiple agents agree. Return ONLY valid JSON."""

    try:
        if ai_provider == "anthropic" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=8000,
                messages=[{"role": "user", "content": synthesis_prompt}]
            )
            response_text = response.content[0].text
        elif ai_provider == "openai" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=8000,
                messages=[{"role": "user", "content": synthesis_prompt}]
            )
            response_text = response.choices[0].message.content
        else:
            return WinnersLosersAnalysis(summary="No AI provider available")

        # Parse JSON
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        data = json.loads(response_text.strip())

        winners = [
            CompanyTrendPosition(
                symbol=w.get('symbol', ''),
                company_name=w.get('company_name', ''),
                position='winner',
                trend=w.get('trend', ''),
                rationale=w.get('rationale', ''),
                confidence=w.get('confidence', 'Medium')
            )
            for w in data.get('winners', [])
        ]

        losers = [
            CompanyTrendPosition(
                symbol=l.get('symbol', ''),
                company_name=l.get('company_name', ''),
                position='loser',
                trend=l.get('trend', ''),
                rationale=l.get('rationale', ''),
                confidence=l.get('confidence', 'Medium')
            )
            for l in data.get('losers', [])
        ]

        result = WinnersLosersAnalysis(
            winners=winners,
            losers=losers,
            summary=data.get('summary', '')
        )

        if record_to_track_record and theme:
            try:
                tickers_universe: List[str] = []
                if universe_df is not None and not universe_df.empty:
                    tickers_universe = [str(s).upper() for s in universe_df.iloc[:, 0].tolist()]
                else:
                    tickers_universe = [w.symbol for w in winners] + [l.symbol for l in losers]

                # Union every ticker that appears in any agent's picks or the synthesis.
                all_symbols = {w.symbol for w in winners} | {l.symbol for l in losers}
                for picks_list in agent_picks.values():
                    for p in picks_list:
                        all_symbols.add(p.ticker)

                price_by_symbol: Dict[str, Optional[float]] = {}
                for sym in all_symbols:
                    try:
                        profile = get_company_profile(sym) or {}
                        price_by_symbol[sym] = profile.get("price")
                    except Exception:
                        price_by_symbol[sym] = None
                companies_for_pricing = [
                    {"symbol": s, "price": price_by_symbol.get(s)} for s in all_symbols
                ]

                synthesis_picks = picks_from_winners_losers(
                    result, agent="synthesis", companies=companies_for_pricing
                )

                # Backfill entry_price on every per-agent pick from the same FMP snapshot.
                combined_picks: List[Pick] = []
                for picks_list in agent_picks.values():
                    for p in picks_list:
                        if p.entry_price is None:
                            p.entry_price = price_by_symbol.get(p.ticker)
                        combined_picks.append(p)
                combined_picks.extend(synthesis_picks)

                record_picks(
                    theme=theme,
                    tickers=tickers_universe,
                    picks=combined_picks,
                    sector=sector,
                    model="claude-sonnet-4-6" if ai_provider == "anthropic" else "gpt-4o-mini",
                    prompt_version="agents-v1",
                )
            except Exception as rec_err:
                logger.warning(f"Failed to record picks to track record: {rec_err}")

        return result

    except Exception as e:
        logger.error(f"Synthesis failed: {e}")
        return WinnersLosersAnalysis(summary=f"Synthesis error: {str(e)}")


def critique_and_refine_synthesis(
    initial: WinnersLosersAnalysis,
    agent_results: Dict[str, str],
    ai_provider: str = "anthropic",
    cost_tracker=None,
) -> WinnersLosersAnalysis:
    """Run a devil's-advocate critique against the initial synthesis, then
    produce a revised synthesis incorporating the critique.

    Returns the revised WinnersLosersAnalysis. On any failure, returns the
    original `initial` unchanged so the pipeline degrades gracefully.
    """
    if ai_provider != "anthropic" or not anthropic_client:
        return initial

    initial_text = (
        f"SUMMARY: {initial.summary}\n\n"
        f"WINNERS:\n" + "\n".join(
            f"- {w.symbol} ({w.company_name}) [{w.confidence}] {w.trend}: {w.rationale}"
            for w in initial.winners
        )
        + "\n\nLOSERS:\n" + "\n".join(
            f"- {l.symbol} ({l.company_name}) [{l.confidence}] {l.trend}: {l.rationale}"
            for l in initial.losers
        )
    )

    critique_prompt = f"""You are a skeptical senior portfolio manager doing a pre-publication review of an analyst's winners/losers call.

PROPOSED CALL:
{initial_text}

UNDERLYING AGENT VIEWS (for your reference):
""" + "\n\n".join(f"=== {k} ===\n{v[:1500]}" for k, v in agent_results.items()) + """

Your job: surface the strongest counter-arguments. For each ticker in the proposed call, list 1-3 specific challenges:
- What evidence weakens this thesis?
- What's the strongest contrarian view?
- Is the confidence level overstated given the evidence?
- Are there obvious risks the analyst is glossing over?

Be sharp and specific. Cite numbers from the agent views. Do not rewrite the call yet — only critique it. Output prose, no JSON, no recap of the original."""

    try:
        critique_resp = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=3000,
            messages=[{"role": "user", "content": critique_prompt}],
        )
        critique_text = critique_resp.content[0].text
        if cost_tracker is not None:
            cost_tracker.record_anthropic("synthesis_critique", critique_resp)
    except Exception as e:
        logger.warning(f"Critique pass failed (keeping initial synthesis): {e}")
        return initial

    refine_prompt = f"""You are revising the winners/losers call based on a critique from a senior PM.

ORIGINAL CALL:
{initial_text}

PM CRITIQUE:
{critique_text}

Produce the revised call. Adjust confidence levels where the critique was valid; remove or move names where the thesis didn't survive; sharpen rationales by addressing the strongest counter-arguments. Return ONLY valid JSON in this shape, no other text:

{{
    "summary": "2-3 sentence synthesis incorporating critique",
    "winners": [
        {{"symbol": "TICKER", "company_name": "Name", "trend": "Key trend", "rationale": "Revised rationale addressing critique", "confidence": "High/Medium/Low"}}
    ],
    "losers": [
        {{"symbol": "TICKER", "company_name": "Name", "trend": "Key trend", "rationale": "Revised rationale addressing critique", "confidence": "High/Medium/Low"}}
    ]
}}"""

    try:
        refine_resp = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=3000,
            messages=[{"role": "user", "content": refine_prompt}],
        )
        refine_text = refine_resp.content[0].text
        if cost_tracker is not None:
            cost_tracker.record_anthropic("synthesis_refine", refine_resp)
    except Exception as e:
        logger.warning(f"Refine pass failed (keeping initial synthesis): {e}")
        return initial

    refine_text = refine_text.strip()
    if refine_text.startswith("```json"):
        refine_text = refine_text[7:]
    elif refine_text.startswith("```"):
        refine_text = refine_text[3:]
    if refine_text.endswith("```"):
        refine_text = refine_text[:-3]
    refine_text = refine_text.strip()

    try:
        data = json.loads(refine_text)
    except json.JSONDecodeError as e:
        logger.warning(f"Refine JSON parse failed (keeping initial): {e}")
        return initial

    winners = [
        CompanyTrendPosition(
            symbol=w.get("symbol", ""),
            company_name=w.get("company_name", ""),
            position="winner",
            trend=w.get("trend", ""),
            rationale=w.get("rationale", ""),
            confidence=w.get("confidence", "Medium"),
        )
        for w in data.get("winners", [])
    ]
    losers = [
        CompanyTrendPosition(
            symbol=l.get("symbol", ""),
            company_name=l.get("company_name", ""),
            position="loser",
            trend=l.get("trend", ""),
            rationale=l.get("rationale", ""),
            confidence=l.get("confidence", "Medium"),
        )
        for l in data.get("losers", [])
    ]
    logger.info(
        f"Synthesis refined: {len(initial.winners)}W/{len(initial.losers)}L -> "
        f"{len(winners)}W/{len(losers)}L"
    )
    return WinnersLosersAnalysis(winners=winners, losers=losers, summary=data.get("summary", initial.summary))


@dataclass
class RobustReportResult:
    """Bundled output of run_robust_report."""
    theme: str
    sector: Optional[str] = None
    findings: Any = None  # WebResearchFindings
    deep_dives: Dict[str, TickerDeepDive] = field(default_factory=dict)
    synthesis: Optional[WinnersLosersAnalysis] = None
    cost: Optional[CostTracker] = None


def run_robust_report(
    theme: str,
    tickers: List[str],
    ticker_names: Optional[Dict[str, str]] = None,
    sector: Optional[str] = None,
    user_directions: str = "",
    cost_cap_usd: float = 12.0,
    lookback_days: int = 120,
    web_search_max: int = 30,
    do_critique: bool = True,
    progress_callback=None,
) -> RobustReportResult:
    """Run the robust report pipeline end-to-end with a hard cost cap.

    Stages (each gated by the budget):
      1. Industry-level web research (~$0.50)
      2. Per-ticker deep dives, one per ticker (~$0.50 each)
      3. 11-agent synthesis on the combined research (~$3-4)
      4. Critique-and-refine pass (~$1)
      5. Track-record recording

    If the budget cap would be exceeded before a stage, that stage is
    skipped (logged as warning) and the pipeline returns what it has so far.
    """
    tracker = CostTracker(cap_usd=cost_cap_usd)
    ticker_names = ticker_names or {}
    result = RobustReportResult(theme=theme, sector=sector, cost=tracker)

    def report(msg: str, pct: float):
        logger.info(f"[robust {pct:.0%}] {msg}")
        if progress_callback:
            progress_callback(msg, pct)

    # Realistic per-stage cost estimates (used by require_remaining as a guard;
    # the tracker records actual usage afterward). Tuned against measured Sonnet
    # 4.6 spend on Sonnet 4.6 + web_search:
    #   industry research (30 searches): ~$0.50
    #   each deep dive (8 searches):     ~$0.25
    #   agents + synthesis (11 + 1):     ~$1.20
    #   critique + refine (2 calls):     ~$0.50
    EST_INDUSTRY = 0.60
    EST_DEEP_DIVE = 0.30
    EST_AGENTS_SYN = 2.50
    EST_CRITIQUE = 1.00

    # ---- Stage 1: industry-level web research ----
    try:
        tracker.require_remaining(estimated_usd=EST_INDUSTRY, stage_label="industry_research")
    except BudgetExceeded as e:
        logger.warning(f"Skipping industry research: {e}")
        return result
    report("Industry web research...", 0.05)
    findings = web_research_industry(
        theme=theme,
        tickers=tickers,
        lookback_days=lookback_days,
        max_searches=web_search_max,
        tracker=tracker,
    )
    result.findings = findings

    # ---- Stage 2: per-ticker deep dives ----
    for i, sym in enumerate(tickers, 1):
        try:
            tracker.require_remaining(estimated_usd=EST_DEEP_DIVE, stage_label=f"deep_dive_{sym}")
        except BudgetExceeded as e:
            logger.warning(f"Skipping deep dive for {sym} and beyond: {e}")
            break
        report(f"Deep dive {sym} ({i}/{len(tickers)})...", 0.05 + 0.35 * i / max(1, len(tickers)))
        dive = deep_research_ticker(
            ticker=sym,
            company_name=ticker_names.get(sym, ""),
            theme=theme,
            lookback_days=lookback_days,
            max_searches=8,
            tracker=tracker,
        )
        result.deep_dives[sym] = dive

    # ---- Stage 3: 11-agent synthesis ----
    pieces = [findings.as_context_block()]
    for sym, dive in result.deep_dives.items():
        pieces.append(dive.as_brief())
    report_content = "\n\n".join(pieces)

    try:
        tracker.require_remaining(estimated_usd=EST_AGENTS_SYN, stage_label="agents_and_synthesis")
    except BudgetExceeded as e:
        logger.warning(f"Skipping agents/synthesis: {e}")
        return result

    report("Running 11-agent synthesis...", 0.50)
    universe = pd.DataFrame(
        {"Symbol": tickers, "Name": [ticker_names.get(t, t) for t in tickers]}
    )
    synthesis = run_all_agents_and_synthesize(
        report_content=report_content,
        universe_df=universe,
        ai_provider="anthropic",
        user_directions=user_directions,
        theme=theme,
        sector=sector,
        record_to_track_record=False,  # we'll record AFTER critique so picks reflect final view
        progress_callback=lambda m, p: report(m, 0.50 + 0.35 * p),
    )

    # ---- Stage 4: critique-and-refine ----
    if do_critique and synthesis and (synthesis.winners or synthesis.losers):
        try:
            tracker.require_remaining(estimated_usd=EST_CRITIQUE, stage_label="critique_refine")
            report("Critique-and-refine pass...", 0.88)
            # Reconstitute agent_results from the synthesis's input by re-running?
            # Cheaper to just pass the synthesis prose + deep-dive briefs as context.
            # We'll synthesize a compact "agent views" block from the briefs themselves.
            agent_views_block = {
                "deep_dive_briefs": "\n\n".join(d.as_brief() for d in result.deep_dives.values()),
                "industry_findings": findings.as_context_block(),
            }
            synthesis = critique_and_refine_synthesis(
                initial=synthesis,
                agent_results=agent_views_block,
                ai_provider="anthropic",
                cost_tracker=tracker,
            )
        except BudgetExceeded as e:
            logger.warning(f"Skipping critique-refine: {e}")

    result.synthesis = synthesis

    # ---- Stage 5: record picks from the FINAL synthesis ----
    if synthesis and (synthesis.winners or synthesis.losers):
        try:
            companies_for_pricing = []
            for sym in tickers:
                try:
                    profile = get_company_profile(sym) or {}
                    companies_for_pricing.append({"symbol": sym, "price": profile.get("price")})
                except Exception:
                    companies_for_pricing.append({"symbol": sym, "price": None})
            picks = picks_from_winners_losers(synthesis, agent="synthesis_robust", companies=companies_for_pricing)
            record_picks(
                theme=theme,
                tickers=tickers,
                picks=picks,
                sector=sector,
                model="claude-sonnet-4-6",
                prompt_version="robust-v1",
            )
        except Exception as rec_err:
            logger.warning(f"Failed to record robust picks: {rec_err}")

    report("Done.", 1.0)
    logger.info(tracker.summary())
    return result


def scan_universe_for_winners_losers(
    universe_df: pd.DataFrame,
    trends_data: Dict[str, Any],
    ai_provider: str = "anthropic"
) -> WinnersLosersAnalysis:
    """Scan stock universe to identify winners and losers based on extracted trends."""

    # Get relevant columns from universe
    symbol_col = None
    name_col = None
    sector_col = None

    for col in universe_df.columns:
        col_lower = str(col).lower()
        if 'symbol' in col_lower or 'ticker' in col_lower or col == 'Symbol':
            symbol_col = col
        elif 'name' in col_lower or 'company' in col_lower or col == 'Company':
            name_col = col
        elif 'sector' in col_lower:
            sector_col = col

    # Fallback: use first column as symbol if not found
    if not symbol_col and len(universe_df.columns) > 0:
        symbol_col = universe_df.columns[0]
        logger.info(f"Using first column '{symbol_col}' as symbol column")

    # Use second column as name if not found
    if not name_col and len(universe_df.columns) > 1:
        name_col = universe_df.columns[1]

    if not symbol_col:
        return WinnersLosersAnalysis(summary="Could not find symbol column in universe file")

    # Build stock list for analysis
    stocks_info = []
    for _, row in universe_df.head(100).iterrows():  # Limit to 100 stocks
        info = f"- {row[symbol_col]}"
        if name_col and pd.notna(row.get(name_col)):
            info += f" ({row[name_col]})"
        if sector_col and pd.notna(row.get(sector_col)):
            info += f" [Sector: {row[sector_col]}]"
        stocks_info.append(info)

    prompt = f"""Based on the following industry trends and themes, analyze the stock universe and identify WINNERS and LOSERS.

INDUSTRY TRENDS & THEMES:
- Industry: {trends_data.get('industry', 'N/A')}
- Key Trends: {', '.join(trends_data.get('key_trends', []))}
- Bullish Factors: {', '.join(trends_data.get('bullish_factors', []))}
- Bearish Factors: {', '.join(trends_data.get('bearish_factors', []))}
- Key Themes: {', '.join(trends_data.get('key_themes', []))}

STOCK UNIVERSE TO ANALYZE:
{chr(10).join(stocks_info)}

Analyze each stock against the trends and categorize as WINNER, LOSER, or skip if not relevant.

Return as JSON:
{{
    "summary": "Brief 2-3 sentence summary of winners/losers dynamics",
    "winners": [
        {{
            "symbol": "TICKER",
            "company_name": "Company Name",
            "trend": "The specific trend they benefit from",
            "rationale": "Why they win (2-3 sentences)",
            "confidence": "High/Medium/Low"
        }}
    ],
    "losers": [
        {{
            "symbol": "TICKER",
            "company_name": "Company Name",
            "trend": "The specific trend hurting them",
            "rationale": "Why they lose (2-3 sentences)",
            "confidence": "High/Medium/Low"
        }}
    ]
}}

Focus on stocks that are CLEARLY positioned as winners or losers. Skip neutral stocks.
Return ONLY valid JSON."""

    try:
        if ai_provider == "anthropic" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.content[0].text
        elif ai_provider == "openai" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.choices[0].message.content
        else:
            return WinnersLosersAnalysis(summary="No AI provider available")

        # Clean up response
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        data = json.loads(response_text.strip())

        winners = [
            CompanyTrendPosition(
                symbol=w.get('symbol', ''),
                company_name=w.get('company_name', ''),
                position='winner',
                trend=w.get('trend', ''),
                rationale=w.get('rationale', ''),
                confidence=w.get('confidence', 'Medium')
            )
            for w in data.get('winners', [])
        ]

        losers = [
            CompanyTrendPosition(
                symbol=l.get('symbol', ''),
                company_name=l.get('company_name', ''),
                position='loser',
                trend=l.get('trend', ''),
                rationale=l.get('rationale', ''),
                confidence=l.get('confidence', 'Medium')
            )
            for l in data.get('losers', [])
        ]

        return WinnersLosersAnalysis(
            winners=winners,
            losers=losers,
            summary=data.get('summary', '')
        )

    except Exception as e:
        logger.error(f"Error scanning universe: {e}")
        return WinnersLosersAnalysis(summary=f"Error: {str(e)}")

# Page configuration
st.set_page_config(
    page_title="Industry Report Generator",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main .block-container {
        border: 1.5px solid black;
        padding: 30px;
        border-radius: 10px;
        background-color: white;
        max-width: 1400px;
        margin: auto;
    }
    .stMetric {
        background-color: #f0f2f6;
        padding: 10px;
        border-radius: 5px;
    }
    .winner-card {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
    }
    .loser-card {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
    }
    .neutral-card {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
    }
    .section-header {
        background-color: #2E86AB;
        color: white;
        padding: 10px 15px;
        border-radius: 5px;
        margin: 20px 0 10px 0;
    }
    div[data-testid="stExpander"] {
        border: 1px solid #ddd;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

# Available sectors (from FMP)
SECTORS = [
    "Technology",
    "Healthcare",
    "Financial Services",
    "Consumer Cyclical",
    "Communication Services",
    "Industrials",
    "Consumer Defensive",
    "Energy",
    "Basic Materials",
    "Real Estate",
    "Utilities"
]

# Common industries by sector
INDUSTRIES_BY_SECTOR = {
    "Technology": [
        "Software - Application",
        "Software - Infrastructure",
        "Semiconductors",
        "Information Technology Services",
        "Computer Hardware",
        "Electronic Components",
        "Scientific & Technical Instruments",
        "Communication Equipment"
    ],
    "Healthcare": [
        "Biotechnology",
        "Drug Manufacturers - General",
        "Medical Devices",
        "Healthcare Plans",
        "Medical Instruments & Supplies",
        "Diagnostics & Research",
        "Medical Care Facilities"
    ],
    "Financial Services": [
        "Banks - Diversified",
        "Banks - Regional",
        "Asset Management",
        "Insurance - Diversified",
        "Capital Markets",
        "Credit Services",
        "Financial Data & Stock Exchanges"
    ],
    "Consumer Cyclical": [
        "Internet Retail",
        "Auto Manufacturers",
        "Restaurants",
        "Home Improvement Retail",
        "Apparel Retail",
        "Specialty Retail",
        "Travel Services"
    ],
    "Communication Services": [
        "Internet Content & Information",
        "Entertainment",
        "Telecom Services",
        "Electronic Gaming & Multimedia",
        "Advertising Agencies"
    ],
    "Industrials": [
        "Aerospace & Defense",
        "Railroads",
        "Industrial Distribution",
        "Farm & Heavy Construction Machinery",
        "Specialty Industrial Machinery",
        "Consulting Services"
    ],
    "Energy": [
        "Oil & Gas Integrated",
        "Oil & Gas E&P",
        "Oil & Gas Midstream",
        "Oil & Gas Refining & Marketing",
        "Uranium"
    ],
    "Consumer Defensive": [
        "Household & Personal Products",
        "Beverages - Non-Alcoholic",
        "Packaged Foods",
        "Discount Stores",
        "Grocery Stores"
    ],
    "Basic Materials": [
        "Specialty Chemicals",
        "Gold",
        "Copper",
        "Steel",
        "Lumber & Wood Production"
    ],
    "Real Estate": [
        "REIT - Residential",
        "REIT - Industrial",
        "REIT - Retail",
        "REIT - Office",
        "Real Estate Services"
    ],
    "Utilities": [
        "Utilities - Regulated Electric",
        "Utilities - Renewable",
        "Utilities - Diversified",
        "Utilities - Independent Power Producers"
    ]
}


def display_logo():
    """Display company logo at the top, centered."""
    logo_path = "company_logo.png"
    if os.path.exists(logo_path):
        st.markdown(
            """
            <div style="display: flex; justify-content: center; align-items: center; flex-direction: column; margin-bottom: 10px;">
            """,
            unsafe_allow_html=True
        )
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image(logo_path, use_container_width=True)
        st.markdown(
            """
            <p style='text-align: center; color: #666; font-style: italic; margin-top: 5px;'>
                Precision Analysis for Informed Investment Decisions
            </p>
            </div>
            """,
            unsafe_allow_html=True
        )


def display_winners_losers(winners_losers: WinnersLosersAnalysis):
    """Display winners and losers in styled cards."""
    if not winners_losers:
        st.info("No winners/losers analysis available.")
        return

    # Summary
    if winners_losers.summary:
        st.markdown(f"**Summary:** {winners_losers.summary}")
        st.markdown("---")

    col1, col2 = st.columns(2)

    # Winners
    with col1:
        st.markdown("### 🏆 Winners")
        if winners_losers.winners:
            for w in winners_losers.winners:
                confidence_color = {"High": "🟢", "Medium": "🟡", "Low": "🔴"}.get(w.confidence, "⚪")
                st.markdown(f"""
                <div class="winner-card">
                    <h4 style="color: #155724; margin: 0;">{w.symbol} - {w.company_name}</h4>
                    <p style="margin: 5px 0;"><strong>Trend:</strong> {w.trend}</p>
                    <p style="margin: 5px 0; font-size: 0.9em;">{w.rationale}</p>
                    <p style="margin: 0; font-size: 0.85em;"><strong>Confidence:</strong> {confidence_color} {w.confidence}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No clear winners identified.")

    # Losers
    with col2:
        st.markdown("### 📉 Losers")
        if winners_losers.losers:
            for l in winners_losers.losers:
                confidence_color = {"High": "🟢", "Medium": "🟡", "Low": "🔴"}.get(l.confidence, "⚪")
                st.markdown(f"""
                <div class="loser-card">
                    <h4 style="color: #721c24; margin: 0;">{l.symbol} - {l.company_name}</h4>
                    <p style="margin: 5px 0;"><strong>Trend:</strong> {l.trend}</p>
                    <p style="margin: 5px 0; font-size: 0.9em;">{l.rationale}</p>
                    <p style="margin: 0; font-size: 0.85em;"><strong>Confidence:</strong> {confidence_color} {l.confidence}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No clear losers identified.")

    # Neutral (if any)
    if winners_losers.neutral:
        st.markdown("### ⚖️ Neutral / Mixed Positioning")
        for n in winners_losers.neutral:
            st.markdown(f"""
            <div class="neutral-card">
                <strong>{n.symbol}</strong> ({n.company_name}): {n.rationale}
            </div>
            """, unsafe_allow_html=True)


def main():
    # Display logo
    display_logo()

    # Title
    st.markdown("<h1 style='text-align: center;'>🏭 Industry Report Generator</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #666;'>Generate comprehensive industry analysis with AI-powered insights and trend-based winners/losers identification</p>", unsafe_allow_html=True)

    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Report Configuration")

        # Debug: Show logo path info
        with st.expander("🔧 Debug Info"):
            st.write(f"**Logo Path:** {LOGO_PATH}")
            st.write(f"**Script Dir:** {SCRIPT_DIR}")
            st.write(f"**CWD:** {os.getcwd()}")
            st.write(f"**PNG files found:** {[str(f) for f in DEBUG_FILES]}")

        # Mode selection - Upload existing report OR generate new
        report_mode = st.radio(
            "Mode:",
            ["📄 Analyze Existing Report", "🔍 Generate New Report"],
            help="Upload a completed report to scan for winners/losers, or generate a new industry report"
        )

        st.markdown("---")

        # AI Provider (common to both modes)
        ai_provider = st.selectbox(
            "AI Provider",
            ["anthropic", "openai"],
            index=0,
            help="Select AI provider for analysis"
        )

        # Variables for new report mode
        selected_sector = None
        selected_industry = None
        company_limit = 20
        analyst_notes_text = ""
        key_themes_text = ""
        investment_thesis = ""
        articles_data = []

        if report_mode == "📄 Analyze Existing Report":
            # ============================================
            # MODE 1: UPLOAD EXISTING REPORT
            # ============================================
            st.markdown("---")
            st.subheader("📄 Upload Industry Report")
            st.markdown("*Upload your completed report to identify winners & losers*")

            # Custom industry/theme name
            custom_industry_name = st.text_input(
                "Industry/Theme Name",
                placeholder="e.g., Tokenization, AI Infrastructure, Clean Energy...",
                help="Enter any industry or investment theme - not limited to standard sectors"
            )

            # Upload report
            uploaded_industry_note = st.file_uploader(
                "Upload Report",
                type=["docx", "pdf", "txt"],
                help="Upload Word, PDF, or text file",
                key="industry_note_upload"
            )

            st.markdown("---")

            # Additional research files
            st.subheader("📎 Additional Research Files")
            st.markdown("*Upload related reports, data files, or articles*")

            additional_research_files = st.file_uploader(
                "Upload additional research documents",
                type=["pdf", "docx", "xlsx", "xls", "txt"],
                accept_multiple_files=True,
                key="additional_research_upload",
                help="Upload PDFs, Word docs, Excel files, or text files with related research"
            )

            if additional_research_files:
                st.success(f"📁 {len(additional_research_files)} file(s) uploaded")
                with st.expander("View uploaded files"):
                    for f in additional_research_files:
                        file_type = get_file_type(f.name)
                        icon = {"pdf": "📄", "word": "📝", "excel": "📊", "text": "📃"}.get(file_type, "📁")
                        st.write(f"{icon} {f.name}")

            st.markdown("---")

            # User Directions / Analysis Focus
            st.subheader("🎯 Analysis Directions")
            st.markdown("*Add specific areas or questions you want analyzed*")

            user_directions = st.text_area(
                "Your Directions & Focus Areas",
                height=150,
                placeholder="""Examples:
• Focus on companies with AI exposure
• Look for margin expansion opportunities
• Identify stocks vulnerable to tariffs
• Find undervalued small-caps in this space
• Which companies have pricing power?
• Analyze impact of new regulations...""",
                key="user_directions_scan",
                help="These directions will guide the AI agents and be included in the final report"
            )

            st.markdown("---")

            # Scan button
            scan_universe_button = st.button(
                "🔍 Scan for Winners & Losers",
                type="primary",
                use_container_width=True,
                disabled=(uploaded_industry_note is None or not custom_industry_name)
            )

            if uploaded_industry_note is None:
                st.caption("⬆️ Upload a report to enable scanning")
            elif not custom_industry_name:
                st.caption("⬆️ Enter an industry/theme name")

            # Set generate_button to False for this mode
            generate_button = False
            uploaded_notes = None

        else:
            # ============================================
            # MODE 2: GENERATE NEW REPORT
            # ============================================
            st.markdown("---")
            st.subheader("🔍 Industry Selection")

            # Selection mode
            selection_mode = st.radio(
                "Select by:",
                ["Sector", "Industry", "Custom Theme", "Market View"],
                help="Choose standard sector/industry, custom theme, or Market View (research-based only)"
            )

            if selection_mode == "Sector":
                selected_sector = st.selectbox(
                    "Select Sector",
                    SECTORS,
                    index=0
                )
            elif selection_mode == "Industry":
                sector_for_industry = st.selectbox(
                    "Filter by Sector",
                    SECTORS,
                    index=0
                )
                industries = INDUSTRIES_BY_SECTOR.get(sector_for_industry, [])
                if industries:
                    selected_industry = st.selectbox(
                        "Select Industry",
                        industries
                    )
                else:
                    selected_industry = st.text_input("Custom Industry Name")
            elif selection_mode == "Custom Theme":
                # Custom theme
                selected_industry = st.text_input(
                    "Custom Theme/Industry",
                    placeholder="e.g., Tokenization, Space Economy, Nuclear Energy..."
                )
            else:
                # Market View mode - research files only, no FMP data
                st.info("📊 **Market View Mode**: Generate a report based solely on your research files and notes. No company data will be fetched from market APIs.")
                selected_industry = st.text_input(
                    "Report Title",
                    value="Market View",
                    placeholder="e.g., Q1 2026 Market View, Weekly Market Commentary..."
                )

            st.markdown("---")

            # Analysis options
            st.subheader("📊 Analysis Options")
            company_limit = st.slider(
                "Max Companies",
                min_value=5,
                max_value=50,
                value=20
            )

            st.markdown("---")

            # Research notes section
            st.subheader("📝 Research Notes (Optional)")

            # Multi-file upload for research documents
            st.markdown("**📎 Upload Research Files**")
            uploaded_research_files = st.file_uploader(
                "Upload research documents",
                type=["pdf", "docx", "xlsx", "xls", "txt"],
                accept_multiple_files=True,
                key="research_files_upload",
                help="Upload PDFs, Word docs, Excel files, or text files to include in analysis"
            )

            if uploaded_research_files:
                st.success(f"📁 {len(uploaded_research_files)} file(s) uploaded")
                with st.expander("View uploaded files"):
                    for f in uploaded_research_files:
                        file_type = get_file_type(f.name)
                        icon = {"pdf": "📄", "word": "📝", "excel": "📊", "text": "📃"}.get(file_type, "📁")
                        st.write(f"{icon} {f.name}")

            with st.expander("Add Analyst Notes"):
                analyst_notes_text = st.text_area(
                    "Analyst Notes (one per line)",
                    height=100,
                    placeholder="Enter analyst observations..."
                )

            with st.expander("Add Key Themes"):
                key_themes_text = st.text_area(
                    "Key Investment Themes (one per line)",
                    height=80,
                    placeholder="AI infrastructure buildout..."
                )

            with st.expander("Add Investment Thesis"):
                investment_thesis = st.text_area(
                    "Investment Thesis",
                    height=100,
                    placeholder="Overall investment thesis..."
                )

            with st.expander("Add Articles"):
                st.markdown("**Add research articles:**")
                num_articles = st.number_input("Number of articles", min_value=0, max_value=5, value=0)
                articles_data = []
                for i in range(int(num_articles)):
                    st.markdown(f"**Article {i+1}**")
                    title = st.text_input(f"Title {i+1}", key=f"art_title_{i}")
                    source = st.text_input(f"Source {i+1}", key=f"art_source_{i}")
                    date = st.text_input(f"Date {i+1}", key=f"art_date_{i}", placeholder="YYYY-MM-DD")
                    content = st.text_area(f"Summary {i+1}", key=f"art_content_{i}", height=60)
                    url = st.text_input(f"URL {i+1}", key=f"art_url_{i}")
                    if title:
                        articles_data.append({
                            "title": title,
                            "source": source,
                            "date": date,
                            "content": content,
                            "url": url
                        })

            # Upload JSON notes file
            st.markdown("---")
            uploaded_notes = st.file_uploader(
                "Or upload notes JSON file",
                type=["json"],
                help="Upload a JSON file with analyst notes and articles"
            )

            st.markdown("---")

            # User Directions / Analysis Focus
            st.subheader("🎯 Analysis Directions")
            st.markdown("*Add specific areas or questions you want analyzed*")

            user_directions_generate = st.text_area(
                "Your Directions & Focus Areas",
                height=150,
                placeholder="""Examples:
• Focus on companies with AI exposure
• Look for margin expansion opportunities
• Identify stocks vulnerable to tariffs
• Find undervalued small-caps in this space
• Which companies have pricing power?
• Analyze impact of new regulations...""",
                key="user_directions_generate",
                help="These directions will guide the AI analysis and be included in the final report"
            )

            st.markdown("---")

            # Generate button for new report mode
            generate_button = st.button(
                "🚀 Generate Report",
                type="primary",
                use_container_width=True
            )

            # Initialize variables not used in this mode
            scan_universe_button = False
            uploaded_industry_note = None
            custom_industry_name = None

    # Main content area

    # Handle Scan Universe button (Analyze Existing Report mode)
    if scan_universe_button and uploaded_industry_note:
        try:
            # Read the uploaded industry note
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_text.text("Reading industry report...")

            # Store original file bytes for Word export (preserve formatting)
            uploaded_industry_note.seek(0)
            original_file_bytes = uploaded_industry_note.read()
            uploaded_industry_note.seek(0)
            original_filename = uploaded_industry_note.name

            note_content = read_uploaded_document(uploaded_industry_note)
            progress_bar.progress(5)

            if not note_content:
                st.error("Could not read the uploaded document. Please try a different format.")
            else:
                st.success(f"Read {len(note_content):,} characters from report")

                # Process additional research files if uploaded
                research_files_list = []
                combined_research_content = ""

                if additional_research_files:
                    status_text.text(f"Processing {len(additional_research_files)} additional research file(s)...")
                    st.info(f"📎 Processing {len(additional_research_files)} additional research file(s)...")

                    def research_progress(msg, pct):
                        status_text.text(msg)
                        progress_bar.progress(int(5 + pct * 10))  # 5-15% for research files

                    research_files_list = process_research_files(
                        additional_research_files,
                        industry_context=custom_industry_name,
                        ai_provider=ai_provider,
                        progress_callback=research_progress
                    )

                    # Combine content from all research files for analysis context
                    for rf in research_files_list:
                        if rf.content:
                            combined_research_content += f"\n\n=== {rf.filename} ===\n{rf.content[:5000]}"

                    st.success(f"✅ Processed {len(research_files_list)} research file(s)")

                progress_bar.progress(15)

                # Combine main note with additional research content for analysis
                full_analysis_content = note_content
                if combined_research_content:
                    full_analysis_content += f"\n\n=== ADDITIONAL RESEARCH DOCUMENTS ===\n{combined_research_content}"

                # Run multi-agent analysis (no stock universe needed)
                def update_progress(msg, pct):
                    status_text.text(msg)
                    progress_bar.progress(int(15 + pct * 85))

                status_text.text("Running multi-agent analysis...")
                st.info("🤖 **11 AI Agents analyzing your report:** Industry Analyst → Competitive Intel → Financial Analyst → Valuation Analyst → Technical/Momentum → Supply Chain → ESG Analyst → Macro/Policy → M&A Strategist → Risk Analyst → Investment Strategist")

                winners_losers_result = run_all_agents_and_synthesize(
                    full_analysis_content,
                    None,  # No universe - analyze report directly
                    ai_provider,
                    progress_callback=update_progress,
                    user_directions=user_directions
                )
                progress_bar.progress(100)

                # Extract trends for display (quick extraction)
                trends_data = extract_trends_from_note(note_content, ai_provider)

                # Store in session state (including original file for Word export)
                st.session_state['universe_scan_results'] = {
                    'industry_name': custom_industry_name,
                    'note_content': note_content,
                    'trends_data': trends_data if "error" not in trends_data else {},
                    'winners_losers': winners_losers_result,
                    'original_file_bytes': original_file_bytes,
                    'original_filename': original_filename,
                    'research_files': research_files_list,  # Store processed research files
                    'user_directions': user_directions  # Store user analysis directions
                }

                status_text.text("")
                st.success(f"✅ Multi-agent analysis complete! Found {len(winners_losers_result.winners)} winners and {len(winners_losers_result.losers)} losers.")

        except Exception as e:
            st.error(f"Error during analysis: {str(e)}")
            logger.exception("Universe scan failed")

    # Display universe scan results if available
    if 'universe_scan_results' in st.session_state:
        scan_results = st.session_state['universe_scan_results']
        trends_data = scan_results.get('trends_data', {})
        industry_name = scan_results.get('industry_name', 'Industry')

        st.markdown("---")
        st.markdown(f"<h2 style='text-align: center;'>📊 {industry_name} Analysis</h2>", unsafe_allow_html=True)

        # Display user directions if provided
        user_dirs_display = scan_results.get('user_directions', '')
        if user_dirs_display and user_dirs_display.strip():
            st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>🎯 Analysis Directions</h3></div>", unsafe_allow_html=True)
            st.info(user_dirs_display)

        # Display extracted trends
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>📈 Extracted Trends & Themes</h3></div>", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**Industry:** {trends_data.get('industry', industry_name)}")
            st.markdown("**Key Trends:**")
            for trend in trends_data.get('key_trends', []):
                st.markdown(f"• {trend}")

            if trends_data.get('key_themes'):
                st.markdown("**Key Themes:**")
                for theme in trends_data.get('key_themes', []):
                    st.markdown(f"• {theme}")

        with col2:
            st.markdown("**Bullish Factors:**")
            for factor in trends_data.get('bullish_factors', []):
                st.markdown(f"🟢 {factor}")

            st.markdown("**Bearish Factors:**")
            for factor in trends_data.get('bearish_factors', []):
                st.markdown(f"🔴 {factor}")

        if trends_data.get('summary'):
            st.info(f"**Summary:** {trends_data['summary']}")

        # Research Files Section (if any)
        research_files = scan_results.get('research_files', [])
        if research_files:
            st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>📎 Research Documents</h3></div>", unsafe_allow_html=True)
            st.markdown(f"*{len(research_files)} research document(s) analyzed*")

            for rf in research_files:
                file_icon = {"pdf": "📄", "word": "📝", "excel": "📊", "text": "📃"}.get(rf.file_type, "📁")
                with st.expander(f"{file_icon} {rf.filename}"):
                    st.markdown(rf.summary)

        # Winners & Losers section
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>🎯 Winners & Losers</h3></div>", unsafe_allow_html=True)
        display_winners_losers(scan_results['winners_losers'])

        # Export options
        st.markdown("---")
        st.markdown("### 📥 Export Results")

        col1, col2, col3 = st.columns(3)

        # Export as text
        with col1:
            wl = scan_results['winners_losers']
            user_dirs = scan_results.get('user_directions', '')
            export_text = f"# {industry_name} - Winners & Losers Analysis\n"
            export_text += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
            if user_dirs:
                export_text += f"## Analysis Directions\n{user_dirs}\n\n"
            export_text += f"## Summary\n{wl.summary}\n\n"
            export_text += "## WINNERS\n"
            for w in wl.winners:
                export_text += f"- {w.symbol} ({w.company_name})\n"
                export_text += f"  Trend: {w.trend}\n"
                export_text += f"  Rationale: {w.rationale}\n"
                export_text += f"  Confidence: {w.confidence}\n\n"
            export_text += "## LOSERS\n"
            for l in wl.losers:
                export_text += f"- {l.symbol} ({l.company_name})\n"
                export_text += f"  Trend: {l.trend}\n"
                export_text += f"  Rationale: {l.rationale}\n"
                export_text += f"  Confidence: {l.confidence}\n\n"

            st.download_button(
                "📄 Download as Text",
                data=export_text,
                file_name=f"{industry_name.replace(' ', '_')}_winners_losers.txt",
                mime="text/plain"
            )

        # Export as JSON
        with col2:
            export_json = {
                "industry": industry_name,
                "generated": datetime.now().isoformat(),
                "user_directions": scan_results.get('user_directions', ''),
                "trends": trends_data,
                "summary": wl.summary,
                "winners": [{"symbol": w.symbol, "company": w.company_name, "trend": w.trend, "rationale": w.rationale, "confidence": w.confidence} for w in wl.winners],
                "losers": [{"symbol": l.symbol, "company": l.company_name, "trend": l.trend, "rationale": l.rationale, "confidence": l.confidence} for l in wl.losers]
            }
            st.download_button(
                "📋 Download as JSON",
                data=json.dumps(export_json, indent=2),
                file_name=f"{industry_name.replace(' ', '_')}_winners_losers.json",
                mime="application/json"
            )

        # Export as CSV
        with col3:
            rows = []
            for w in wl.winners:
                rows.append({"Type": "Winner", "Symbol": w.symbol, "Company": w.company_name, "Trend": w.trend, "Rationale": w.rationale, "Confidence": w.confidence})
            for l in wl.losers:
                rows.append({"Type": "Loser", "Symbol": l.symbol, "Company": l.company_name, "Trend": l.trend, "Rationale": l.rationale, "Confidence": l.confidence})
            export_df = pd.DataFrame(rows)
            st.download_button(
                "📊 Download as CSV",
                data=export_df.to_csv(index=False),
                file_name=f"{industry_name.replace(' ', '_')}_winners_losers.csv",
                mime="text/csv"
            )

        # PDF and Email row
        st.markdown("---")
        st.markdown("### 📑 Professional Report")

        # Check original file type
        original_filename = scan_results.get('original_filename', '')
        is_word_doc = original_filename.lower().endswith('.docx')
        is_pdf_doc = original_filename.lower().endswith('.pdf')

        if is_word_doc:
            st.info("📝 **Original document formatting preserved** - Download as Word to keep all bullet points, indentation, and formatting from your uploaded document.")
        elif is_pdf_doc:
            st.info("📄 **Original PDF preserved** - Download as PDF to keep your original document with Winners & Losers appended.")

        col1, col2, col3 = st.columns(3)

        with col1:
            # Generate Word doc with original formatting preserved or with note content
            research_files_for_export = scan_results.get('research_files', [])
            user_dirs_for_export = scan_results.get('user_directions', '')
            if is_word_doc and scan_results.get('original_file_bytes'):
                original_buffer = BytesIO(scan_results['original_file_bytes'])
                word_buffer = generate_winners_losers_word(
                    industry_name,
                    wl,
                    original_file_buffer=original_buffer,
                    research_files=research_files_for_export,
                    user_directions=user_dirs_for_export
                )
            else:
                # Generate Word doc with note content included
                word_buffer = generate_winners_losers_word(
                    industry_name,
                    wl,
                    original_file_buffer=None,
                    note_content=scan_results.get('note_content'),
                    research_files=research_files_for_export,
                    user_directions=user_dirs_for_export
                )
            st.download_button(
                "📥 Download Word Report",
                data=word_buffer,
                file_name=f"{industry_name.replace(' ', '_')}_Winners_Losers_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

        with col2:
            # PDF option - merge with original if PDF was uploaded
            if is_pdf_doc and scan_results.get('original_file_bytes'):
                # Merge original PDF with Winners/Losers appendix
                appendix_buffer = generate_winners_losers_appendix_pdf(industry_name, wl)
                pdf_buffer = merge_pdf_with_appendix(
                    scan_results['original_file_bytes'],
                    appendix_buffer,
                    industry_name
                )
                st.download_button(
                    "📄 Download PDF (Original + W&L)",
                    data=pdf_buffer,
                    file_name=f"{industry_name.replace(' ', '_')}_Full_Report_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                    type="primary"
                )
            else:
                # Generate standalone PDF with research note
                pdf_buffer = generate_winners_losers_pdf(
                    industry_name,
                    trends_data,
                    wl,
                    scan_results.get('note_content'),
                    scan_results.get('user_directions', '')
                )
                st.download_button(
                    "📄 Download PDF Report",
                    data=pdf_buffer,
                    file_name=f"{industry_name.replace(' ', '_')}_Report_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf"
                )

        with col3:
            col3a, col3b = st.columns(2)
            with col3a:
                if st.button("📧 Email Word", type="primary", key="email_word"):
                    with st.spinner("Sending Word..."):
                        # Always send Word document with research note included
                        research_files_for_email = scan_results.get('research_files', [])
                        user_dirs_for_email = scan_results.get('user_directions', '')
                        if is_word_doc and scan_results.get('original_file_bytes'):
                            original_buffer = BytesIO(scan_results['original_file_bytes'])
                            word_buffer = generate_winners_losers_word(
                                industry_name,
                                wl,
                                original_file_buffer=original_buffer,
                                research_files=research_files_for_email,
                                user_directions=user_dirs_for_email
                            )
                        else:
                            word_buffer = generate_winners_losers_word(
                                industry_name,
                                wl,
                                original_file_buffer=None,
                                note_content=scan_results.get('note_content'),
                                research_files=research_files_for_email,
                                user_directions=user_dirs_for_email
                            )
                        success, message = send_email_with_attachment(word_buffer, industry_name, "docx")
                        if success:
                            st.success(message)
                        else:
                            st.error(message)
            with col3b:
                if st.button("📧 Email PDF", key="email_pdf"):
                    with st.spinner("Sending PDF..."):
                        # Merge with original if PDF was uploaded
                        if is_pdf_doc and scan_results.get('original_file_bytes'):
                            appendix_buffer = generate_winners_losers_appendix_pdf(industry_name, wl)
                            pdf_for_email = merge_pdf_with_appendix(
                                scan_results['original_file_bytes'],
                                appendix_buffer,
                                industry_name
                            )
                        else:
                            pdf_for_email = generate_winners_losers_pdf(
                                industry_name,
                                trends_data,
                                wl,
                                scan_results.get('note_content'),
                                scan_results.get('user_directions', '')
                            )
                        success, message = send_email_with_attachment(pdf_for_email, industry_name, "pdf")
                        if success:
                            st.success(message)
                        else:
                            st.error(message)

        # Clear results button
        st.markdown("---")
        if st.button("🗑️ Clear Results", key="clear_scan_results"):
            del st.session_state['universe_scan_results']
            st.rerun()

    if generate_button:
        target = selected_industry if selected_industry else selected_sector

        if not target:
            st.error("Please select a sector or industry.")
            return

        # Build research notes
        research_notes = None
        if uploaded_notes:
            try:
                notes_data = json.load(uploaded_notes)
                articles = []
                for a in notes_data.get('articles', []):
                    articles.append(Article(
                        title=a.get('title', ''),
                        source=a.get('source', ''),
                        date=a.get('date', ''),
                        content=a.get('content', ''),
                        url=a.get('url', '')
                    ))
                research_notes = ResearchNotes(
                    analyst_notes=notes_data.get('analyst_notes', []),
                    key_themes=notes_data.get('key_themes', []),
                    investment_thesis=notes_data.get('investment_thesis', ''),
                    articles=articles
                )
                st.success("Loaded research notes from uploaded file.")
            except Exception as e:
                st.error(f"Error loading notes file: {e}")
        else:
            # Build from form inputs
            analyst_notes = [n.strip() for n in analyst_notes_text.split('\n') if n.strip()] if analyst_notes_text else []
            key_themes = [t.strip() for t in key_themes_text.split('\n') if t.strip()] if key_themes_text else []
            articles = [Article(**a) for a in articles_data] if articles_data else []

            if analyst_notes or key_themes or investment_thesis or articles:
                research_notes = ResearchNotes(
                    analyst_notes=analyst_notes,
                    key_themes=key_themes,
                    investment_thesis=investment_thesis,
                    articles=articles
                )

        # Process uploaded research files
        processed_research_files = []
        if uploaded_research_files:
            with st.spinner(f"Processing {len(uploaded_research_files)} research file(s)..."):
                st.info(f"📎 Processing {len(uploaded_research_files)} research file(s)...")

                def research_progress(msg, pct):
                    st.text(msg)

                processed_research_files = process_research_files(
                    uploaded_research_files,
                    industry_context=target,
                    ai_provider=ai_provider,
                    progress_callback=research_progress
                )
                st.success(f"✅ Processed {len(processed_research_files)} research file(s)")

            # Add research files to research_notes
            if processed_research_files:
                if research_notes is None:
                    research_notes = ResearchNotes()
                research_notes.research_files = processed_research_files

        with st.spinner(f"Generating report for {target}..."):
            try:
                progress = st.progress(0)

                # Check if Market View mode
                is_market_view = (selection_mode == "Market View")

                if is_market_view:
                    # Market View mode - skip company fetching
                    if not research_notes or (not research_notes.research_files and not research_notes.analyst_notes and not research_notes.articles):
                        st.error("Market View mode requires research files or notes. Please upload documents or add notes.")
                        return

                    st.text("Generating AI analysis from research content...")
                    companies = []
                    sector_data = {'sector': 'Market View'}
                    winners_losers = None

                    # Generate market view analysis from research content
                    ai_analysis = generate_market_view_analysis(
                        target,
                        research_notes,
                        ai_provider=ai_provider
                    )
                    progress.progress(70)

                else:
                    # Standard mode - fetch companies from FMP
                    st.text("Fetching companies...")

                    if selected_industry:
                        companies = get_stocks_by_industry(selected_industry, limit=company_limit)
                    else:
                        companies = get_stocks_by_sector(selected_sector, limit=company_limit)

                    if not companies:
                        st.error(f"No companies found for {target}")
                        return

                    # Sort by market cap
                    companies = sorted(companies, key=lambda x: x.get('marketCap', 0), reverse=True)
                    progress.progress(20)

                    # Get sector data
                    st.text("Fetching sector metrics...")
                    sector_pe_data = get_sector_pe_ratios()
                    sector_data = {}
                    if sector_pe_data:
                        for item in sector_pe_data:
                            if selected_sector and item.get('sector', '').lower() == selected_sector.lower():
                                sector_data = item
                                break
                        if not sector_data and companies:
                            sector_data = {'sector': companies[0].get('sector', 'N/A')}
                    progress.progress(40)

                    # Generate AI analysis
                    st.text("Generating AI analysis...")
                    ai_analysis = generate_industry_analysis(
                        target,
                        companies,
                        sector_data,
                        ai_provider=ai_provider
                    )
                    progress.progress(60)

                    # Identify winners and losers
                    st.text("Identifying winners and losers from trends...")
                    winners_losers = None
                    if ai_analysis.get('trends'):
                        winners_losers = identify_winners_losers(
                            target,
                            companies,
                            ai_analysis['trends'],
                            ai_provider=ai_provider
                        )

                progress.progress(80)

                # Generate PDF
                st.text("Generating PDF report...")
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_name = target.replace(" ", "_").replace("/", "_")[:30]
                output_path = f"output/{safe_name}_Report_{timestamp}.pdf"

                pdf_path = generate_industry_pdf(
                    target,
                    companies,
                    sector_data,
                    ai_analysis,
                    output_path=output_path,
                    research_notes=research_notes,
                    winners_losers=winners_losers,
                    market_view_mode=is_market_view
                )
                progress.progress(100)

                # Store in session state
                st.session_state['report_data'] = {
                    'target': target,
                    'companies': companies,
                    'sector_data': sector_data,
                    'ai_analysis': ai_analysis,
                    'winners_losers': winners_losers,
                    'pdf_path': pdf_path
                }

                st.success(f"Report generated successfully!")

            except Exception as e:
                st.error(f"Error generating report: {str(e)}")
                logger.exception("Report generation failed")
                return

    # Display results if available
    if 'report_data' in st.session_state:
        data = st.session_state['report_data']
        target = data['target']
        companies = data['companies']
        sector_data = data['sector_data']
        ai_analysis = data['ai_analysis']
        winners_losers = data['winners_losers']
        pdf_path = data['pdf_path']

        # Download button
        if os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                st.download_button(
                    label="📥 Download PDF Report",
                    data=f.read(),
                    file_name=os.path.basename(pdf_path),
                    mime="application/pdf",
                    type="primary"
                )

        st.markdown("---")

        # Report header
        st.markdown(f"<h2 style='text-align: center;'>Industry Report: {target}</h2>", unsafe_allow_html=True)
        st.markdown(f"<p style='text-align: center; color: #666;'>Generated: {datetime.now().strftime('%B %d, %Y')}</p>", unsafe_allow_html=True)

        # Key metrics
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>📊 Key Metrics</h3></div>", unsafe_allow_html=True)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Companies Analyzed", len(companies))
        with col2:
            total_mcap = sum(c.get('marketCap', 0) for c in companies)
            st.metric("Total Market Cap", format_currency(total_mcap))
        with col3:
            avg_mcap = total_mcap / len(companies) if companies else 0
            st.metric("Avg Market Cap", format_currency(avg_mcap))
        with col4:
            st.metric("Sector P/E", f"{sector_data.get('pe', 'N/A')}")

        # Industry Overview
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>📋 Industry Overview</h3></div>", unsafe_allow_html=True)
        if ai_analysis.get('overview'):
            st.markdown(ai_analysis['overview'])

        # Top Companies Table
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>🏢 Top Companies</h3></div>", unsafe_allow_html=True)

        company_df = pd.DataFrame([
            {
                'Rank': i + 1,
                'Symbol': c.get('symbol', 'N/A'),
                'Company': c.get('companyName', 'N/A')[:40],
                'Market Cap': format_currency(c.get('marketCap', 0)),
                'Price': f"${c.get('price', 0):.2f}",
                'Beta': f"{c.get('beta', 0):.2f}" if c.get('beta') else "N/A"
            }
            for i, c in enumerate(companies[:15])
        ])
        st.dataframe(company_df, use_container_width=True, hide_index=True)

        # Key Trends
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>📈 Key Industry Trends</h3></div>", unsafe_allow_html=True)
        if ai_analysis.get('trends'):
            st.markdown(ai_analysis['trends'])

        # Winners & Losers
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>🎯 Winners & Losers from Trends</h3></div>", unsafe_allow_html=True)
        display_winners_losers(winners_losers)

        # Industry Risks
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>⚠️ Industry Risks</h3></div>", unsafe_allow_html=True)
        if ai_analysis.get('risks'):
            st.markdown(ai_analysis['risks'])

        # 12-Month Outlook
        st.markdown("<div class='section-header'><h3 style='margin:0; color: white;'>🔮 12-Month Outlook</h3></div>", unsafe_allow_html=True)
        if ai_analysis.get('outlook'):
            st.markdown(ai_analysis['outlook'])

        # Footer
        st.markdown("---")
        st.markdown(f"<p style='text-align: center; color: #999; font-size: 0.8em;'>Report generated by Industry Report Generator | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
