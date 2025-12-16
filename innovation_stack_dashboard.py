"""
Innovation Stack Dashboard
Interactive Streamlit dashboard for AI-powered article summaries from top innovation sources
"""
import streamlit as st
import feedparser
from datetime import datetime
import pytz
import os
import re
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from openai import OpenAI
import anthropic
from docx import Document
from docx.shared import Pt
import io

load_dotenv()

st.set_page_config(
    page_title="Innovation Stack",
    page_icon="💡",
    layout="wide"
)

# API Keys
def get_api_key(key_name):
    key = os.getenv(key_name)
    if key:
        return key
    try:
        return st.secrets.get(key_name)
    except:
        return None

OPENAI_API_KEY = get_api_key("OPENAI_API_KEY")
ANTHROPIC_API_KEY = get_api_key("ANTHROPIC_API_KEY")

# RSS Feeds
FEEDS = {
    "Stratechery": "https://stratechery.com/feed/",
    "Not Boring": "https://www.notboring.co/feed",
    "Lenny's Newsletter": "https://www.lennysnewsletter.com/feed",
    "Paul Graham": "http://www.aaronsw.com/2002/feeds/pgessays.rss",
    "Seth Godin": "https://seths.blog/feed/atom/",
    "MIT Tech Review": "https://www.technologyreview.com/feed/",
    "TechCrunch": "https://techcrunch.com/feed/",
}


def fetch_article_content(url):
    """Fetch and extract main content from an article URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            element.decompose()

        article = soup.find('article') or soup.find('main') or soup.find('div', class_=re.compile('content|article|post'))
        text = article.get_text(separator=' ', strip=True) if article else soup.get_text(separator=' ', strip=True)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:5000] if len(text) > 5000 else text
    except:
        return None


def summarize_article(title, content, use_claude=True):
    """Summarize article with AI."""
    if not content:
        return "Unable to fetch article content."

    prompt = f"""Summarize this article in ONE concise paragraph (3-5 sentences). Focus on key insights and actionable takeaways.

Article Title: {title}
Content: {content}

Provide only the summary paragraph."""

    if use_claude and ANTHROPIC_API_KEY:
        try:
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            response = client.messages.create(
                model="claude-3-5-haiku-latest",
                max_tokens=300,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text.strip()
        except:
            pass

    if OPENAI_API_KEY:
        try:
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300
            )
            return response.choices[0].message.content.strip()
        except:
            pass

    return "Summary unavailable."


def fetch_feed_entries(selected_feeds, articles_per_feed):
    """Fetch entries from selected RSS feeds."""
    all_entries = []

    for name in selected_feeds:
        url = FEEDS.get(name)
        if not url:
            continue

        try:
            feed = feedparser.parse(url)
            for entry in feed.entries[:articles_per_feed]:
                if hasattr(entry, "published_parsed") and entry.published_parsed:
                    pub_date = datetime(*entry.published_parsed[:6])
                elif hasattr(entry, "updated_parsed") and entry.updated_parsed:
                    pub_date = datetime(*entry.updated_parsed[:6])
                else:
                    pub_date = datetime.now()

                pub_date = pub_date.replace(tzinfo=pytz.UTC).astimezone(pytz.timezone("America/New_York"))

                all_entries.append({
                    "date": pub_date,
                    "source": name,
                    "title": entry.title,
                    "link": entry.link,
                    "summary": None
                })
        except Exception as e:
            st.warning(f"Error fetching {name}: {e}")

    all_entries.sort(key=lambda x: x["date"], reverse=True)
    return all_entries


def create_word_doc(entries):
    """Create Word document with summaries."""
    doc = Document()
    doc.add_heading("Daily Innovation Stack", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph()

    for i, entry in enumerate(entries, 1):
        doc.add_heading(f"{i}. [{entry['source']}] {entry['title']}", level=2)
        if entry['summary']:
            doc.add_paragraph(entry['summary'])
        doc.add_paragraph(f"Link: {entry['link']}")
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Main App
st.title("💡 Innovation Stack Dashboard")
st.markdown("AI-powered article summaries from top innovation sources")

# Check API keys
if not OPENAI_API_KEY and not ANTHROPIC_API_KEY:
    st.error("Missing API keys. Add OPENAI_API_KEY or ANTHROPIC_API_KEY to your secrets.")
    st.stop()

# Sidebar
st.sidebar.header("Settings")
selected_feeds = st.sidebar.multiselect(
    "Select Sources",
    options=list(FEEDS.keys()),
    default=list(FEEDS.keys())[:4]
)
articles_per_feed = st.sidebar.slider("Articles per source", 1, 5, 3)
use_claude = st.sidebar.checkbox("Use Claude (faster)", value=True)

if st.sidebar.button("🚀 Generate Innovation Stack", type="primary"):
    if not selected_feeds:
        st.error("Please select at least one source")
        st.stop()

    # Fetch articles
    with st.spinner("Fetching articles from RSS feeds..."):
        entries = fetch_feed_entries(selected_feeds, articles_per_feed)

    if not entries:
        st.error("No articles found")
        st.stop()

    st.success(f"Found {len(entries)} articles")

    # Generate summaries
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, entry in enumerate(entries):
        status_text.text(f"Summarizing: {entry['title'][:50]}...")
        content = fetch_article_content(entry['link'])
        entry['summary'] = summarize_article(entry['title'], content, use_claude)
        progress_bar.progress((i + 1) / len(entries))

    status_text.text("Done!")

    # Display results
    st.header("📰 Today's Innovation Stack")

    for i, entry in enumerate(entries, 1):
        with st.expander(f"**{i}. [{entry['source']}]** {entry['title']}", expanded=(i <= 3)):
            st.write(f"📅 {entry['date'].strftime('%B %d, %Y')}")
            st.write("---")
            st.write(entry['summary'])
            st.markdown(f"[Read full article →]({entry['link']})")

    # Download
    st.divider()
    doc = create_word_doc(entries)
    st.download_button(
        "📥 Download Innovation Stack (Word)",
        doc,
        file_name=f"Innovation_Stack_{datetime.now().strftime('%Y-%m-%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

else:
    st.info("👈 Select sources and click 'Generate Innovation Stack' to start")

    st.markdown("""
    ### Sources Available:
    - **Stratechery** - Ben Thompson's tech strategy analysis
    - **Not Boring** - Packy McCormick's business & tech deep dives
    - **Lenny's Newsletter** - Product management insights
    - **Paul Graham** - Startup wisdom from YC founder
    - **Seth Godin** - Marketing & leadership thoughts
    - **MIT Tech Review** - Emerging technology coverage
    - **TechCrunch** - Startup & tech news
    """)
